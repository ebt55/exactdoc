"""Narrow OOXML contracts for the Google Docs writer profile."""
import base64
import tempfile
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document

from exactdoc.docxout import WriteCtx, write_docx, write_figure, write_para, write_table
from exactdoc.infer import para_from_lines
from exactdoc.layout import (Cell, Chunk, DocLayout, FigureEl, PageLayout,
                             Para, Run, TableEl)
from exactdoc.model import Line, Span


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _xml(docx_path):
    with zipfile.ZipFile(docx_path) as z:
        return ET.fromstring(z.read("word/document.xml"))


def _run(text, size=12.0):
    return Run(text=text, font="Helvetica", size=size, color="#000000")


def _cover_layout():
    band = TableEl(
        rows=[[Cell(paras=[Para(runs=[_run("Cover title")], space_before=20.0,
                                 left_indent=30.0),
                            Para(runs=[_run("Cover subtitle")], space_before=4.0)],
                    pad=(0.0, 30.0, 0.0, 0.0))]],
        col_widths=[604.0], row_heights=[150.0], role="band",
    )
    return DocLayout(cover_band=band, cover_top=0.0,
                     pages=[PageLayout(1, [Chunk(elements=[Para(runs=[_run("body")])])])])


def _first_table_values(docx_path):
    root = _xml(docx_path)
    tbl = root.find(".//" + W + "tbl")
    tc = tbl.find(".//" + W + "tc")
    mar_left = tc.find(".//" + W + "tcMar/" + W + "left").get(W + "w")
    par = tc.find(W + "p")
    spacing = par.find(".//" + W + "spacing")
    ind = par.find(".//" + W + "ind")
    return mar_left, spacing.get(W + "before"), None if ind is None else ind.get(W + "left")


class GoogleDocsWriterFixes(unittest.TestCase):
    def test_gdocs_cover_moves_left_padding_and_compensates_first_before(self):
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            standard = tmp_path / "standard.docx"
            gdocs = tmp_path / "gdocs.docx"
            write_docx(_cover_layout(), str(standard), output_profile="standard")
            write_docx(_cover_layout(), str(gdocs), output_profile="gdocs")

            # Standard preserves the original cell-padding form.  Google Docs receives
            # the same horizontal position as paragraph indentation and 290 fewer twips
            # only on the first real cover paragraph.
            self.assertEqual(_first_table_values(standard), ("1960", "400", None))
            self.assertEqual(_first_table_values(gdocs), ("0", "110", "1960"))


    def test_gdocs_does_not_move_padding_for_ordinary_tables(self):
        with tempfile.TemporaryDirectory() as td:
            table = TableEl(rows=[[Cell(paras=[Para(runs=[_run("ordinary")], space_before=20.0)],
                                        pad=(0.0, 30.0, 0.0, 0.0))]], col_widths=[200.0])
            doc = Document()
            write_table(doc, table, 200.0, ctx=WriteCtx(output_profile="gdocs"))
            path = Path(td) / "ordinary.docx"
            doc.save(path)
            self.assertEqual(_first_table_values(path), ("600", "400", None))

    def test_gdocs_cover_retains_a_source_indent_beyond_cell_padding(self):
        table = TableEl(
            rows=[[Cell(paras=[Para(runs=[_run("offset cover")], left_indent=48.0)],
                        pad=(0.0, 30.0, 0.0, 0.0))]],
            col_widths=[200.0], role="band",
        )
        with tempfile.TemporaryDirectory() as td:
            doc = Document()
            write_table(doc, table, 200.0, ctx=WriteCtx(output_profile="gdocs"),
                        cover_band=True)
            path = Path(td) / "offset-cover.docx"
            doc.save(path)
            _, _, indent = _first_table_values(path)
            self.assertEqual(indent, "960")


_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4z8DwHwAFgAI/ScL9xQAAAABJRU5ErkJggg=="
)


def _figure_spacing(tmp_path, profile):
    doc = Document()
    fig = FigureEl(page_no=1, clip=(0, 0, 1, 1), width=20, height=40)
    write_figure(doc, fig, ctx=WriteCtx(output_profile=profile,
                                         render_clip=lambda *_: _PNG))
    path = tmp_path / (profile + ".docx")
    doc.save(path)
    return _xml(path).find(".//" + W + "pPr/" + W + "spacing")


def _line(text, width, baseline):
    span = Span(text=text, font="Helvetica", size=12.0, color="#000000",
                bold=False, italic=False, mono=False, serif=False,
                superscript=False, bbox=(0.0, baseline - 10, width, baseline),
                origin=(0.0, baseline))
    return Line(spans=[span], bbox=span.bbox)


class GoogleDocsInferenceFixes(unittest.TestCase):
    def test_gdocs_inline_figure_omits_redundant_atleast_height(self):
        with tempfile.TemporaryDirectory() as td:
            tmp_path = Path(td)
            standard = _figure_spacing(tmp_path, "standard")
            gdocs = _figure_spacing(tmp_path, "gdocs")
            self.assertEqual(standard.get(W + "lineRule"), "atLeast")
            self.assertEqual(standard.get(W + "line"), "800")
            self.assertIsNone(gdocs.get(W + "lineRule"))
            self.assertIsNone(gdocs.get(W + "line"))

    def test_metadata_rows_are_preserved_only_when_inset_is_impossible(self):
        metadata = para_from_lines([
            _line("From: Operations", 77.8, 100),
            _line("To: All team leads", 80.0, 115),
            _line("Date: July 21, 2026", 86.2, 130),
        ], 0.0, 500.0)
        # Inference retains the standard flowing form and records a Google-only
        # row representation.  The writer is where that alternate form applies.
        self.assertEqual(metadata.align, "justify")
        self.assertEqual(metadata.right_indent, 420.0)
        self.assertFalse(metadata.line_breaks)
        self.assertEqual(len(metadata.gdocs_rows), 3)
        with tempfile.TemporaryDirectory() as td:
            standard_doc = Document()
            write_para(standard_doc, metadata, 500.0,
                       ctx=WriteCtx(output_profile="standard"))
            standard_path = Path(td) / "metadata-standard.docx"
            standard_doc.save(standard_path)
            standard_xml = _xml(standard_path)
            standard_ppr = standard_xml.find(".//" + W + "pPr")
            self.assertEqual(standard_ppr.find(W + "jc").get(W + "val"), "both")
            self.assertEqual(standard_ppr.find(W + "ind").get(W + "right"), "8400")
            self.assertEqual(len(standard_xml.findall(".//" + W + "br")), 0)

            gdocs_doc = Document()
            write_para(gdocs_doc, metadata, 500.0, ctx=WriteCtx(output_profile="gdocs"))
            gdocs_path = Path(td) / "metadata-gdocs.docx"
            gdocs_doc.save(gdocs_path)
            gdocs_xml = _xml(gdocs_path)
            gdocs_ppr = gdocs_xml.find(".//" + W + "pPr")
            self.assertEqual(gdocs_ppr.find(W + "jc").get(W + "val"), "left")
            self.assertIsNone(gdocs_ppr.find(W + "ind"))
            self.assertEqual(len(gdocs_xml.findall(".//" + W + "br")), 2)

        ordinary = para_from_lines([
            _line("A justified line", 100.0, 100),
            _line("Another justified line", 100.0, 115),
            _line("ragged last line", 60.0, 130),
        ], 0.0, 500.0)
        self.assertEqual(ordinary.align, "justify")
        self.assertEqual(ordinary.right_indent, 400.0)
        self.assertFalse(ordinary.line_breaks)
        self.assertNotIn("\n", ordinary.text)
