"""The refusal contract, and the image degradation that must never become one.

Two failure modes the expansion corpus exposed, held down from opposite ends:

  - An image exactdoc cannot embed must degrade, never crash. A 126-page
    document is not allowed to die of two JPEGs.
  - An input exactdoc cannot honestly convert must be refused with a typed
    error and a stable exit code, never converted into a plausible-looking
    wrong answer.

Synthetic PDFs are built with PyMuPDF rather than reportlab, which is not
installed on every developer machine this suite runs on. The real corpus
fixtures are used where the *evidence* is the point -- a threshold argued from
invented documents would be a threshold argued from nothing.
"""
import io
import os
from pathlib import Path
import subprocess
import sys
import tempfile
import unittest

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from exactdoc import errors                                     # noqa: E402
from exactdoc.cli import EXIT_CODES                             # noqa: E402
from exactdoc.convert import convert                            # noqa: E402
from exactdoc.docxout import _embeddable, _to_png, write_docx   # noqa: E402
from exactdoc.errors import (ConfigurationError, InteractiveFormError,  # noqa: E402
                             PageLimitError, ResourceLimitError,
                             UnsupportedInputError)
from exactdoc.layout import Chunk, DocLayout, ImageEl, PageLayout  # noqa: E402
from exactdoc.options import RAW                                # noqa: E402
from exactdoc.scan import (FORM_PAGE_SHARE, FORM_PAGE_WIDGETS,  # noqa: E402
                           MAX_PAGES_PER_DOCUMENT, ScanReport, is_form,
                           page_cap, refusal)

FIXTURES = ROOT / "testkit" / "fixtures"
EXPANSION = ROOT / "testkit" / "fixtures_expansion"


def _have(module):
    import importlib.util
    return importlib.util.find_spec(module) is not None


def _adobe_app14_jpeg():
    """A valid JPEG with an Adobe APP14 marker -- the shape that used to crash.

    `FF D8 FF EE` instead of JFIF's `FF D8 FF E0`. python-docx matches signatures
    against the first 32 bytes and has no entry for this, so it refuses a file
    every other decoder on the machine reads without comment. Antenna House and
    the Adobe toolchain emit it; two are in y06_irs_1040_instructions.pdf.
    """
    from PIL import Image
    buf = io.BytesIO()
    # Pillow writes the Adobe APP14 segment for CMYK/YCCK JPEGs.
    Image.new("CMYK", (24, 16), (0, 0, 0, 40)).save(buf, format="JPEG")
    data = buf.getvalue()
    assert data[:4] == b"\xff\xd8\xff\xee", data[:4].hex()
    return data


def _jfif_jpeg():
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (24, 16), (200, 30, 30)).save(buf, format="JPEG")
    data = buf.getvalue()
    assert data[:4] == b"\xff\xd8\xff\xe0", data[:4].hex()
    return data


def _layout_with_images(*images):
    """One page, one chunk, one ImageEl per blob -- the writer's real shape."""
    elements = [ImageEl(data=d, ext="jpeg", width=48, height=32)
                for d in images]
    return DocLayout(pages=[PageLayout(1, chunks=[Chunk(elements=elements)])])


def _synthetic_pdf(path, pages=1, widgets_per_page=0, text="synthetic page"):
    """A small PDF with a controlled number of interactive form widgets."""
    import fitz
    doc = fitz.open()
    for _ in range(pages):
        page = doc.new_page(width=612, height=792)
        page.insert_text((72, 72), text, fontsize=11)
        for i in range(widgets_per_page):
            w = fitz.Widget()
            w.field_type = fitz.PDF_WIDGET_TYPE_TEXT
            w.field_name = "f%d" % i
            w.rect = fitz.Rect(72, 120 + i * 18, 300, 134 + i * 18)
            page.add_widget(w)
    doc.save(str(path))
    doc.close()
    return str(path)


# --------------------------------------------------------------- the crash
@unittest.skipUnless(_have("PIL"), "Pillow is required to build the fixtures")
class UnembeddableImageTests(unittest.TestCase):
    """FAILURE 1: one bad image must cost one image, not the document."""

    def test_adobe_app14_jpeg_is_reencoded_not_dropped(self):
        report = {}
        data = _embeddable(_adobe_app14_jpeg(), report)
        self.assertIsNotNone(data, "a JPEG Pillow reads must never be dropped")
        self.assertEqual(report, {"reencoded": 1})
        self.assertEqual(data[:8], b"\x89PNG\r\n\x1a\n")

    def test_embeddable_bytes_pass_through_untouched(self):
        report = {}
        original = _jfif_jpeg()
        self.assertEqual(_embeddable(original, report), original)
        self.assertEqual(report, {"embedded": 1})

    def test_unreadable_bytes_are_dropped_and_recorded(self):
        for label, blob in (("garbage", b"\x00\x01\x02not an image at all"),
                            ("empty", b""),
                            ("truncated-jpeg", _jfif_jpeg()[:12])):
            with self.subTest(blob=label):
                report = {}
                self.assertIsNone(_embeddable(blob, report))
                self.assertEqual(report, {"dropped": 1})

    def test_to_png_returns_none_rather_than_raising(self):
        self.assertIsNone(_to_png(b"definitely not an image"))

    def test_write_docx_survives_a_bad_image_and_keeps_the_good_one(self):
        """The whole point: the document is written, and says what it lost."""
        report = {}
        lay = _layout_with_images(_jfif_jpeg(), b"\x00 unreadable",
                                  _adobe_app14_jpeg())
        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "mixed.docx")
            write_docx(lay, out, image_report=report)
            self.assertTrue(os.path.getsize(out) > 0)
            from docx import Document
            # Two images survived; the unreadable one is absent, not a crash.
            self.assertEqual(len(Document(out).inline_shapes), 2)
        self.assertEqual(report, {"embedded": 1, "dropped": 1, "reencoded": 1})

    def test_the_ledger_is_per_write_not_cumulative(self):
        """The refine loop writes the same layout once per round."""
        report = {}
        lay = _layout_with_images(b"\x00 unreadable")
        with tempfile.TemporaryDirectory() as td:
            for rnd in range(3):
                write_docx(lay, os.path.join(td, "r%d.docx" % rnd),
                           image_report=report)
        self.assertEqual(report, {"dropped": 1},
                         "three rounds must not report three dropped images")

    @unittest.skipUnless((EXPANSION / "y06_irs_1040_instructions.pdf").exists(),
                         "y06 expansion fixture is not present")
    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_y06_the_regression_document_converts_end_to_end(self):
        source = EXPANSION / "y06_irs_1040_instructions.pdf"
        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "y06.docx")
            convert(str(source), out, options=RAW)
            from docx import Document
            document = Document(out)
            # 126 pages of text used to be lost to two images.
            self.assertGreater(sum(len(p.text) for p in document.paragraphs),
                               100000)
            self.assertGreater(len(document.inline_shapes), 0)

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_every_crashing_corpus_document_recovers_losslessly(self):
        """All four documents that raised UnrecognizedImageError, by ledger.

        The crash was measured on four: three Antenna House and one Adobe PDF
        Library 11.0, which is why the fix keys on the image bytes rather than
        the producer. y08 is the discriminating one -- it carries accepted JFIF
        and PNG images *alongside* a rejected Adobe-APP14 JPEG, so a single write
        exercises the pass-through and the re-encode paths together.

        Nothing here may be dropped: every one of these is a JPEG Pillow reads.
        A drop would mean the re-encode ladder regressed to the old behaviour
        with the crash merely swallowed.
        """
        from exactdoc.backend import get_backend
        from exactdoc.dialect import normalize
        from exactdoc.infer import infer
        from exactdoc.input import parse as parse_input
        backend = get_backend("pymupdf")
        expected = {"y06_irs_1040_instructions": {"reencoded": 1},
                    "y08_nist_sp80088r1": {"embedded": 3, "reencoded": 1},
                    "y12_irs_pub15": {"embedded": 1, "reencoded": 2},
                    "y13_irs_pub501": {"reencoded": 2}}
        with tempfile.TemporaryDirectory() as td:
            for name, ledger in expected.items():
                source = EXPANSION / (name + ".pdf")
                if not source.exists():
                    continue
                with self.subTest(source=name):
                    report = {}
                    lay = infer(normalize(parse_input(backend, str(source))))
                    out = os.path.join(td, name + ".docx")
                    write_docx(lay, out, dpi=RAW.dpi, backend=backend,
                               output_profile=RAW.output_profile,
                               image_report=report)
                    self.assertEqual(report, ledger)
                    self.assertNotIn("dropped", report)
                    self.assertTrue(os.path.getsize(out) > 0)


# ---------------------------------------------------------- form detection
class FormThresholdTests(unittest.TestCase):
    """FAILURE 2a: the threshold, and the documents that must NOT trip it."""

    def test_threshold_constants_are_the_ones_argued_in_scan(self):
        self.assertEqual((FORM_PAGE_WIDGETS, FORM_PAGE_SHARE), (12, 0.10))

    def test_negative_cases_a_few_widgets_are_not_a_form(self):
        cases = {"no widgets at all": [0] * 8,
                 "one signature field": [1],
                 "a signature page in a contract": [0, 0, 4],
                 # y03 NIST FIPS 197: 16 stray widgets over 46 pages.
                 "stray annotations": [2, 1, 2] + [0] * 43,
                 # A dense form page buried in a long booklet is not the booklet.
                 "one worksheet in a manual": [0] * 125 + [30]}
        for label, widgets in cases.items():
            with self.subTest(case=label):
                self.assertFalse(is_form(widgets))

    def test_positive_cases_the_form_dominates(self):
        cases = {"IRS 1040 (y07)": [128, 71],
                 "IRS W-9 (y14): one form page, five of instructions": [23, 0, 0, 0, 0, 0],
                 "an application form": [20, 18]}
        for label, widgets in cases.items():
            with self.subTest(case=label):
                self.assertTrue(is_form(widgets))

    def test_an_unavailable_census_is_not_read_as_zero_widgets(self):
        self.assertFalse(is_form(None))
        self.assertFalse(ScanReport("digital", 2, 100).census_available)

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_synthetic_one_widget_letter_converts(self):
        """A letter with a single signature widget must not be refused."""
        with tempfile.TemporaryDirectory() as td:
            source = _synthetic_pdf(Path(td) / "letter.pdf", pages=1,
                                    widgets_per_page=1)
            out = os.path.join(td, "letter.docx")
            convert(source, out, options=RAW)
            self.assertTrue(os.path.exists(out))

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_synthetic_dense_form_is_refused_before_any_output(self):
        with tempfile.TemporaryDirectory() as td:
            source = _synthetic_pdf(Path(td) / "form.pdf", pages=2,
                                    widgets_per_page=FORM_PAGE_WIDGETS)
            out = os.path.join(td, "form.docx")
            with self.assertRaises(InteractiveFormError) as raised:
                convert(source, out, options=RAW)
            self.assertEqual(raised.exception.code, "interactive-form")
            self.assertFalse(os.path.exists(out),
                             "a refusal must not leave an output behind")

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_both_backends_count_the_same_widgets(self):
        """A refusal that depends on the parser is not a contract."""
        if not _have("pypdfium2"):
            self.skipTest("PDFium is an optional candidate backend")
        from exactdoc.backend import get_backend
        with tempfile.TemporaryDirectory() as td:
            source = _synthetic_pdf(Path(td) / "form.pdf", pages=3,
                                    widgets_per_page=5)
            counts = [get_backend(n).form_widgets(source)
                      for n in ("pymupdf", "pdfium")]
            self.assertEqual(counts[0], counts[1])
            self.assertEqual(counts[0], [5, 5, 5])


class RealFormFixtureTests(unittest.TestCase):
    """The threshold against the documents it was calibrated on."""

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_real_forms_refuse_and_real_documents_do_not(self):
        from exactdoc.backend import get_backend
        backend = get_backend("pymupdf")
        cases = {"y07_irs_f1040_form.pdf": True,
                 "y14_irs_fw9_form.pdf": True,
                 "y03_nist_fips197.pdf": False,
                 "y06_irs_1040_instructions.pdf": False}
        for name, expected in cases.items():
            source = EXPANSION / name
            if not source.exists():
                continue
            with self.subTest(source=name):
                self.assertEqual(is_form(backend.form_widgets(str(source))),
                                 expected)

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_the_gated_corpus_contains_no_form_and_nothing_over_cap(self):
        """The 16 metric inputs must be untouched by either refusal.

        If a gated fixture ever tripped one of these, every published number
        would silently start describing a smaller corpus.
        """
        from exactdoc.backend import get_backend
        backend = get_backend("pymupdf")
        sources = sorted(FIXTURES.glob("*.pdf"))
        self.assertEqual(len(sources), 16, "the gate corpus is 16 documents")
        for source in sources:
            with self.subTest(source=source.name):
                widgets = backend.form_widgets(str(source))
                self.assertFalse(is_form(widgets))
                self.assertLessEqual(len(widgets), MAX_PAGES_PER_DOCUMENT)


# --------------------------------------------------------------- page cap
class PageCapTests(unittest.TestCase):
    """FAILURE 2b: the same bound whether the document arrives alone or not."""

    def test_the_single_file_cap_is_the_batch_cap(self):
        from exactdoc import batch
        self.assertEqual(batch.MAX_PAGES_PER_DOCUMENT, MAX_PAGES_PER_DOCUMENT)

    def test_cap_resolution(self):
        self.assertEqual(page_cap(None), MAX_PAGES_PER_DOCUMENT)
        self.assertEqual(page_cap(600), 600)
        self.assertIsNone(page_cap(0), "0 means no cap")

    def test_refusal_is_typed_and_lifted_by_an_explicit_override(self):
        report = ScanReport("digital", MAX_PAGES_PER_DOCUMENT + 1, 5000)
        error = refusal(report)
        self.assertIsInstance(error, PageLimitError)
        self.assertIsInstance(error, ResourceLimitError)
        self.assertEqual(error.code, "page-limit")
        self.assertIn("--max-pages", error.message)
        self.assertIsNone(refusal(report, max_pages=600))
        self.assertIsNone(refusal(report, max_pages=0))

    def test_a_document_at_the_cap_is_not_refused(self):
        self.assertIsNone(refusal(ScanReport("digital", MAX_PAGES_PER_DOCUMENT, 10)))

    def test_max_pages_is_validated_where_every_other_option_is(self):
        for bad in (-1, 2.5, "many", True):
            with self.subTest(value=bad):
                with self.assertRaises(ConfigurationError):
                    RAW.replace(max_pages=bad)
        self.assertEqual(RAW.replace(max_pages=0).max_pages, 0)

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_synthetic_over_cap_refuses_then_converts_under_override(self):
        with tempfile.TemporaryDirectory() as td:
            source = _synthetic_pdf(Path(td) / "long.pdf", pages=6)
            out = os.path.join(td, "long.docx")
            tiny = RAW.replace(max_pages=5)
            with self.assertRaises(PageLimitError) as raised:
                convert(source, out, options=tiny)
            self.assertIn("6 pages", raised.exception.message)
            self.assertFalse(os.path.exists(out))
            convert(source, out, options=RAW.replace(max_pages=6))
            self.assertTrue(os.path.exists(out))

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_the_cap_is_enforced_before_the_parse(self):
        """A bound spent after a 492-page extraction has protected nothing."""
        from unittest import mock
        with tempfile.TemporaryDirectory() as td:
            source = _synthetic_pdf(Path(td) / "long.pdf", pages=4)
            with mock.patch("exactdoc.convert.parse_input",
                            side_effect=AssertionError("parsed before refusing")):
                with self.assertRaises(PageLimitError):
                    convert(source, os.path.join(td, "o.docx"),
                            options=RAW.replace(max_pages=3))


# --------------------------------------------------------------- the CLI
class ExitCodeTests(unittest.TestCase):
    def test_every_error_class_has_an_exit_code(self):
        """A new error falling through to 1 looks like an ordinary failure.

        errors.py has promised this test exists for some time; it did not.
        """
        # ExactdocError itself is the generic root and maps to 1 by design.
        missing = sorted(set(errors.BY_CODE) - set(EXIT_CODES) - {"error"})
        self.assertEqual(missing, [], "error codes with no exit code")
        # batch-partial is a CLI outcome, not an exception class.
        unknown = sorted(set(EXIT_CODES) - set(errors.BY_CODE) - {"batch-partial"})
        self.assertEqual(unknown, [], "exit codes with no error class")

    def test_the_new_codes_extend_the_numbering_without_renumbering(self):
        self.assertEqual(EXIT_CODES["ocr-required"], 17)
        self.assertEqual(EXIT_CODES["batch-partial"], 18)
        self.assertEqual(EXIT_CODES["interactive-form"], 19)
        self.assertEqual(EXIT_CODES["page-limit"], 20)
        self.assertEqual(len(set(EXIT_CODES.values())), len(EXIT_CODES))

    def test_the_form_error_is_an_unsupported_input(self):
        self.assertTrue(issubclass(InteractiveFormError, UnsupportedInputError))
        self.assertTrue(issubclass(PageLimitError, ResourceLimitError))

    def _cli(self, *args):
        return subprocess.run([sys.executable, "-m", "exactdoc.cli"] + list(args),
                              cwd=str(ROOT), text=True, capture_output=True)

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_cli_exit_19_for_a_form_and_20_over_the_page_cap(self):
        with tempfile.TemporaryDirectory() as td:
            form = _synthetic_pdf(Path(td) / "form.pdf", pages=2,
                                  widgets_per_page=FORM_PAGE_WIDGETS)
            long_doc = _synthetic_pdf(Path(td) / "long.pdf", pages=4)
            out = os.path.join(td, "out.docx")
            for source, code, extra in ((form, 19, []),
                                        (long_doc, 20, ["--max-pages", "3"])):
                with self.subTest(exit_code=code):
                    proc = self._cli(source, "--refine", "0", "-o", out, *extra)
                    self.assertEqual(proc.returncode, code, proc.stderr)
                    self.assertIn("error:", proc.stderr)
                    self.assertNotIn("Traceback", proc.stderr)
                    self.assertFalse(os.path.exists(out))
            # The override is what makes the cap a decision rather than a wall.
            proc = self._cli(long_doc, "--refine", "0", "-o", out,
                             "--max-pages", "0")
            self.assertEqual(proc.returncode, 0, proc.stderr)

    @unittest.skipUnless(_have("fitz"), "PyMuPDF is the measured backend")
    def test_scan_only_reports_both_conditions_without_converting(self):
        with tempfile.TemporaryDirectory() as td:
            # Dense enough to be a form AND longer than the cap we pass.
            source = _synthetic_pdf(Path(td) / "both.pdf", pages=4,
                                    widgets_per_page=FORM_PAGE_WIDGETS)
            proc = self._cli(source, "--scan-only", "--max-pages", "3")
            self.assertEqual(proc.returncode, 19, proc.stderr)
            self.assertIn("form", proc.stdout)
            self.assertIn("pages: 4", proc.stdout)
            self.assertIn("over the page cap", proc.stdout,
                          "the page cap must be reported even when the form "
                          "refusal is the one that exits")
            self.assertIn("form widgets: 48", proc.stdout)
            self.assertEqual(
                [n for n in os.listdir(td) if n.endswith(".docx")], [],
                "--scan-only must never write a document")

    def test_max_pages_is_rejected_for_batch(self):
        proc = self._cli("--input-dir", "x", "--out-dir", "y", "--max-pages", "9")
        self.assertEqual(proc.returncode, 2)
        self.assertIn("--max-pages is not supported for batch", proc.stderr)


if __name__ == "__main__":
    unittest.main()
