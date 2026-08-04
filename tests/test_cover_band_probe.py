"""Contracts for the cover-band floor probe.

Run with ``python tests/test_cover_band_probe.py``.  Hermetic: the analyser is
exercised against a PDF built in-process, so no renderer, no network, and no
Google export is needed to know the probe still reads geometry correctly.
"""
import os
import sys
import tempfile
import unittest

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, "testkit"))

import probe_cover_band as P  # noqa: E402

try:
    import fitz
except ImportError:  # pragma: no cover - optional-backend installs
    fitz = None

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class MarkerTests(unittest.TestCase):
    def test_marker_round_trips_including_a_fractional_point(self):
        for family, value in P.VARIANTS:
            self.assertEqual(P.parse_marker(P._marker(family, value)),
                             (family, value))

    def test_the_fractional_variant_survives_encoding(self):
        # 14.4 is the whole point of the probe; a marker scheme that rounded it
        # to 14 would answer a different question than the one being asked.
        self.assertEqual(P._marker("TOP", 14.4), "PROBETOP0144")
        self.assertEqual(P.parse_marker("PROBETOP0144"), ("TOP", 14.4))

    def test_unmarked_text_is_not_invented_into_a_variant(self):
        for junk in ("", None, "ordinary text", "PROBE", "PROBEXXX0000"):
            self.assertIsNone(P.parse_marker(junk))


class BuildTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.path = os.path.join(self.tmp.name, "probe.docx")
        P.build(self.path)
        from docx import Document
        self.doc = Document(self.path)

    def tearDown(self):
        self.tmp.cleanup()

    def test_one_section_and_one_band_per_variant(self):
        self.assertEqual(len(self.doc.sections), len(P.VARIANTS))
        self.assertEqual(len(self.doc.tables), len(P.VARIANTS))

    def test_each_section_carries_only_its_own_variable(self):
        for sec, (family, value) in zip(self.doc.sections, P.VARIANTS):
            top = sec.top_margin.pt
            side = sec.left_margin.pt
            hdr = sec.header_distance.pt
            self.assertAlmostEqual(sec.right_margin.pt, side, places=1)
            if family == "TOP":
                self.assertAlmostEqual(top, value, places=1)
                self.assertAlmostEqual(side, 4.0, places=1)
                self.assertAlmostEqual(hdr, 0.0, places=1)
            elif family == "HDR":
                self.assertAlmostEqual(top, 0.0, places=1)
                self.assertAlmostEqual(side, 4.0, places=1)
                self.assertAlmostEqual(hdr, value, places=1)
            else:
                self.assertAlmostEqual(top, 0.0, places=1)
                self.assertAlmostEqual(side, value, places=1)
                self.assertAlmostEqual(hdr, 0.0, places=1)

    def test_every_band_is_shaded_and_labelled_with_its_variant(self):
        seen = []
        for tbl, (family, value) in zip(self.doc.tables, P.VARIANTS):
            cell = tbl.cell(0, 0)
            shd = cell._tc.find(W + "tcPr").find(W + "shd")
            self.assertIsNotNone(shd, "band has no shading to measure")
            self.assertEqual(shd.get(W + "fill"), P.BAND_FILL)
            seen.append(P.parse_marker(cell.text))
        self.assertEqual(seen, list(P.VARIANTS))

    def test_no_content_precedes_the_first_band(self):
        """A stray leading paragraph would invalidate the 0pt variant.

        python-docx starts every document with an empty paragraph. Left in
        place it sits above the first band, so the one variant that asks for a
        zero top offset would silently measure the paragraph's height instead.
        """
        body = self.doc.element.body
        first = None
        for child in body:
            if child.tag in (W + "p", W + "tbl"):
                first = child
                break
        self.assertEqual(first.tag, W + "tbl")

    def test_band_width_spans_the_page_less_its_margins(self):
        for tbl, (family, value) in zip(self.doc.tables, P.VARIANTS):
            side = value if family == "SIDE" else 4.0
            want = int(round((P.PAGE_W - 2 * side) * 20))
            tblW = tbl._tbl.tblPr.find(W + "tblW")
            self.assertEqual(int(tblW.get(W + "w")), want)


@unittest.skipIf(fitz is None, "PyMuPDF not installed")
class AnalyseTests(unittest.TestCase):
    def _pdf(self, path, bands):
        """bands: [(marker|None, x0, y0, x1, y1, colour)] one page each."""
        doc = fitz.open()
        for marker, x0, y0, x1, y1, colour in bands:
            page = doc.new_page(width=P.PAGE_W, height=P.PAGE_H)
            if x1 > x0:
                page.draw_rect(fitz.Rect(x0, y0, x1, y1), color=colour,
                               fill=colour)
            if marker:
                page.insert_text((x0 + 20, y0 + 40), marker, fontsize=18,
                                 color=(1, 1, 1))
        doc.save(path)
        doc.close()

    def test_reads_band_geometry_and_computes_the_right_delta(self):
        dark = (0.12, 0.22, 0.39)
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            self._pdf(path, [
                (P._marker("TOP", 14.4), 4.0, 20.0, 608.0, 170.0, dark),
                (P._marker("SIDE", 0.0), 3.9, 0.0, 606.8, 150.0, dark),
            ])
            rows = P.analyse(path)
            self.assertEqual(len(rows), 2)

            top = rows[0]
            self.assertEqual((top["family"], top["requested_pt"]), ("TOP", 14.4))
            self.assertAlmostEqual(top["top"], 20.0, places=1)
            self.assertAlmostEqual(top["height"], 150.0, places=1)
            # asked 14.4, rendered 20.0 -> Docs held it 5.6pt lower
            self.assertAlmostEqual(top["delta"], 5.6, places=1)

            side = rows[1]
            self.assertEqual(side["family"], "SIDE")
            # a SIDE variant is judged on its left edge, not its top
            self.assertAlmostEqual(side["left"], 3.9, places=1)
            self.assertAlmostEqual(side["delta"], 3.9, places=1)

    def test_a_pale_fill_is_not_mistaken_for_the_band(self):
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            self._pdf(path, [(P._marker("TOP", 0.0), 4.0, 30.0, 608.0, 180.0,
                              (0.97, 0.97, 0.99))])
            rows = P.analyse(path)
            self.assertNotIn("top", rows[0])
            self.assertIn("no band rectangle found", P.render(rows))

    def test_missing_marker_or_band_degrades_without_raising(self):
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            self._pdf(path, [(None, 4.0, 10.0, 608.0, 160.0, (0.1, 0.2, 0.4))])
            rows = P.analyse(path)
            self.assertIsNone(rows[0]["family"])
            self.assertIn("no probe marker found", P.render(rows))

    def test_render_names_the_lowest_top_achieved(self):
        dark = (0.12, 0.22, 0.39)
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            # every requested top clamps to the same 14.4pt floor
            self._pdf(path, [
                (P._marker("TOP", 0.0), 4.0, 14.4, 608.0, 164.4, dark),
                (P._marker("TOP", 4.0), 4.0, 14.4, 608.0, 164.4, dark),
                (P._marker("TOP", 20.0), 4.0, 20.0, 608.0, 170.0, dark),
            ])
            text = P.render(P.analyse(path))
            self.assertIn("lowest band top achieved: 14.40pt", text)
            self.assertIn("20.0", text)


if __name__ == "__main__":
    unittest.main()
