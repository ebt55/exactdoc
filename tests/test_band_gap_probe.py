"""Contracts for the band-to-body gap probe.

Run with ``python tests/test_band_gap_probe.py``.  Hermetic: the analyser is
exercised against PDFs built in-process.
"""
import os
import sys
import tempfile
import unittest

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, "testkit"))

import probe_band_gap as P  # noqa: E402

try:
    import fitz
except ImportError:  # pragma: no cover
    fitz = None

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class DesignTests(unittest.TestCase):
    def test_marker_round_trips(self):
        for i in range(len(P.VARIANTS)):
            self.assertEqual(P.parse_marker(P.marker(i)), i)
        for junk in ("", None, "BGPROBE", "nothing"):
            self.assertIsNone(P.parse_marker(junk))

    def test_each_family_varies_exactly_one_thing(self):
        """One variable per page, or the residual cannot be attributed."""
        tc = [v for v in P.VARIANTS if v[0] == "tcmar"]
        self.assertEqual(len({v[1] for v in tc}), len(tc))   # tcMar varies
        self.assertEqual(len({v[2] for v in tc}), 1)         # before fixed
        sp = [v for v in P.VARIANTS if v[0] == "space"]
        self.assertEqual(len({v[2] for v in sp}), len(sp))   # before varies
        self.assertEqual(len({v[1] for v in sp}), 1)         # tcMar fixed

    def test_there_is_a_no_table_control(self):
        # without it, a residual at the boundary cannot be blamed on the table
        self.assertTrue(any(not v[5] for v in P.VARIANTS))

    def test_labels_are_distinct(self):
        labels = [P.label(i) for i in range(len(P.VARIANTS))]
        self.assertEqual(len(set(labels)), len(labels))


class BuildTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.path = os.path.join(self.tmp.name, "p.docx")
        P.build(self.path)
        from docx import Document
        self.doc = Document(self.path)

    def tearDown(self):
        self.tmp.cleanup()

    def test_one_section_per_variant_and_a_band_for_each_banded_one(self):
        self.assertEqual(len(self.doc.sections), len(P.VARIANTS))
        self.assertEqual(len(self.doc.tables),
                         sum(1 for v in P.VARIANTS if v[5]))

    def test_every_page_carries_both_measurement_anchors(self):
        text = "\n".join(p.text for p in self.doc.paragraphs)
        for tbl in self.doc.tables:
            text += "\n" + tbl.cell(0, 0).text
        self.assertEqual(text.count(P.BAND_TEXT), len(P.VARIANTS))
        self.assertEqual(text.count(P.BODY_TEXT), len(P.VARIANTS))

    def test_requested_tcmar_bottom_reaches_the_xml(self):
        banded = [v for v in P.VARIANTS if v[5]]
        for tbl, spec in zip(self.doc.tables, banded):
            tcPr = tbl.cell(0, 0)._tc.find(W + "tcPr")
            bottom = tcPr.find(W + "tcMar").find(W + "bottom")
            self.assertEqual(int(bottom.get(W + "w")), int(round(spec[1] * 20)))


@unittest.skipIf(fitz is None, "PyMuPDF not installed")
class AnalyseTests(unittest.TestCase):
    def _pdf(self, path, idx, band_bottom=200.0, band_base=180.0,
             body_base=260.0, band_top=40.0, with_band=True):
        doc = fitz.open()
        page = doc.new_page(width=P.PAGE_W, height=P.PAGE_H)
        page.insert_text((P.MARGIN, 20), "%s  %s" % (P.marker(idx),
                                                     P.label(idx)), fontsize=8)
        if with_band:
            page.draw_rect(fitz.Rect(P.MARGIN, band_top, P.PAGE_W - P.MARGIN,
                                     band_bottom),
                           color=(0.12, 0.22, 0.39), fill=(0.12, 0.22, 0.39))
        page.insert_text((P.MARGIN, band_base), P.BAND_TEXT,
                         fontsize=P.BAND_SIZE, color=(1, 1, 1))
        page.insert_text((P.MARGIN, body_base), P.BODY_TEXT,
                         fontsize=P.BODY_SIZE)
        doc.save(path)
        doc.close()

    def test_reads_both_baselines_and_the_band_rectangle(self):
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            self._pdf(path, 1)          # tcMar=12 before=24
            r = P.analyse(path)[0]
            self.assertEqual(r["family"], "tcmar")
            self.assertAlmostEqual(r["band_bottom"], 200.0, places=1)
            self.assertAlmostEqual(r["below_line"], 20.0, places=1)
            self.assertAlmostEqual(r["gap"], 60.0, places=1)
            # residuals subtract what was requested plus the line's own metric
            self.assertAlmostEqual(
                r["resid_bottom"], 20.0 - (12.0 + r["band_desc"]), places=2)
            self.assertAlmostEqual(
                r["resid_gap"], 60.0 - (24.0 + r["body_asc"]), places=2)

    def test_a_larger_rendered_gap_shows_up_as_a_larger_residual(self):
        with tempfile.TemporaryDirectory() as work:
            a = os.path.join(work, "a.pdf")
            b = os.path.join(work, "b.pdf")
            self._pdf(a, 1, body_base=260.0)
            self._pdf(b, 1, body_base=265.0)
            ra, rb = P.analyse(a)[0], P.analyse(b)[0]
            self.assertAlmostEqual(rb["resid_gap"] - ra["resid_gap"], 5.0,
                                   places=1)

    def test_the_no_table_control_still_measures_a_gap(self):
        idx = [i for i, v in enumerate(P.VARIANTS) if not v[5]][0]
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            self._pdf(path, idx, with_band=False)
            r = P.analyse(path)[0]
            self.assertFalse(r["has_band"])
            self.assertNotIn("band_bottom", r)
            # falls back to the band line's own descent as the anchor
            self.assertIn("resid_gap", r)

    def test_missing_marker_degrades_without_raising(self):
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            doc = fitz.open()
            doc.new_page(width=P.PAGE_W, height=P.PAGE_H)
            doc.save(path)
            doc.close()
            rows = P.analyse(path)
            self.assertIsNone(rows[0]["variant"])
            self.assertIn("no probe marker found", P.render(rows))

    def test_render_contrasts_the_table_and_no_table_cases(self):
        with tempfile.TemporaryDirectory() as work:
            rows = []
            for i in (1, [j for j, v in enumerate(P.VARIANTS) if not v[5]][0]):
                path = os.path.join(work, "p%d.pdf" % i)
                self._pdf(path, i, with_band=P.VARIANTS[i][5])
                rows.extend(P.analyse(path))
            text = P.render(rows)
            self.assertIn("table present", text)
            self.assertIn("no table", text)


if __name__ == "__main__":
    unittest.main()
