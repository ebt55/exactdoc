"""Contracts for the Docs font-metrics probe.

Run with ``python tests/test_font_metrics_probe.py``.  Hermetic: the analyser is
exercised against PDFs built in-process, so the probe stays testable without a
renderer, a network, or a Google export.
"""
import os
import sys
import tempfile
import unittest

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)
sys.path.insert(0, os.path.join(ROOT, "testkit"))

import probe_font_metrics as P  # noqa: E402

from exactdoc.fonts import GDOCS_NATIVE, METRIC_REFERENCE  # noqa: E402

try:
    import fitz
except ImportError:  # pragma: no cover - optional-backend installs
    fitz = None

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
FONTS = r"C:\Windows\Fonts"


class ReferenceTests(unittest.TestCase):
    def test_the_copied_reference_string_has_not_drifted(self):
        """The probe deliberately does not import the converter.

        That copy is only safe while something checks it, because a probe
        measuring a different character set than fonts.FAMILY_METRICS was built
        from would produce numbers that look comparable and are not.
        """
        self.assertEqual(P.METRIC_REFERENCE, METRIC_REFERENCE)

    def test_every_candidate_is_a_family_docs_is_said_to_have(self):
        for fam in P.CANDIDATES:
            self.assertIn(fam.lower(), GDOCS_NATIVE, fam)

    def test_the_controls_are_families_measurable_offline(self):
        for fam in P.KNOWN_ADVANCE:
            self.assertIn(fam, P.CANDIDATES)

    def test_segments_fit_one_line_and_preserve_the_reference_words(self):
        segs = P.segments()
        self.assertTrue(segs)
        for s in segs:
            self.assertLessEqual(len(s), P.MAX_SEG_CHARS)
            self.assertFalse(s.startswith(" ") or s.endswith(" "))
        self.assertEqual(" ".join(segs).split(), METRIC_REFERENCE.split())

    def test_marker_round_trips_and_junk_does_not(self):
        for i in range(len(P.CANDIDATES)):
            self.assertEqual(P.parse_marker(P.marker(i)), i)
        for junk in ("", None, "FMPROBE", "no marker here"):
            self.assertIsNone(P.parse_marker(junk))


class BuildTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.path = os.path.join(self.tmp.name, "probe.docx")
        P.build(self.path)
        from docx import Document
        self.doc = Document(self.path)

    def tearDown(self):
        self.tmp.cleanup()

    def test_one_section_per_candidate(self):
        self.assertEqual(len(self.doc.sections), len(P.CANDIDATES))

    def test_each_page_declares_its_family_and_carries_every_segment(self):
        segs = P.segments()
        paras = [p for p in self.doc.paragraphs if p.text.strip()]
        # marker paragraph + body paragraph per family
        self.assertEqual(len(paras), 2 * len(P.CANDIDATES))
        for i, family in enumerate(P.CANDIDATES):
            head, body = paras[2 * i], paras[2 * i + 1]
            self.assertEqual(P.parse_marker(head.text), i)
            self.assertIn(family, head.text)
            for seg in segs:
                self.assertIn(seg, body.text)

    def test_measurement_runs_request_the_candidate_family(self):
        paras = [p for p in self.doc.paragraphs if p.text.strip()]
        for i, family in enumerate(P.CANDIDATES):
            body = paras[2 * i + 1]
            asciis = {rf.get(W + "ascii")
                      for rf in body._p.iter(W + "rFonts")}
            self.assertEqual(asciis, {family}, family)

    def test_the_marker_uses_a_family_google_certainly_has(self):
        # if the candidate is missing the page must still identify itself
        paras = [p for p in self.doc.paragraphs if p.text.strip()]
        for i in range(len(P.CANDIDATES)):
            asciis = {rf.get(W + "ascii")
                      for rf in paras[2 * i]._p.iter(W + "rFonts")}
            self.assertEqual(asciis, {"Arial"})

    def test_measurement_lines_carry_no_paragraph_spacing(self):
        """Pitch must be the family's natural line height and nothing else."""
        paras = [p for p in self.doc.paragraphs if p.text.strip()]
        for i in range(len(P.CANDIDATES)):
            pf = paras[2 * i + 1].paragraph_format
            self.assertEqual(pf.space_before.pt, 0)
            self.assertEqual(pf.space_after.pt, 0)
            self.assertEqual(pf.line_spacing, 1.0)


@unittest.skipIf(fitz is None, "PyMuPDF not installed")
class AnalyseTests(unittest.TestCase):
    def _pdf(self, path, pages, pitch=13.0, size=10.0,
             fontfile="DejaVuSerif.ttf"):
        """pages: [(index|None, fontfile|None)] -- one probe page each."""
        segs = P.segments()
        doc = fitz.open()
        for idx, ff in pages:
            page = doc.new_page(width=P.PAGE_W, height=P.PAGE_H)
            if idx is not None:
                page.insert_text((P.MARGIN, 40), "%s  %s"
                                 % (P.marker(idx), P.CANDIDATES[idx]),
                                 fontsize=9)
            src = os.path.join(FONTS, ff or fontfile)
            if not os.path.exists(src):
                continue
            page.insert_font(fontname="probe", fontfile=src)
            for j, seg in enumerate(segs):
                page.insert_text((P.MARGIN, 80 + j * pitch), seg,
                                 fontname="probe", fontsize=size)
        doc.save(path)
        doc.close()

    def test_measures_advance_and_pitch_against_the_dejavu_target(self):
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            # index 0 is Noto Serif, but the page is drawn in DejaVu Serif
            self._pdf(path, [(0, "DejaVuSerif.ttf")], pitch=13.0, size=10.0)
            rows = P.analyse(path)
            self.assertEqual(len(rows), 1)
            r = rows[0]
            self.assertEqual(r["family"], "Noto Serif")
            self.assertEqual(r["lines_found"], len(P.segments()))
            # drawn in the target face, so the ratio must land on 1.0
            self.assertAlmostEqual(r["ratio"], 1.0, places=2)
            self.assertAlmostEqual(r["factor"], 13.0 / 10.0, places=2)

    def test_a_face_that_is_not_the_requested_family_is_flagged(self):
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            self._pdf(path, [(0, "DejaVuSerif.ttf")])
            r = P.analyse(path)[0]
            self.assertTrue(r["substituted"])
            # and a flagged row must never win the decision, however good it looks
            verdict = P.decide([r])
            self.assertEqual(verdict["action"], "inconclusive")

    def test_pages_without_a_marker_or_without_text_degrade_quietly(self):
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e.pdf")
            self._pdf(path, [(None, "DejaVuSerif.ttf")])
            rows = P.analyse(path)
            self.assertIsNone(rows[0]["family"])
            self.assertIn("no probe marker found", P.render(rows))
        with tempfile.TemporaryDirectory() as work:
            path = os.path.join(work, "e2.pdf")
            doc = fitz.open()
            page = doc.new_page(width=P.PAGE_W, height=P.PAGE_H)
            page.insert_text((P.MARGIN, 40), P.marker(3), fontsize=9)
            doc.save(path)
            doc.close()
            rows = P.analyse(path)
            self.assertEqual(rows[0]["family"], P.CANDIDATES[3])
            self.assertNotIn("ratio", rows[0])
            self.assertIn("no measurement lines matched", P.render(rows))


class DecisionRuleTests(unittest.TestCase):
    def _row(self, family, ratio, factor=1.30, substituted=False):
        return {"page": 1, "family": family, "ratio": ratio, "factor": factor,
                "advance": ratio * P.TARGET_ADVANCE, "face": family,
                "substituted": substituted, "lines_found": 6}

    def _noto(self):
        return self._row("Noto Serif",
                         P.KNOWN_ADVANCE["Noto Serif"] / P.TARGET_ADVANCE,
                         factor=1.360)

    def test_a_family_inside_the_adopt_bar_is_adopted(self):
        v = P.decide([self._noto(), self._row("Merriweather", 0.99)])
        self.assertEqual(v["action"], "adopt")
        self.assertEqual(v["best"]["family"], "Merriweather")

    def test_a_family_that_halves_the_residual_is_adopted(self):
        # 3% deviation: outside the 2% bar, but less than half of Noto's 6.2%
        v = P.decide([self._noto(), self._row("Bitter", 0.97)])
        self.assertEqual(v["action"], "adopt")
        self.assertIn("halves", v["reason"])

    def test_a_marginal_gain_does_not_justify_changing_typeface(self):
        v = P.decide([self._noto(), self._row("Lora", 0.96)])
        self.assertEqual(v["action"], "marginal")

    def test_nothing_better_than_noto_serif_keeps_noto_serif(self):
        v = P.decide([self._noto(), self._row("PT Serif", 0.90)])
        self.assertEqual(v["action"], "keep-noto-serif")

    def test_no_family_close_enough_closes_the_book(self):
        """The outcome that ends the search instead of prolonging it."""
        v = P.decide([self._row("Merriweather", 0.95),
                      self._row("Bitter", 0.93)])
        self.assertEqual(v["action"], "close-the-book")
        self.assertIn("waiver", v["reason"])

    def test_a_candidate_without_a_measured_line_height_cannot_win(self):
        """The guard the last substitution earned the hard way.

        Adopting a family on advance width alone, without its natural line
        height, is what rendered l1 at a 17.48pt pitch against a 14.70pt
        source. A candidate whose pitch did not come back is not adoptable.
        """
        v = P.decide([self._noto(),
                      self._row("Merriweather", 1.0, factor=None)])
        self.assertEqual(v["action"], "keep-noto-serif")

    def test_a_substituted_candidate_cannot_win(self):
        v = P.decide([self._noto(),
                      self._row("Merriweather", 1.0, substituted=True)])
        self.assertEqual(v["action"], "keep-noto-serif")

    def test_no_usable_measurement_is_inconclusive_not_a_decision(self):
        v = P.decide([self._row("Merriweather", 1.0, substituted=True)])
        self.assertEqual(v["action"], "inconclusive")
        self.assertIsNone(v["best"])

    def test_thresholds_are_ordered(self):
        self.assertLess(P.ADOPT_DEV, P.CLOSE_BOOK_DEV)
        self.assertLess(P.NOTO_DEV * P.IMPROVE_FACTOR, P.CLOSE_BOOK_DEV)
        self.assertGreater(P.NOTO_DEV, P.ADOPT_DEV)

    def test_render_names_the_control_disagreement(self):
        rows = [self._noto()]
        rows[0]["advance"] = P.KNOWN_ADVANCE["Noto Serif"]
        text = P.render(rows, P.decide(rows))
        self.assertIn("control check", text)
        self.assertIn("Noto Serif", text)


if __name__ == "__main__":
    unittest.main()
