"""DocLayout -> .docx writer (Google Docs-safe vocabulary).

Constructs used: styled paragraphs, fixed-layout tables with per-side borders
and cell shading, section geometry + true column sections, inline images,
headers/footers with PAGE/NUMPAGES fields, hyperlinks, tab stops.
No floating text boxes, no embedded fonts, no VML.
"""
import copy
import dataclasses
import io
import re
from typing import Any, Callable, Dict, Optional, List

from docx import Document
from docx.shared import Pt, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .layout import (DocLayout, Para, Run, Cell, TableEl, FigureEl, ImageEl,
                     RuleEl, ColBreak, HFPart)
from .fonts import map_font
from .metrics import source_line_width


@dataclasses.dataclass(frozen=True)
class WriteCtx:
    """Everything a write needs to know that is not in the DocLayout.

    This replaces a module global. `LINE_MODE` was set by `write_docx` and
    restored in a `finally`, so two conversions running concurrently with
    different targets could each observe the other's line-height encoding --
    silently, and only in the overlap. A frozen object passed down the call tree
    cannot do that.

    `render_clip(page_no, clip, dpi) -> png bytes | None` is how a figure region
    reaches the writer. It used to be an open MuPDF document handed down five
    call levels, which is what put `import fitz` at the top of this module and
    made a wheel without PyMuPDF fail while *importing the writer* -- before any
    backend selection could happen.
    """

    line_mode: str = "exact"
    dpi: int = 240
    render_clip: Optional[Callable] = None
    # Keep the output profile as an explicit writer concern.  `line_mode` is
    # intentionally a lossy rendering choice (another profile may choose the
    # same encoding), while a handful of Google Docs workarounds really are
    # profile-specific.  It comes last to preserve the old positional shape.
    output_profile: str = "standard"
    # Internal-link plumbing, filled in by _write_docx once it has planned the
    # bookmarks. Appended after output_profile so the positional shape above is
    # still the old one. {LinkDest: anchor name} and {anchor name: w:id}.
    dest_anchors: Dict[Any, str] = dataclasses.field(default_factory=dict)
    anchor_ids: Dict[str, int] = dataclasses.field(default_factory=dict)
    #: Optional mutable tally of what happened to each extracted raster:
    #: ``{"embedded": n, "reencoded": n, "dropped": n}``.  A degradation nobody
    #: can observe is indistinguishable from a lie, and dropping an image is a
    #: degradation.  Defaults to None -- the writer keeps no global ledger, so
    #: two concurrent conversions cannot accumulate into each other's counts.
    image_report: Optional[dict] = None


_DEFAULT_CTX = WriteCtx()

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
TABAL = {"left": WD_TAB_ALIGNMENT.LEFT, "center": WD_TAB_ALIGNMENT.CENTER,
         "right": WD_TAB_ALIGNMENT.RIGHT}


def _hex(c: str) -> str:
    return (c or "#000000").lstrip("#").upper()


def _set_borders(el_pr, borders: dict, tag: str):
    """Apply border dict to tcPr/pPr. tag='w:tcBorders' or 'w:pBdr'."""
    bel = OxmlElement(tag)
    order = ("top", "left", "bottom", "right")
    for side in order:
        spec = borders.get(side)
        b = OxmlElement("w:" + side)
        if spec:
            w, color = spec[0], spec[1]
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(max(2, int(round(w * 8)))))
            b.set(qn("w:space"), str(int(spec[2])) if len(spec) > 2 else "0")
            b.set(qn("w:color"), _hex(color))
        else:
            b.set(qn("w:val"), "nil")
        bel.append(b)
    el_pr.append(bel)


def _style_run(r, run: Run):
    f = r.font
    fam = map_font(run.font, mono=run.mono, serif=run.serif)
    f.name = fam
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), fam)
    f.size = Pt(round(run.size * 2) / 2)
    f.bold = run.bold
    f.italic = run.italic
    if run.underline:
        f.underline = True
    if run.superscript:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)
    if abs(getattr(run, "char_spacing", 0.0)) > 0.004:
        # w:spacing on rPr is character tracking, in twentieths of a point
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(round(run.char_spacing * 20))))
        rpr.append(sp)
    try:
        f.color.rgb = RGBColor.from_string(_hex(run.color))
    except Exception:
        pass


def _add_field(par, instr: str, sample: str, style_from: Run):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " %s " % instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = sample
    r.append(t)
    fld.append(r)
    par._p.append(fld)
    # style the inner run
    from docx.text.run import Run as DRun
    dr = DRun(r, par)
    _style_run(dr, style_from)


def _add_hyperlink(par, url: str, runs_and_styles):
    part = par.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    par._p.append(h)
    from docx.text.run import Run as DRun
    for text, style in runs_and_styles:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        h.append(r)
        _style_run(DRun(r, par), style)


def _add_internal_hyperlink(par, anchor: str, runs_and_styles):
    """A link to a bookmark in this document: w:hyperlink w:anchor.

    Deliberately built the same way as _add_hyperlink rather than through
    python-docx's helper, and deliberately WITHOUT a w:rStyle: the Hyperlink
    character style would repaint the text blue and underline it, and the source
    span already carries the styling the producer chose. c8_toc_links sets its
    table of contents in #123a5e with `text-decoration: none`, so borrowing
    Word's link styling would visibly recolour text that is not blue in the PDF.
    _style_run applies the run's own formatting, exactly as for external links.
    """
    h = OxmlElement("w:hyperlink")
    h.set(qn("w:anchor"), anchor)
    par._p.append(h)
    from docx.text.run import Run as DRun
    for text, style in runs_and_styles:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        h.append(r)
        _style_run(DRun(r, par), style)


def _bookmark_pair(name: str, bid: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    return start, end


def _wrap_paragraph_bookmark(par, name: str, bid: int):
    """Bracket a paragraph's content with a bookmark, inside the w:p.

    bookmarkStart has to follow w:pPr -- pPr must be the first child of w:p --
    so this inserts after it rather than at index 0.
    """
    start, end = _bookmark_pair(name, bid)
    p = par._p
    ppr = p.find(qn("w:pPr"))
    p.insert(list(p).index(ppr) + 1 if ppr is not None else 0, start)
    p.append(end)


def _add_block_bookmark(container, name: str, bid: int):
    """A zero-height bookmark between block elements.

    Used when a destination resolves to a table, figure, image or rule rather
    than a paragraph. bookmarkStart/End are range markers and are legal as
    direct children of w:body, so this costs no paragraph and therefore no
    vertical space -- which is the whole reason internal links can be added
    without moving a single measured number.
    """
    start, end = _bookmark_pair(name, bid)
    container.element.body.append(start)
    container.element.body.append(end)


def _el_extent(el):
    """(top, bottom) of a flow element in source page points, or None."""
    if isinstance(el, Para):
        bb = el.bbox
    elif isinstance(el, TableEl):
        bb = el.bbox
    elif isinstance(el, FigureEl):
        bb = el.clip
    else:
        bb = getattr(el, "_bbox", None)
    return (bb[1], bb[3]) if bb else None


def _iter_runs(el):
    if isinstance(el, Para):
        for r in el.runs:
            yield r
        for row in getattr(el, "gdocs_rows", None) or []:
            for r in row:
                yield r
    elif isinstance(el, TableEl):
        for row in el.rows:
            for cell in row:
                if not cell:
                    continue
                for p in cell.paras:
                    for r in p.runs:
                        yield r


def _anchor_name(dest) -> str:
    """Deterministic, collision-free, and a legal Word bookmark name.

    Word bookmark names must begin with a letter and may contain only letters,
    digits and underscores. Two destinations that resolve to the same point get
    the same name on purpose -- they are the same anchor -- and the hundredths
    of a point keep two genuinely different points apart.
    """
    return "exactdoc_dest_p%d_%d" % (int(dest.page), int(round(dest.y * 100)))


def _plan_bookmarks(lay: DocLayout):
    """Resolve every referenced destination to a flow element and name it.

    THE ANCHORING RULE, in order:

      1. the element whose vertical extent CONTAINS the destination y;
      2. otherwise the element whose top edge is nearest at-or-below it;
      3. otherwise (the destination sits below all content) the last element.

    Ties are broken by flow order, earliest first, so the result does not
    depend on dictionary or sort stability.

    Step 1 is not decoration. A /XYZ destination names the point that should
    come to the top of the window, and producers put it at the target's top
    edge -- which lands a hair INSIDE the element once font ascent is taken into
    account, not above it. Measured on a ReportLab file whose destination is the
    baseline of the text it names, "nearest at-or-below" alone skipped that text
    and anchored to the following paragraph; containment gets it right.

    Anchoring is by y only. A destination's x is recorded in the IR but says
    nothing about which paragraph is meant -- /XYZ's `left` is a horizontal
    scroll position, and on this corpus it is 0 for every destination.
    """
    wanted = set()
    for pg in lay.pages:
        for ch in pg.chunks:
            for el in ch.elements:
                for run in _iter_runs(el):
                    if run.dest is not None:
                        wanted.add(run.dest)
    if not wanted:
        return {}, {}

    per_page = {}
    for pi, pg in enumerate(lay.pages):
        seq = []
        for ch in pg.chunks:
            for el in ch.elements:
                ext = _el_extent(el)
                if ext is not None:
                    seq.append((ext[0], ext[1], el))
        per_page[pi] = seq

    dest_anchors, targets, named_element = {}, {}, {}
    for dest in sorted(wanted, key=lambda d: (d.page, d.y, d.x)):
        seq = per_page.get(int(dest.page)) or []
        if not seq:
            continue
        hit = next((el for top, bot, el in seq if top <= dest.y <= bot), None)
        if hit is None:
            below = [(top, el) for top, _, el in seq if top >= dest.y]
            hit = min(below, key=lambda t: t[0])[1] if below else seq[-1][2]
        # ONE bookmark per element, whatever the destination was called. Two
        # destinations a few points apart routinely land on the same paragraph;
        # minting a name each would leave the element carrying only the last of
        # them, and every other anchor pointing at a bookmark never written.
        name = named_element.get(id(hit))
        if name is None:
            name = _anchor_name(dest)
            named_element[id(hit)] = name
            targets[name] = hit
        dest_anchors[dest] = name

    anchor_ids = {name: i + 1 for i, name in enumerate(sorted(targets))}
    for name, el in targets.items():
        el._bookmark = name
    return dest_anchors, anchor_ids


# Half-point wrap correction: implemented and measured, OFF by default.
#
# It works -- line-break agreement on the WeasyPrint sample goes 0.599 -> 0.796,
# and a parameter sweep puts the optimum narrowing at ~1%, exactly the
# 10.0/10.1 size ratio. But restoring the correct line breaks makes paragraphs
# their correct (taller) height, and the writer currently has no way to keep a
# page from overflowing, so that sample goes 10 -> 11 pages and word placement
# drops 0.985 -> 0.706. Net regression, so it stays off.
#
# Turning it on is blocked on overflow control (the closed-loop second pass):
# render, find pages that spilled, shrink discretionary space, re-emit. Enable
# both together, never this alone.
WRAP_CORRECTION = False


# --- line-height encoding, per render target ------------------------------
# Word and LibreOffice honour w:spacing lineRule="exact" literally, and the
# whole fidelity model (THEORY 3.1) is built on it: paragraph height is
# n_lines x leading, exactly.
#
# Google Docs has no "exact" line spacing in its own document model -- only
# multiples -- so its importer must translate, and the translation is wrong in
# a way that scales with font size. Measured (testkit/docs_quirks.py, three
# lines per probe, error vs LibreOffice):
#
#     size/leading      exact      atLeast    multiple
#     10pt / 12.0pt     -2.0pt     -2.0pt     --
#     18pt / 21.0pt    +45.2pt     -1.4pt     -0.3pt
#     22pt / 25.5pt    +84.3pt     -1.1pt     -0.6pt
#
# So a heading with exact leading gains ~28pt in Docs, which is precisely the
# "+28pt after the first heading" that the closed loop had been compensating
# per document. Emitting the same intent as a multiple makes it a static fix.
#
# Natural line height as a fraction of font size, per family, measured in
# Google Docs (testkit/docs_quirks.py h5: four bare 20pt lines, no spacing
# properties; factor = rendered_gap / (4 x 20)). The gdocs translation divides
# by this, so a single constant silently drifts on any document set in a font
# whose metric differs -- Roboto is 4% taller than Arial, which is exactly the
# kind of quiet assumption this project exists to avoid.
NATURAL_FACTORS = {
    "arial": 1.144, "times new roman": 1.144, "courier new": 1.127,
    "georgia": 1.130, "roboto": 1.194,
    # Added when the metric fit began substituting these families. Docs' live
    # pass 2 rendered l1_word_native in Noto Serif at a 17.48pt pitch where the
    # source used 14.70pt: dividing by the 1.144 default inflated every line by
    # 19%, which is the whole of that document's remaining dy. Deriving the
    # factor back out of that export gives 1.144 * 17.48 / 14.70 = 1.360.
    #
    # The font files agree and explain the whole table. With
    #     (hhea.ascender - hhea.descender + hhea.lineGap) / unitsPerEm
    # Arial reads 1.150, Times New Roman 1.150, Courier New 1.133 and Georgia
    # 1.136 -- each exactly 0.006 above its Docs-measured value here, a constant
    # offset across four independently probed families. Noto Serif reads 1.362
    # by the same formula, so 1.356 predicted against 1.360 observed.
    "noto serif": 1.360, "noto sans": 1.356, "verdana": 1.209,
    # Also measured inside Docs rather than from a font file, by
    # testkit/probe_font_metrics.py in live pass 3 -- the family is not
    # installed here. The probe's own controls recovered Noto Serif at 1.362,
    # Times New Roman at 1.150 and Georgia at 1.136 against the 1.360/1.144/
    # 1.130 in this table, so a pitch read this way is good to about 0.006.
    "libre baskerville": 1.240,
}
NATURAL_DEFAULT = 1.144
# The two encodings. Which one is used is a per-write decision carried in
# WriteCtx.line_mode, not a module global -- see WriteCtx.
LINE_MODES = ("exact", "multiple")


def line_mode_for(output_profile: str) -> str:
    """Word and LibreOffice honour lineRule="exact"; Google Docs mistranslates it.

    Keyed on the OUTPUT PROFILE, not on which renderer the refinement loop talks
    to. Those were one field, so "write OOXML that survives Google Docs" was
    inseparable from "upload this document to Google" -- and the offline
    Docs-safe profile, which is what this project intends to ship, could not be
    expressed at all.
    """
    return "multiple" if output_profile == "gdocs" else "exact"

# Floor on compressing a table row's leading to make it fit its source height.
# Below this the text starts to collide with its neighbours, and an honestly
# too-tall row beats an unreadable one.
MIN_ROW_SHRINK = 0.55

# A continuous section break is carried by a real paragraph, and that paragraph
# occupies flow height the source page never spent. The "crush" further down
# used to be described as making it consume none; it does not.
#
# Measured on c2_paper2col through the canonical LibreOffice by putting marker
# text in the paragraph and reading the render back: its advance is exactly its
# w:line, with a floor near 1pt. Raising w:line from 20 twips to 400 moved every
# following line down by 19.00pt -- exactly the 19pt added -- so the height is
# honoured, and 1pt is the least it can be made to cost. Dropping it to 1 twip
# changed nothing, which is the floor showing. Folding the sectPr into the
# preceding paragraph instead does NOT work: that collapses the paragraph's own
# exact height, which cost 13.7pt on c2.
#
# So it cannot be removed, only accounted for -- the gap hoisted ahead of the
# break is emitted that much shorter.
SECT_BREAK_PARA_TWIPS = 20
SECT_BREAK_PARA_PT = SECT_BREAK_PARA_TWIPS / 20.0


def _sect_break_comp(ctx) -> float:
    """How much of the section-break paragraph's height to pre-subtract.

    Deliberately gdocs-only, and scoped by output profile rather than by
    whether the correction loop runs.

    The standard profile ships with the LibreOffice loop (refine3), and that
    loop has already absorbed this 1pt empirically: correcting it at source as
    well makes the loop over-correct, and it was measured doing so --
    c2's product-lane mean_ssim fell 0.824 -> 0.793 and its dy_p50 went
    0.85 -> 1.75. The recorded 0.85 encodes the loop cancelling this bug, and
    unbundling that is not worth a shipping regression.

    The gdocs profile ships with no loop at all (refine0) and exists precisely
    as the static-translation layer for corrections a loop cannot supply there,
    so this is exactly the kind of correction it is for. Scoping on the output
    profile rather than on refine_rounds also leaves the raw gate lane alone --
    raw runs the standard profile open-loop -- so neither gate lane moves.
    """
    return SECT_BREAK_PARA_PT if getattr(ctx, "output_profile", "") == "gdocs" \
        else 0.0


def _natural_factor(family: str) -> float:
    return NATURAL_FACTORS.get((family or "").lower(), NATURAL_DEFAULT)


def _apply_leading(pf, leading: float, size: float, mode: str = "exact",
                   family: str = ""):
    """Encode a line height the way the chosen target actually honours."""
    if mode == "multiple" and size and size > 0.5 and leading > 1.0:
        natural = size * _natural_factor(family)
        pf.line_spacing = max(0.06, leading / natural)   # w:line as a multiple
        return
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(round(leading, 1))


def _quantised_size(sz: float) -> float:
    """OOXML stores font size in half-points (w:sz is in half-points, integer)."""
    return round(sz * 2) / 2


def _wrap_correction(p: Para, content_w: float) -> float:
    """Extra right indent that cancels the half-point font quantisation.

    A 10.1pt source font can only be emitted at 10.0 or 10.5. At 10.0 the
    glyphs are 1% narrow, so ~1% more text fits per line and the paragraph
    re-wraps; at 9.5 (from 9.7) they are 2% narrow the other way. Either way
    the line breaks move, and every paragraph below shifts.

    Line breaking is scale-invariant: shrink every advance by k and the same
    breaks return if the wrap width also shrinks by k. So narrow the column by
    (1 - emitted/source). This is exact for a uniform-size paragraph and a good
    approximation for mixed runs, where the dominant size is used.

    Only applied where it can help and cannot hurt: the paragraph must actually
    wrap (a single-line paragraph has no breaks to preserve) and must be
    left/justified (on centred or right-aligned text a right indent *moves* the
    text instead of only changing where it wraps).
    """
    if not WRAP_CORRECTION:
        return 0.0
    if p.align not in ("left", "justify") or not p.runs:
        return 0.0
    weight = {}
    for r in p.runs:
        if r.text and not r.is_tab:
            weight[r.size] = weight.get(r.size, 0) + len(r.text)
    if not weight:
        return 0.0
    src = max(weight, key=weight.get)
    if src < 1.0:
        return 0.0
    k = _quantised_size(src) / src
    if abs(k - 1.0) < 1e-6:
        return 0.0
    wrap_w = max(1.0, content_w - p.left_indent - p.right_indent)
    # a paragraph that fits on one line has no wrap to preserve
    if p.bbox is not None and (p.bbox[2] - p.bbox[0]) < 0.92 * wrap_w and \
            not p.line_breaks:
        est_lines = sum(len(r.text) for r in p.runs) * 0.5 * src / max(1.0, wrap_w)
        if est_lines < 1.2:
            return 0.0
    return wrap_w * (1.0 - k)


def write_para(container, p: Para, content_w: float, par=None, ctx=None,
               space_before: Optional[float] = None):
    """Write a Para into container (doc/cell/header). Returns the paragraph.

    `space_before` overrides the paragraph's own gap for this write only. It is
    how `_absorb_page_spill` spends a page's slack without touching the layout
    -- see the note below on why nothing here may be mutated.
    """
    ctx = ctx or _DEFAULT_CTX
    if par is None:
        par = container.add_paragraph()
    # NB: local, not `p.right_indent +=`. The refine loop writes the same
    # layout more than once, and mutating it here would compound the
    # correction on every pass.
    gdocs_rows = p.gdocs_rows if ctx.output_profile == "gdocs" else []
    right_indent = 0.0 if gdocs_rows else p.right_indent + _wrap_correction(p, content_w)
    pf = par.paragraph_format
    par.alignment = ALIGN.get("left" if gdocs_rows else p.align, WD_ALIGN_PARAGRAPH.LEFT)
    gap = p.space_before if space_before is None else space_before
    if gap > 0.05:
        pf.space_before = Pt(round(gap, 1))
    else:
        pf.space_before = Pt(0)
    pf.space_after = Pt(round(max(0.0, p.space_after), 1))
    if p.leading and p.leading > 1:
        dom, fam = 0.0, ""
        if p.runs:
            w = {}
            for r in p.runs:
                if r.text and not r.is_tab:
                    key = (r.size, map_font(r.font, mono=r.mono, serif=r.serif))
                    w[key] = w.get(key, 0) + len(r.text)
            if w:
                dom, fam = max(w, key=w.get)
        _apply_leading(pf, p.leading, dom, mode=ctx.line_mode, family=fam)
    if p.left_indent > 0.05:
        pf.left_indent = Pt(round(p.left_indent, 1))
    if abs(p.first_indent) > 0.05:
        pf.first_line_indent = Pt(round(p.first_indent, 1))
    if right_indent > 0.05:
        pf.right_indent = Pt(round(right_indent, 1))
    for pos, al in p.tab_stops:
        pf.tab_stops.add_tab_stop(Pt(round(pos, 1)), TABAL.get(al, WD_TAB_ALIGNMENT.LEFT))
    # keep heading with following content
    if p.heading:
        pf.keep_with_next = True
        ppr = par._p.get_or_add_pPr()
        lvl = OxmlElement("w:outlineLvl")
        lvl.set(qn("w:val"), str(min(8, p.heading - 1)))
        ppr.append(lvl)
    # paragraph borders from header/footer rules
    bt = getattr(p, "border_top", None)
    bb = getattr(p, "border_bottom", None)
    if bt or bb:
        ppr = par._p.get_or_add_pPr()
        bd = {}
        if bt:
            bd["top"] = bt
        if bb:
            bd["bottom"] = bb
        _set_borders(ppr, bd, "w:pBdr")

    i = 0
    if gdocs_rows:
        runs = []
        for row_i, row in enumerate(gdocs_rows):
            runs.extend(row)
            if row_i < len(gdocs_rows) - 1:
                style = row[0] if row else Run(text="", font="Helvetica", size=10,
                                               color="#000000")
                runs.append(Run(text="\n", font=style.font, size=style.size,
                                color=style.color, bold=style.bold, italic=style.italic,
                                mono=style.mono, serif=style.serif))
    else:
        runs = p.runs
    while i < len(runs):
        run = runs[i]
        if run.link:
            grp = []
            while i < len(runs) and runs[i].link == run.link:
                grp.append((runs[i].text, runs[i]))
                i += 1
            _add_hyperlink(par, run.link, grp)
            continue
        if run.dest is not None:
            grp = []
            while i < len(runs) and runs[i].dest == run.dest:
                grp.append((runs[i].text, runs[i]))
                i += 1
            anchor = ctx.dest_anchors.get(run.dest)
            if anchor:
                _add_internal_hyperlink(par, anchor, grp)
            else:
                # A destination whose page holds no flow element to anchor to
                # (an all-figure page, say). Write the text plainly rather than
                # a hyperlink pointing at a bookmark that was never emitted.
                for text, style in grp:
                    _style_run(par.add_run(text), style)
            continue
        if run.field:
            _add_field(par, run.field, "1", run)
            i += 1
            continue
        if run.is_tab:
            r = par.add_run()
            r.add_tab()
            _style_run(r, run)
            i += 1
            continue
        # split on newlines -> soft breaks
        parts = run.text.split("\n")
        for j, chunk in enumerate(parts):
            if chunk:
                r = par.add_run(chunk)
                _style_run(r, run)
            if j < len(parts) - 1:
                br = par.add_run()
                br.add_break(WD_BREAK.LINE)
                _style_run(br, run)
        i += 1
    bookmark = getattr(p, "_bookmark", None)
    if bookmark and bookmark in ctx.anchor_ids:
        _wrap_paragraph_bookmark(par, bookmark, ctx.anchor_ids[bookmark])
    return par


def _spacer(container, height_pt: float):
    """Tiny exact-height paragraph used as vertical spacing before tables."""
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(max(1.0, round(height_pt, 1)))
    r = par.add_run("")
    r.font.size = Pt(1)
    return par


def _cell_text_width(cell) -> float:
    """Widest single source line in the cell, in pt. 0 when unmeasurable.

    This used to re-shape the text through MuPDF's base-14 metric tables, and
    that was both the writer's only hard dependency on PyMuPDF and a worse answer
    than the one already in the IR. `infer` records the width of every source line
    from its bbox, so for the question this function exists to answer -- is this
    column too narrow for content that occupied exactly one line in the source? --
    the source's own measurement is what actually happened rather than a
    prediction of what will happen. The font mapping is metric-compatible by
    design (Helvetica->Arial, Times->Times New Roman) precisely so the two agree.

    Unmeasurable still returns 0, and the caller still declines to resize on 0. It
    is reached differently now: not "this font has no base-14 equivalent" but
    "this paragraph wrapped in the source, so its width is the column's and says
    nothing about what the content needs", or "the cell was built by a path that
    records no line widths". An absent fact must not be read as a width of zero,
    which is why `source_line_width` returns None and this converts it here.
    """
    widest = 0.0
    for p in cell.paras:
        if "\n" in p.text:
            continue                     # multi-line in source: allowed to wrap
        w = source_line_width(p)
        if w is not None:
            widest = max(widest, w)
    return widest


def _fit_col_widths(t: TableEl, content_w: float = 0.0) -> List[float]:
    """Widen any column too narrow for its own single-line content, funded by
    columns with slack. Table width is unchanged.

    Column boundaries are inferred from text x-position clustering, and on
    dense numeric tables (booktabs especially) they land tight: measured on an
    arXiv paper, '23.75' needed 22.5pt in a 15.6pt column while column 0 sat
    on 44pt of slack. Every such cell wraps to two lines, the row doubles, and
    a 13-row results table gains ~250pt -- the single largest contributor to
    LaTeX page inflation. A cell that occupied one line in the source must get
    a column wide enough to stay one line.
    """
    n = len(t.col_widths)
    widths = list(t.col_widths)
    if n == 0:
        return widths
    # The requirement must include each cell's OWN pads: pads encode the
    # source x-alignment (left pad = text x minus column boundary) and run
    # 7-8pt on booktabs cells. A flat allowance under-asks, the funded width
    # still wraps, and the fix silently does nothing -- measured exactly so on
    # its first run.
    need = [0.0] * n
    for row in t.rows:
        ci = 0
        for cell in row:
            if cell is None:
                ci += 1
                continue
            span = max(1, getattr(cell, "col_span", 1))
            w = _cell_text_width(cell)
            if w > 0 and span == 1 and ci < n:
                pads = cell.pad[1] + cell.pad[3] if len(cell.pad) >= 4 else 8.0
                need[ci] = max(need[ci], w + pads + 1.0)
            ci += span
    deficit = [max(0.0, need[i] - widths[i]) for i in range(n)]
    surplus = [max(0.0, widths[i] - need[i] - 1.0) if need[i] > 0
               else max(0.0, widths[i] - 12.0) for i in range(n)]
    if sum(deficit) <= 0.01:
        return widths
    # A column must be funded FULLY or not at all: a partially-widened column
    # still wraps, so the width is spent and nothing is fixed. (The first
    # version scaled every deficit proportionally when funds ran short --
    # measured effect: zero.) Funds are the gap up to the container width
    # first -- growing the table costs nothing visually, the source usually
    # leaves room -- then slack shaved from over-wide columns.
    grow = max(0.0, (content_w or 0.0) - sum(widths)) if content_w else 0.0
    have = grow + sum(surplus)
    order = sorted((i for i in range(n) if deficit[i] > 0),
                   key=lambda i: deficit[i])
    funded = []
    for i in order:
        if deficit[i] <= have + 0.01:
            funded.append(i)
            have -= deficit[i]
    spend = sum(deficit[i] for i in funded)
    for i in funded:
        widths[i] += deficit[i]
    from_slack = max(0.0, spend - grow)
    tot_sur = sum(surplus)
    if from_slack > 0.01 and tot_sur > 0.01:
        for i in range(n):
            if surplus[i] > 0:
                widths[i] -= from_slack * (surplus[i] / tot_sur)
    return widths


# Google adds ~14.8pt of its own above a page-leading band, and it is an
# ADDITION, not a clamp. Measured by testkit/probe_cover_band.py in live pass 2,
# one variable per page, against a LibreOffice control that honoured all twelve
# variants exactly:
#
#     requested top   0.0    4.0    8.0   14.4   20.0
#     Docs rendered  14.55  18.83  22.83  29.23  34.83
#
# so rendered = requested + ~14.8 throughout, with no value of the section top
# margin reaching the page edge. Asking for a header distance instead lands in
# the same place (header 14.4 with top 0 differs by 0.43pt).
#
# Two consequences. Compensation can only subtract where the source actually
# left 14.8pt or more above the band; below that the remainder is a floor and
# `max(0, ...)` accepts it rather than inventing negative spacing. And a true
# top bleed is UNREACHABLE in Google Docs -- roughly 14.6pt of white above a
# page-one band is a documented limitation of the target, not a defect here.
_GDOCS_COVER_BEFORE_COMP_TWIPS = 296  # 14.8pt, measured above

# ---- Google Docs paragraph boundaries: nothing to compensate --------------
# This profile briefly subtracted 3.0pt of space_before at every flow-element
# boundary, on the theory (testkit/docs_quirks.py) that Docs' importer adds
# ~3pt per boundary. The consented live pass of 2026-08-04 measured the theory
# directly against Google's OWN exported PDFs and it is wrong.
#
# Method: for each pair of consecutive flow elements, compare the gap Google
# rendered with the gap the source had, as
#     gap_delta = dy(first line of el i+1) - dy(last line of el i)
# The file asked for space_before - comp, so Docs' own contribution is
# A = gap_delta + comp. Differencing two dy values cancels the paragraph's
# internal line-height error, which a whole-page regression on cumulative
# compensation cannot do (there the two are collinear and the fit blames the
# boundary for all of it -- that is how 3.0pt survived review).
#
# Measured over 187 single-column boundaries in 12 corpus documents:
#     all boundaries      A = +0.10pt   95% CI [+0.04, +0.21]
#     into a body para    A = +0.04pt   95% CI [-2.47, +0.07]   n=147
#     into a heading      A = +0.75pt                            n= 40
#     into a table        A = +1.04pt   95% CI [+0.49, +1.67]   n= 17
# Boundaries that received the full 3.0pt subtraction rendered a gap 2.90pt
# SMALLER than the source: Docs honoured the subtraction and added nothing.
#
# So the subtraction was pure loss, and it accumulated -- c6_long carries 17.4
# boundaries per page, and its dy_p50 went 25.84pt with every word pulled UP
# (dy is negative in all 16 documents). The earlier ~3pt reading came from
# probes written with lineRule="exact"; this profile already retranslates
# exact leading into a multiple, so compensating again double-counted the same
# height twice.
#
# The heading and table residuals are real but sub-point, and this writer
# already quantises font sizes to the half-point (_quantised_size), which moves
# a baseline by up to ~0.5pt on its own. Encoding a +0.75pt constant would be
# encoding its own noise floor, so they are recorded here and not applied.
# Re-measure them if the corpus ever needs the last point of vertical fidelity.


# --- column-break emission --------------------------------------------------
# An explicit column break says "column one ends HERE". That is right only if
# column one's content actually reaches the bottom of the column and no
# further. Measured on an isolated OOXML matrix at y12_irs_pub15's own
# geometry:
#
#   column 1 content     with the break        without it
#   under-fills          19/66  (correct)      66/28  (columns MERGE)
#   just fits            66/66                 66/66  (identical)
#   OVERFLOWS            66/10  (column 2      66/66  (degrades by lines,
#                               abandoned)            not by a column)
#
# So the break protects the source's split when column one under-fills, and
# destroys a column when it overflows: the spill enters column two, and the
# break then fires from column two and advances to the next page. On
# y12_irs_pub15 that is the whole defect -- 59 source pages rendering as 114
# with every second column empty.
#
# The break is therefore emitted only when column one is predicted NOT to
# overflow. The prediction is `ladder.predict_lines`, the same greedy first-fit
# the quality ladder uses; it returns None when it cannot be trusted (non
# base-14 family, unmeasurable glyph), and None keeps the break, which is the
# behaviour that shipped.
#
# The boundary is biased toward keeping the break: the matrix shows the
# "just fits" case works either way, so the slack costs nothing there and buys
# safety against a prediction that is a line optimistic.
COL_OVERFLOW_SLACK_PT = 6.0


def _line_height(p: Para) -> float:
    """The height one rendered line of this paragraph occupies.

    `leading` is the source's own measured line pitch and is preferred whenever
    it is a real measurement; the 1.15 fallback is Word's default single spacing
    for a paragraph that arrived without one.
    """
    return p.leading if (p.leading and p.leading > 1) else \
        (max((r.size for r in p.runs if r.text), default=10.0) * 1.15)


def _element_height(el, avail_w: float, metrics) -> Optional[float]:
    """How tall `el` is predicted to render, gap included. None = unpredictable.

    A Para is measured by re-wrapping it: `predict_lines` is the greedy
    first-fit Word itself uses, so this is how many lines the WRITER will
    produce, which is the whole point -- the source's own line count is what
    the layout already budgeted for. Anything else is measured by its source
    bbox, which is what it will occupy because the writer pins its size.
    """
    if not isinstance(el, Para):
        bb = getattr(el, "bbox", None) or getattr(el, "clip", None) \
            or getattr(el, "_bbox", None)
        if bb is None:
            return None
        return (el.space_before or 0.0) + (bb[3] - bb[1]) + (el.space_after or 0.0)
    n = predict_lines_for(el, avail_w, metrics)
    if n is None:
        return None
    return el.space_before + n * _line_height(el) + el.space_after


def predict_lines_for(p: Para, avail_w: float, metrics) -> Optional[int]:
    """`ladder.predict_lines` against a positive available width."""
    from .ladder import predict_lines
    if avail_w <= 1.0:
        return None
    return predict_lines(p, avail_w, metrics)


def _text_metrics():
    """Shaping metrics, or None when this installation cannot shape text.

    None is the answer a non-base-14 font has always produced, and every caller
    here already treats it as "do not act": the permissive backend degrades to
    `NullMetrics` rather than raising, and the predictions above then decline.
    """
    try:
        from .metrics import get_metrics
        return get_metrics("mupdf")
    except Exception:
        return None


def _column_one_overflows(ch, content_w: float, lay: DocLayout) -> bool:
    """Is the first column's content predicted to outgrow its column?"""
    if ch.n_cols < 2:
        return False
    metrics = _text_metrics()
    if metrics is None:
        return False
    gap = ch.col_gap or 0.0
    col_w = (content_w - gap * (ch.n_cols - 1)) / ch.n_cols
    if col_w <= 1.0:
        return False
    capacity = lay.page_h - lay.margin_t - lay.margin_b - max(0.0, ch.pre_gap)
    used = 0.0
    for el in ch.elements:
        if isinstance(el, ColBreak):
            break
        if not isinstance(el, Para):
            bb = getattr(el, "bbox", None) or getattr(el, "clip", None) \
                or getattr(el, "_bbox", None)
            used += (bb[3] - bb[1]) if bb else 0.0
            continue
        n = predict_lines_for(el, col_w - el.left_indent - el.right_indent,
                              metrics)
        if n is None:
            return False          # not predictable: leave the break alone
        used += el.space_before + n * _line_height(el) + el.space_after
    return used > capacity + COL_OVERFLOW_SLACK_PT


# --- page-spill absorption ---------------------------------------------------
# Every source page ends in an explicit page break, so the reconstruction has no
# slack at the bottom. When a page's content renders one or two lines taller
# than the page box, those lines flow to a new rendered page -- and the hard
# break then fires and advances again, stranding them there alone. A one-line
# overflow costs a whole page. Measured on the expansion corpus at e5e7f30
# (testkit/probe_thin_pages.py), the excess pages ARE those stranded lines:
#
#     document   arm       page_err  thin  <=2 body lines
#     y02        pymupdf        +59    59             35
#     y02        pdfium         +48    47             30
#     y01        pdfium         +34    40             15
#
# `refine.py` already corrects this from a render -- reclaim the page's gap
# slack, largest gaps first, each keeping a floor. That correction is right and
# arrives too late for the profiles that ship no loop at all (the gdocs profile
# is refine0 by construction). So the same correction is made here from a
# PREDICTION instead of a measurement, spending the same currency by the same
# rule, and it is gated the way `_column_one_overflows` is gated: predict, act
# only on a prediction we trust, and bias the boundary toward doing nothing.
#
# The cap is on the STRANDED LINES, not on the overflow in points, and that
# distinction is the whole design. Measured on y02 source page 20: the flow runs
# 38pt past the page box -- three lines' worth -- yet exactly ONE line is
# stranded, because the last element is a running head sitting behind a 91pt
# gap, and that gap is what carries it over. A gap at the top of a page is
# dropped rather than rendered, so the overflow in points says nothing about how
# much content actually lands on the extra page. Capping on the overflow refused
# that page and 57 others like it on y02 alone; capping on what is stranded
# accepts it, and "stranded lines" is the same quantity
# testkit/probe_thin_pages.py counts on the render -- predicted here, observed
# there, so the fix can be held to the sizing.
#
# Two lines is where the mass is. On y02's reference arm the rendered spill
# pages carry one body line 32 times and two body lines twice; past that a page
# is not spilling, it is a page that genuinely does not fit.
SPILL_MAX_LINES = 2
# Bias at the page boundary, in the spirit of COL_OVERFLOW_SLACK_PT. A line
# poking this far past the bottom is read as fitting, so a prediction that is
# marginally pessimistic finds nothing stranded and the page keeps the
# behaviour that shipped.
SPILL_EDGE_SLACK_PT = 6.0
# Pay slightly more than predicted. Sizes are quantised to the half point and
# gaps to a tenth, so a payment of exactly the predicted overflow lands the page
# on the boundary it was trying to clear.
SPILL_SAFETY_PT = 2.0
# The gap floors are `refine.MIN_GAP_SCALE` and its absolute companion, kept
# numerically identical so the open- and closed-loop corrections cannot crush a
# page to two different depths.
SPILL_MIN_GAP_SCALE = 0.30
SPILL_GAP_FLOOR_PT = 2.0
# The page-break paragraph carries `line_spacing exactly 1pt` and continues onto
# the page it opens, so the body box is that much shorter than the margins say.
PAGE_BREAK_PARA_PT = 1.0


def _hf_height(part) -> float:
    """How much of the margin a header or footer part actually claims."""
    if part is None:
        return 0.0
    h = 0.0
    for el in part.elements:
        if isinstance(el, Para):
            n = max(1, el.src_lines or 1)
            h += (el.space_before or 0.0) + n * _line_height(el) \
                + (el.space_after or 0.0)
            continue
        bb = getattr(el, "bbox", None) or getattr(el, "clip", None) \
            or getattr(el, "_bbox", None)
        h += ((bb[3] - bb[1]) if bb else 0.0) \
            + (getattr(el, "space_before", 0.0) or 0.0)
    return h


def _body_capacity(lay: DocLayout) -> float:
    """The flow height a page really offers, footer and header included.

    Not `page_h - margin_t - margin_b`. `w:pgMar/@footer` is the distance from
    the bottom of the page to the bottom of the footer, and the footer grows
    upward: when it reaches past the bottom margin the renderer shortens the
    BODY to make room. An inferred bottom margin can easily be smaller than the
    footer distance -- y02 comes out at 14pt against the writer's default 36pt
    -- and a capacity taken from the margins alone would then be some 35pt too
    generous.

    It changes no firing on the four documents measured for this fix: none of
    them ends up with a footer part at all, because inference leaves their
    running heads in the body. It is here because the margins are not the box,
    and a document whose footer IS lifted would otherwise be modelled with a
    page that does not exist.

    The header is the same construct upside down and is modelled the same way.
    """
    hd = lay.header_default.distance if lay.header_default else 0.0
    fd = lay.footer_default.distance if lay.footer_default else 0.0
    top = max(lay.margin_t, hd + _hf_height(lay.header_default))
    bottom = max(lay.margin_b, fd + _hf_height(lay.footer_default))
    return lay.page_h - top - bottom - PAGE_BREAK_PARA_PT


def _page_spill(pg, content_w: float, lay: DocLayout):
    """-> (overflow_pt, stranded_lines) for one source page, or None.

    `overflow_pt` is how far the whole flow runs past the page box -- what the
    page's gaps would have to give up for nothing to be stranded.
    `stranded_lines` is how many rendered lines land past the bottom, which is
    what actually appears on the extra page. The two are different numbers and
    the second is the one that says whether this is a spill: see SPILL_MAX_LINES.

    `None` means the page cannot be predicted and must be left exactly as it is
    written today. Only single-column, non-continuation pages are answered: a
    multi-column page is already governed by `_column_one_overflows`, and two
    predictions correcting the same page against different capacity models is
    how a fix starts fighting itself; a continuation page has had its break
    dropped deliberately and has nothing to strand.
    """
    if getattr(pg, "continuation_only", False) or not pg.chunks:
        return None
    if any(ch.n_cols > 1 for ch in pg.chunks):
        return None
    metrics = _text_metrics()
    if metrics is None:
        return None
    capacity = _body_capacity(lay)
    bottom = capacity + SPILL_EDGE_SLACK_PT
    used, stranded = 0.0, 0
    for ch in pg.chunks:
        used += max(0.0, ch.pre_gap)
        for el in ch.elements:
            if isinstance(el, ColBreak):
                return None       # a column break on a one-column page: unmodelled
            if not isinstance(el, Para):
                h = _element_height(el, content_w, metrics)
                if h is None:
                    return None   # no box to measure: leave the page alone
                used += h
                if used > bottom:
                    # A block crossing the boundary is not a two-line spill,
                    # and how a renderer splits one is not modelled here.
                    return None
                continue
            n = predict_lines_for(
                el, content_w - el.left_indent - el.right_indent, metrics)
            if n is None:
                return None       # not predictable: leave the page alone
            lead = _line_height(el)
            used += el.space_before
            for _ in range(n):
                used += lead
                if used > bottom:
                    stranded += 1
            used += el.space_after
    return used - capacity, stranded


def _absorb_page_spill(pg, content_w: float, lay: DocLayout) -> dict:
    """Plan the gap reductions that keep a small spill on its own page.

    Returns `{id(element): new_space_before}`, empty when the page is to be
    written exactly as it is today. **Nothing is mutated**: the refine loop
    writes the same layout once per round and a correction applied in place
    would compound on every pass, which is the same reason `write_para` keeps
    its wrap correction local.

    Slack is taken from paragraph gaps only. They are the bulk of a text page's
    slack, they are the gaps whose loss the eye forgives, and confining the plan
    to them keeps the whole change inside one writer signature. A page whose
    paragraph gaps cannot cover the overflow in full is left alone: a partial
    payment spends the spacing and still loses the page.
    """
    got = _page_spill(pg, content_w, lay)
    if got is None:
        return {}
    overflow, stranded = got
    if stranded <= 0 or stranded > SPILL_MAX_LINES or overflow <= 0.0:
        return {}
    paras = [el for ch in pg.chunks for el in ch.elements
             if isinstance(el, Para)]
    if not paras:
        return {}
    want = overflow + SPILL_SAFETY_PT
    slack = []
    for p in paras:
        gap = p.space_before or 0.0
        floor = max(SPILL_GAP_FLOOR_PT, gap * SPILL_MIN_GAP_SCALE)
        take = gap - floor
        if take > 0.05:
            slack.append((take, gap, p))
    if sum(t for t, _, _ in slack) < want:
        return {}
    plan = {}
    # Largest gaps first, exactly as `refine._apply` reclaims them: a 40pt
    # section break and a 4pt paragraph gap are not equally elastic, and the eye
    # notices the section break shrinking long after it notices the other.
    for take, gap, p in sorted(slack, key=lambda t: -t[1]):
        if want <= 0.05:
            break
        paid = min(take, want)
        plan[id(p)] = round(gap - paid, 1)
        want -= paid
    return plan


def _band_accent_as_row(t: TableEl) -> TableEl:
    """Re-express a band cell's accent border as a shaded row of its own.

    Inference already records a cover block's accent stripe -- a thin
    full-width fill flush against the block -- as a bottom border on the band
    cell, and that is a faithful representation: LibreOffice renders it at the
    right colour and thickness from `w:tcBorders`.

    Google Docs does not. The 4pt orange stripe on 01_whitepaper is absent from
    every Docs render across live passes 2, 3 and 4, while the navy block above
    it -- which is `w:shd` cell shading on the same cell -- comes back every
    time. So shading is demonstrably honoured by Docs where a thick table
    border is not, and this trades one construct for the other on the profile
    that needs it.

    The stripe becomes a second row of the same table, shaded with the accent
    colour and pinned to the stripe's own height, and the border is dropped so
    the two cannot both render.
    """
    if len(t.rows) != 1 or len(t.rows[0]) != 1 or not t.rows[0][0]:
        return t
    cell = t.rows[0][0]
    spec = (cell.borders or {}).get("bottom")
    if not spec:
        return t
    thickness, color = spec
    if not color or thickness <= 0:
        return t
    main = copy.copy(cell)
    main.borders = {k: v for k, v in (cell.borders or {}).items() if k != "bottom"}
    # The stripe is carved OUT of the band, never added underneath it: the
    # source block is navy 170 + accent 4 = 174 total, not 174 + 4.
    #
    # Reducing row_heights[0] alone does nothing, and that was the bug live
    # pass 5 caught. Row 0 carries the cover text, and the writer deliberately
    # leaves text rows content-driven -- it pins w:trHeight only on rows with
    # no text -- so row_heights[0] is never emitted for this row. The height
    # that actually has to give is the cell's bottom padding, which
    # build_band_table sized against the whole block including the stripe.
    pad = tuple(cell.pad) if len(cell.pad) >= 4 else (0.0, 0.0, 0.0, 0.0)
    main.pad = (pad[0], pad[1], max(0.0, round(pad[2] - thickness, 1)), pad[3])
    stripe = Cell(shading=color, borders={}, pad=(0.0, 0.0, 0.0, 0.0),
                  paras=[Para(runs=[Run(text="", font="Helvetica", size=2,
                                        color="#000000")],
                              leading=max(1.0, thickness))])
    band_h = t.row_heights[0] if t.row_heights else None
    heights = [max(1.0, band_h - thickness) if band_h else None, thickness]
    out = copy.copy(t)
    out.rows = [[main], [stripe]]
    out.row_heights = heights
    return out


def write_table(container, t: TableEl, content_w: float, ctx=None,
                cover_band: bool = False):
    ctx = ctx or _DEFAULT_CTX
    # A coloured page-one cover band is the one table Google Docs treats
    # differently.  Do not infer this from ``role == 'band'``: header bands
    # share that role and must retain their ordinary table treatment.
    gdocs_cover = cover_band and ctx.output_profile == "gdocs"
    if gdocs_cover:
        t = _band_accent_as_row(t)
    n_rows = len(t.rows)
    n_cols = len(t.col_widths)
    if n_rows == 0 or n_cols == 0:
        return None
    t = copy.copy(t)
    t.col_widths = _fit_col_widths(t, content_w)
    if t.space_before > 0.5:
        _spacer(container, t.space_before)
    try:
        tbl = container.add_table(rows=n_rows, cols=n_cols)
    except TypeError:  # header/footer/cell containers require a width argument
        tbl = container.add_table(rows=n_rows, cols=n_cols,
                                  width=Emu(int(sum(t.col_widths) * 12700)))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = tbl._tbl.tblPr
    # fixed layout
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(int(round(sum(t.col_widths) * 20))))
    tw.set(qn("w:type"), "dxa")
    tblPr.append(tw)
    if t.left_indent > 0.5:
        ind = OxmlElement("w:tblInd")
        ind.set(qn("w:w"), str(int(round(t.left_indent * 20))))
        ind.set(qn("w:type"), "dxa")
        tblPr.append(ind)
    # no default borders / spacing; zero default cell margins
    tb = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + side)
        b.set(qn("w:val"), "nil")
        tb.append(b)
    tblPr.append(tb)
    mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement("w:" + side)
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tblPr.append(mar)
    # grid
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc, wpt in zip(grid.findall(qn("w:gridCol")), t.col_widths):
        gc.set(qn("w:w"), str(int(round(wpt * 20))))

    first_cover_text = True
    for ri, rowspec in enumerate(t.rows):
        row = tbl.rows[ri]
        if ri < t.repeat_header_rows:
            # Only inference backed by a repeated source row may request a
            # Word repeating header.  Inventing one changes ordinary PDFs.
            trPr = row._tr.get_or_add_trPr()
            hdr = OxmlElement("w:tblHeader")
            hdr.set(qn("w:val"), "true")
            trPr.append(hdr)
        h = t.row_heights[ri] if ri < len(t.row_heights) else None
        # Only pin height on rows with no text: text rows are content-driven
        # (cell pads + exact-leading paragraphs sum to the source height,
        # which renders identically in Word, Google Docs and LibreOffice).
        row_has_text = any(c and any(p.text.strip() for p in c.paras)
                           for c in rowspec)
        if h and not row_has_text:
            trPr = row._tr.get_or_add_trPr()
            th = OxmlElement("w:trHeight")
            th.set(qn("w:val"), str(int(round(h * 20))))
            th.set(qn("w:hRule"), "atLeast")
            trPr.append(th)
        # The content-driven model assumes pads + exact-leading paragraphs sum
        # to the source row height. That holds for ordinary tables and fails
        # for maths: a row of stacked sub/superscripts can occupy 5.2pt in the
        # source while its text carries an 11.6pt leading, so the row renders
        # at more than twice its height. Measured on an arXiv paper: a 73.6pt
        # seven-row table rendered 31.4pt (43%) taller, entirely from three
        # such rows. Where the source row is shorter than its own content,
        # the leading is compressed to fit rather than left to overflow.
        row_shrink = 1.0
        if h and row_has_text:
            need = 0.0
            for c in rowspec:
                if not c:
                    continue
                cell_h = (c.pad[0] + c.pad[2]) if len(c.pad) >= 4 else 0.0
                for p in c.paras:
                    lead = p.leading or (p.runs[0].size * 1.2 if p.runs else 11.0)
                    cell_h += max(1, p.src_lines or 1) * lead
                need = max(need, cell_h)
            if need > h + 0.5:
                row_shrink = max(MIN_ROW_SHRINK, h / need)
        for ci in range(n_cols):
            spec = rowspec[ci] if ci < len(rowspec) else None
            cell = tbl.cell(ri, ci)
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcw = OxmlElement("w:tcW")
            tcw.set(qn("w:w"), str(int(round(t.col_widths[ci] * 20))))
            tcw.set(qn("w:type"), "dxa")
            tcPr.append(tcw)
            if spec is None:
                _blank_cell(cell)
                _set_borders(tcPr, {}, "w:tcBorders")
                continue
            if spec.shading:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), _hex(spec.shading))
                tcPr.append(shd)
            _set_borders(tcPr, spec.borders, "w:tcBorders")
            tmar = OxmlElement("w:tcMar")
            pads = spec.pad  # (top, left, bottom, right)
            # Google ignores tcMar/left: first measured on the bleed cover
            # table, and the c7_code Google evidence shows the same signature
            # on ordinary cells (dx_p50 ~ the 10.7pt code-cell tcMar left).
            # Under the gdocs profile, move (rather than duplicate) the left
            # padding of EVERY cell to its paragraphs below, so a future
            # importer which starts honouring tcMar does not double it.  The
            # relocation is exact for renderers that do honour tcMar (Word,
            # LibreOffice): tcMar_left + max(0, indent - pad) and
            # max(indent, pad) land text at the same x.
            gdocs_cellpad = ctx.output_profile == "gdocs"
            emitted_pads = (pads[0], 0.0, pads[2], pads[3]) \
                if gdocs_cellpad else pads
            for side, val in zip(("top", "left", "bottom", "right"),
                                  emitted_pads):
                m = OxmlElement("w:" + side)
                m.set(qn("w:w"), str(max(0, int(round(val * 20)))))
                m.set(qn("w:type"), "dxa")
                tmar.append(m)
            tcPr.append(tmar)
            if spec.valign and spec.valign != "top":
                va = OxmlElement("w:vAlign")
                va.set(qn("w:val"), spec.valign)
                tcPr.append(va)
            if spec.paras:
                # Cell paragraphs carry indents measured from the CELL EDGE
                # (text x minus cell x), and the same distance is already
                # emitted as tcMar above -- so an unadjusted indent applies the
                # pad twice and the text area loses it twice. Measured: a
                # 35.8pt column left 14.6pt for '24.6' (17.5pt), so the number
                # wrapped char-by-char and the row doubled. Word indents are
                # measured from the tcMar edge; make the paragraphs agree.
                def _depadded(p, _s=row_shrink):
                    q = copy.copy(p)
                    if gdocs_cellpad:
                        # Standard rendering lands at max(source indent,
                        # tcMar left).  With tcMar moved to zero, carry that
                        # effective position exactly; adding would double the
                        # common case where inference already included the pad.
                        q.left_indent = max(0.0, p.left_indent, pads[1])
                    else:
                        q.left_indent = max(0.0, p.left_indent - pads[1])
                    q.right_indent = max(0.0, p.right_indent - pads[3])
                    if _s < 0.999 and p.leading:
                        q.leading = max(2.0, p.leading * _s)
                    return q
                first = cell.paragraphs[0]
                for pi, p in enumerate(spec.paras):
                    q = _depadded(p)
                    if gdocs_cover and first_cover_text and p.text.strip():
                        # The extra space is a direct paragraph before-spacing
                        # translation in Google Docs.  Preserve all other
                        # paragraph properties and floor at zero for short bands.
                        q = copy.copy(q)
                        before = max(0, int(round(p.space_before * 20)) -
                                     _GDOCS_COVER_BEFORE_COMP_TWIPS)
                        q.space_before = before / 20.0
                        first_cover_text = False
                    write_para(cell, q, t.col_widths[ci],
                               par=first if pi == 0 else None, ctx=ctx)
            else:
                _blank_cell(cell)
    return tbl


def _blank_cell(cell):
    par = cell.paragraphs[0]
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(2)
    r = par.add_run("")
    r.font.size = Pt(1)


def write_figure(container, fig: FigureEl, ctx=None, dpi: int = None):
    """Rasterise a figure region through the conversion's backend.

    `ctx.render_clip` replaces an open MuPDF document that used to be threaded
    down from `_write_docx`. A figure clip is a *rendering* operation, which the
    backend seam has always declared (`Backend.render_clip`) and which this writer
    was reaching around.

    Returns None if there is no renderer, and the caller then omits the figure --
    an honest empty space rather than a crash, and a warning once REL-01 lands.
    """
    ctx = ctx or _DEFAULT_CTX
    dpi = ctx.dpi if dpi is None else dpi
    if ctx.render_clip is None:
        return None
    data = ctx.render_clip(fig.page_no, fig.clip, dpi)
    if not data:
        return None
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(round(max(0.0, fig.space_before), 1))
    pf.space_after = Pt(0)
    if ctx.output_profile != "gdocs":
        # Word/LibreOffice need this guard for a paragraph containing only an
        # inline drawing.  Google Docs reserves the inline drawing itself and
        # treats the duplicate atLeast height as extra page-flow pressure.
        pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        pf.line_spacing = Pt(round(fig.height, 1))
    par.alignment = ALIGN.get(fig.align, WD_ALIGN_PARAGRAPH.CENTER)
    if fig.align == "left" and fig.left_indent > 0.5:
        pf.left_indent = Pt(round(fig.left_indent, 1))
    r = par.add_run()
    r.add_picture(io.BytesIO(data), width=Emu(int(fig.width * 12700)),
                  height=Emu(int(fig.height * 12700)))
    return par


def _docx_accepts(data: bytes) -> bool:
    """Whether python-docx will embed these bytes, asked without a Document.

    python-docx matches a small signature table (`docx.image.SIGNATURES`) against
    the first 32 bytes and raises `UnrecognizedImageError` for anything outside
    it. Its JPEG entries are JFIF (`FF D8 FF E0`) and Exif (`FF D8 FF E1`) only,
    so an **Adobe APP14 JPEG** -- `FF D8 FF EE`, what Antenna House and the rest
    of the Adobe toolchain emit -- is a perfectly valid JPEG that python-docx
    refuses. `Image.from_blob` runs exactly that check plus the chosen header
    parser, so a truncated header of a *recognised* format is caught here too
    rather than midway through serialising the package.
    """
    from docx.image.image import Image as _DocxImage
    try:
        _DocxImage.from_blob(data)
        return True
    except Exception:
        return False


def _to_png(data: bytes):
    """Re-encode through Pillow to PNG, or None if Pillow cannot read it either.

    PNG rather than a re-saved JPEG on purpose: the source is already lossily
    encoded, and a second lossy pass would quietly degrade the pixels to work
    around a *container* problem. Pillow reads the Adobe-APP14 JPEG above fine;
    only the signature table objected.
    """
    try:
        from PIL import Image as _PILImage
    except ImportError:            # Pillow is not a hard dependency of the writer
        return None
    img = None
    try:
        img = _PILImage.open(io.BytesIO(data))
        img.load()
        if img.mode not in ("1", "L", "LA", "P", "RGB", "RGBA"):
            # CMYK, I;16 and friends have no PNG representation.
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None
    finally:
        if img is not None:
            try:
                img.close()
            except Exception:
                pass


def _embeddable(data: bytes, report):
    """Bytes python-docx will take, or None -- and record which of the three.

    One image exactdoc cannot embed must never take the document down with it.
    `y06_irs_1040_instructions.pdf` is 126 pages of text carrying two Adobe-APP14
    JPEGs, and `add_picture` raising `UnrecognizedImageError` used to lose all
    126 pages rather than the two images.

    The ladder is: embed as extracted, else re-encode losslessly, else drop --
    and a drop is *counted*, not swallowed.
    """
    outcome, out = "embedded", data
    if not data:
        outcome, out = "dropped", None
    elif not _docx_accepts(data):
        png = _to_png(data)
        if png is not None and _docx_accepts(png):
            outcome, out = "reencoded", png
        else:
            outcome, out = "dropped", None
    if report is not None:
        report[outcome] = report.get(outcome, 0) + 1
    return out


def write_image(container, im: ImageEl, ctx=None):
    """Place an extracted raster. Returns None when the image had to be dropped.

    Returning None so the caller omits the element is `write_figure`'s contract
    for a visual it cannot produce, and this follows it: an honest empty space,
    tallied in `ctx.image_report`, rather than a crash.
    """
    ctx = ctx or _DEFAULT_CTX
    data = _embeddable(im.data, ctx.image_report)
    if data is None:
        return None
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(round(max(0.0, im.space_before), 1))
    pf.space_after = Pt(0)
    if ctx.output_profile != "gdocs":
        pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        pf.line_spacing = Pt(round(im.height, 1))
    par.alignment = ALIGN.get(im.align, WD_ALIGN_PARAGRAPH.CENTER)
    if im.align == "left" and im.left_indent > 0.5:
        pf.left_indent = Pt(round(im.left_indent, 1))
    r = par.add_run()
    r.add_picture(io.BytesIO(data), width=Emu(int(im.width * 12700)),
                  height=Emu(int(im.height * 12700)))
    return par


def write_rule(container, rule: RuleEl, content_w: float):
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(round(max(0.0, rule.space_before), 1))
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(2)
    length = rule.length or (content_w * rule.width_pct / 100.0)
    if rule.left_indent > 0.5:
        pf.left_indent = Pt(round(rule.left_indent, 1))
    ri = content_w - rule.left_indent - length
    if ri > 0.5:
        pf.right_indent = Pt(round(ri, 1))
    r = par.add_run("")
    r.font.size = Pt(1)
    ppr = par._p.get_or_add_pPr()
    _set_borders(ppr, {"bottom": (rule.thickness, rule.color)}, "w:pBdr")
    return par


# ------------------------------------------------------------------ sections
def _config_section(sec, lay: DocLayout, margin_t=None, cols: int = 1,
                    col_gap: float = 24.0, margin_lr=None):
    sec.page_width = Emu(int(lay.page_w * 12700))
    sec.page_height = Emu(int(lay.page_h * 12700))
    ml = lay.margin_l if margin_lr is None else margin_lr
    mr = lay.margin_r if margin_lr is None else margin_lr
    sec.left_margin = Emu(int(ml * 12700))
    sec.right_margin = Emu(int(mr * 12700))
    sec.top_margin = Emu(int((lay.margin_t if margin_t is None else margin_t) * 12700))
    sec.bottom_margin = Emu(int(lay.margin_b * 12700))
    hd = lay.header_default.distance if lay.header_default else 36.0
    fd = lay.footer_default.distance if lay.footer_default else 36.0
    sec.header_distance = Emu(int(max(0.0, hd) * 12700))
    sec.footer_distance = Emu(int(max(0.0, fd) * 12700))
    sectPr = sec._sectPr
    cols_el = sectPr.find(qn("w:cols"))
    if cols_el is None:
        cols_el = OxmlElement("w:cols")
        sectPr.append(cols_el)
    if cols > 1:
        cols_el.set(qn("w:num"), str(cols))
        cols_el.set(qn("w:space"), str(int(round(col_gap * 20))))
        cols_el.set(qn("w:equalWidth"), "1")
    else:
        cols_el.set(qn("w:num"), "1")
        for a in ("w:space", "w:equalWidth"):
            if cols_el.get(qn(a)):
                cols_el.attrib.pop(qn(a))


def _shifted_part(part: Optional[HFPart], dl: float, dr: float) -> Optional[HFPart]:
    """Clone a header/footer part with indents/tabs shifted (bleed sections)."""
    if part is None:
        return None
    import copy
    np = copy.deepcopy(part)
    for el in np.elements:
        if isinstance(el, Para):
            el.left_indent = round(el.left_indent + dl, 1)
            el.right_indent = round((el.right_indent or 0.0) + dr, 1)
            el.tab_stops = [(round(p + dl, 1), a) for p, a in el.tab_stops]
        elif isinstance(el, TableEl):
            el.left_indent = round(el.left_indent + dl, 1)
    return np


def _fill_hf(hf_obj, part: Optional[HFPart], lay: DocLayout, ctx=None):
    """Fill a python-docx header/footer object with an HFPart."""
    ctx = ctx or _DEFAULT_CTX
    hf_obj.is_linked_to_previous = False
    # clear default paragraph content
    first = hf_obj.paragraphs[0]
    first.clear()
    if part is None or not part.elements:
        pf = first.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(2)
        return
    used_first = False
    for el in part.elements:
        if isinstance(el, TableEl):
            write_table(hf_obj, el, lay.content_w, ctx=ctx)
        elif isinstance(el, Para):
            if not used_first:
                write_para(hf_obj, el, lay.content_w, par=first, ctx=ctx)
                used_first = True
            else:
                write_para(hf_obj, el, lay.content_w, ctx=ctx)
        elif isinstance(el, RuleEl):
            write_rule(hf_obj, el, lay.content_w)
    if not used_first:
        # first paragraph unused: make it invisible
        pf = first.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(2)


# ------------------------------------------------------------------ main
def write_docx(lay: DocLayout, out_path: str, dpi: int = 240,
               output_profile: str = "standard", backend=None, ctx=None,
               image_report=None) -> str:
    """Render a DocLayout to a .docx. Pure: `lay` is never modified.

    `output_profile` selects the line-height encoding: Word and LibreOffice
    honour lineRule="exact", Google Docs mistranslates it in a way that scales
    with font size, so the gdocs profile emits the same intent as a multiple
    instead. That choice now travels in a `WriteCtx` rather than in a module
    global that this function set and restored -- two concurrent conversions with
    different profiles could each observe the other's encoding.

    This is a pure serialisation setting. It writes different bytes; it does not
    contact anything. Choosing the Google-safe profile costs no network, no
    credentials and no upload.

    `backend` supplies figure rasterisation. Pass the same backend the parse used;
    without one, figure regions are omitted rather than rendered through a parser
    nobody selected. This is what removed `import fitz` from the top of this
    module, and with it the reason a wheel installed without PyMuPDF could not
    write a DOCX at all.

    The cover-band path shifts every page-1 element by the bleed delta, and
    those shifts are *accumulating* assignments (`el.left_indent + delta_l`,
    `c.pad[1] + delta_l`). Writing the same layout twice therefore used to
    double-shift the second document — silently, and only on cover-band
    documents, which is why it survived: 2 of 16 corpus documents were not
    reproducible on a second write. Callers must not have to know this, so the
    copy lives here and purity is part of the contract, verified by
    tests/test_purity.py.

    `image_report`, when given, is cleared and refilled with this write's raster
    tally (`embedded`/`reencoded`/`dropped`). Cleared rather than accumulated
    because the refine loop writes the same layout once per round, and a ledger
    that summed over rounds would report four dropped images for one.
    """
    if image_report is not None:
        image_report.clear()
    if ctx is None:
        render_clip = None
        if backend is not None and lay.src_path:
            def render_clip(page_no, clip, at_dpi, _bk=backend, _p=lay.src_path):
                try:
                    return _bk.render_clip(_p, page_no, clip, dpi=at_dpi)
                except Exception:
                    return None
        ctx = WriteCtx(output_profile=output_profile,
                       line_mode=line_mode_for(output_profile), dpi=dpi,
                       render_clip=render_clip, image_report=image_report)
    return _write_docx(lay, out_path, ctx)


def _write_docx(lay: DocLayout, out_path: str, ctx: WriteCtx) -> str:
    lay = copy.deepcopy(lay)
    # After the deepcopy: the plan marks the elements this function will write.
    dest_anchors, anchor_ids = _plan_bookmarks(lay)
    if dest_anchors or anchor_ids:
        ctx = dataclasses.replace(ctx, dest_anchors=dest_anchors,
                                  anchor_ids=anchor_ids)
    if ctx.output_profile == "gdocs":
        # Substitute fonts whose measured advance width does not match the
        # source's, and track out what remains, so paragraphs wrap where they
        # wrapped in the PDF. Safe here: the layout above is already a copy,
        # and this rewrites run properties only -- never element identity, so
        # the bookmark plan above stays valid.
        from .gdocs_metrics import apply_metric_fit
        apply_metric_fit(lay)
    doc = Document()
    dpi = ctx.dpi
    content_w = lay.content_w

    # neutralize the template's Normal style (1.08 line, 8pt after) so nothing
    # inherits spacing we didn't ask for
    try:
        npf = doc.styles["Normal"].paragraph_format
        npf.space_before = Pt(0)
        npf.space_after = Pt(0)
        npf.line_spacing = 1.0
    except Exception:
        pass
    if lay.hyphenated:
        # source justifies with hyphenation: let Word/Docs hyphenate too so
        # line packing (and therefore paragraph heights) stay comparable
        st = doc.settings.element
        if st.find(qn("w:autoHyphenation")) is None:
            ah = OxmlElement("w:autoHyphenation")
            ah.set(qn("w:val"), "1")
            st.append(ah)

    sec = doc.sections[0]
    has_cover = lay.cover_band is not None
    # Cover-band section side margins. The 4pt was a hedge against renderers
    # refusing a zero margin, and it cost a visible 3.9pt white frame down both
    # edges of every Google cover page -- the source band bleeds to the paper.
    # probe_cover_band's SIDE family measured Docs honouring side margins
    # exactly (requests of 0/2/4/8 tracked one-for-one), so under the gdocs
    # profile it can simply ask for zero and get a true bleed. Other targets
    # keep the hedge, which also keeps the standard profile byte-identical.
    band_bleed = 0.0 if ctx.output_profile == "gdocs" else 4.0
    if has_cover:
        _config_section(sec, lay, margin_t=lay.cover_top, cols=1, margin_lr=band_bleed)
        sec.header_distance = Emu(0)  # body must start flush at the band
    else:
        _config_section(sec, lay, cols=1)

    # headers/footers for section 1 (never create empty parts: an empty header
    # still reserves a line and pushes the body down)
    if has_cover:
        dl = lay.margin_l - band_bleed
        dr = lay.margin_r - band_bleed
        if lay.header_first is not None:
            _fill_hf(sec.header, _shifted_part(lay.header_first, dl, dr), lay,
                     ctx=ctx)
        if (lay.footer_first or lay.footer_default) is not None:
            _fill_hf(sec.footer,
                     _shifted_part(lay.footer_first or lay.footer_default, dl, dr),
                     lay, ctx=ctx)
    else:
        if lay.header_default is not None:
            _fill_hf(sec.header, lay.header_default, lay, ctx=ctx)
        if lay.footer_default is not None:
            _fill_hf(sec.footer, lay.footer_default, lay, ctx=ctx)
        if lay.different_first:
            sec.different_first_page_header_footer = True
            _fill_hf(sec.first_page_header, lay.header_first, lay, ctx=ctx)
            _fill_hf(sec.first_page_footer,
                     lay.footer_first or lay.footer_default, lay, ctx=ctx)

    cur_cols = 1
    # config of the currently-open section; re-applied after each break because
    # python-docx's add_section clones sectPr elements in ways that can shuffle
    # previously-applied properties between sections
    cur_cfg = {"margin_t": (lay.cover_top if has_cover else None), "cols": 1,
               "gap": 24.0, "margin_lr": (band_bleed if has_cover else None),
               "hdr0": has_cover}

    def new_section(kind, cols, gap=24.0, margin_t=None, margin_lr=None):
        nonlocal sec, cur_cols, cur_cfg
        doc.add_section(kind)
        # re-apply the finished section's geometry to whatever element now
        # represents it
        fin = doc.sections[-2]
        _config_section(fin, lay, margin_t=cur_cfg["margin_t"], cols=cur_cfg["cols"],
                        col_gap=cur_cfg["gap"], margin_lr=cur_cfg["margin_lr"])
        if cur_cfg.get("hdr0"):
            fin.header_distance = Emu(0)
        sec = doc.sections[-1]
        _config_section(sec, lay, margin_t=margin_t, cols=cols, col_gap=gap,
                        margin_lr=margin_lr)
        cur_cfg = {"margin_t": margin_t, "cols": cols, "gap": gap,
                   "margin_lr": margin_lr, "hdr0": False}
        cur_cols = cols
        # Shrink section-break paragraphs to the least height a renderer will
        # give them. That is SECT_BREAK_PARA_PT, not zero -- see the constant.
        for p_el in doc.element.body.findall(qn("w:p")):
            ppr = p_el.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                if ppr.find(qn("w:spacing")) is None:
                    sp = OxmlElement("w:spacing")
                    sp.set(qn("w:before"), "0")
                    sp.set(qn("w:after"), "0")
                    sp.set(qn("w:line"), str(SECT_BREAK_PARA_TWIPS))
                    sp.set(qn("w:lineRule"), "exact")
                    ppr.append(sp)
        return sec

    if has_cover:
        # The whole cover page lives in ONE bleed-margin section (renderers do
        # not honor L/R margin changes at mid-page continuous breaks). The band
        # spans the bleed width; every other page-1 element is shifted right by
        # indents so it keeps its original x-position.
        delta_l = lay.margin_l - band_bleed
        delta_r = lay.margin_r - band_bleed
        band = lay.cover_band
        band.col_widths = [lay.page_w - 2 * band_bleed]
        for c in band.rows[0]:
            if c:
                c.pad = (c.pad[0], round(c.pad[1] + delta_l, 1), c.pad[2], c.pad[3])
        for ch in lay.pages[0].chunks:
            for el in ch.elements:
                if isinstance(el, Para):
                    el.left_indent = round(el.left_indent + delta_l, 1)
                    el.right_indent = round((el.right_indent or 0.0) + delta_r, 1)
                    el.tab_stops = [(round(p + delta_l, 1), a) for p, a in el.tab_stops]
                elif isinstance(el, TableEl):
                    el.left_indent = round(el.left_indent + delta_l, 1)
                elif isinstance(el, (FigureEl, ImageEl)):
                    if el.align == "left":
                        el.left_indent = round(el.left_indent + delta_l, 1)
                elif isinstance(el, RuleEl):
                    el.left_indent = round(el.left_indent + delta_l, 1)
        write_table(doc, band, lay.page_w - 2 * band_bleed, ctx=ctx,
                    cover_band=True)

    last_el_par = None
    for pi, pg in enumerate(lay.pages):
        if pi > 0 and not pg.continuation_only:
            # page boundary
            after_cover = has_cover and pi == 1
            next_cols = pg.chunks[0].n_cols if pg.chunks else 1
            if after_cover or cur_cols != next_cols:
                gap = pg.chunks[0].col_gap if pg.chunks else 24.0
                pre = pg.chunks[0].pre_gap if pg.chunks else 0.0
                mt = (lay.margin_t + pre) if (next_cols > 1 and pre > 0.5) else None
                s = new_section(WD_SECTION.NEW_PAGE, next_cols, gap, margin_t=mt)
                if after_cover:
                    _fill_hf(s.header, lay.header_default, lay, ctx=ctx)
                    _fill_hf(s.footer, lay.footer_default, lay, ctx=ctx)
            else:
                par = doc.add_paragraph()
                pf = par.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(1)
                par.add_run().add_break(WD_BREAK.PAGE)
        cw_ctx = (lay.page_w - 2 * band_bleed) if (has_cover and pi == 0) else content_w
        # A one- or two-line spill is absorbed into this page rather than
        # stranded on one of its own by the break that follows. The plan is
        # `{id(element): gap}` and is applied at write time only: `lay` is
        # written once per refine round and a gap reduced in place would
        # compound on every pass. The cover page keeps its own bleed geometry
        # and is never asked. See `_absorb_page_spill`.
        spill_plan = {} if (has_cover and pi == 0) \
            else _absorb_page_spill(pg, cw_ctx, lay)
        for ci, ch in enumerate(pg.chunks):
            if ch.n_cols != cur_cols:
                if ch.pre_gap > 0.5:
                    _spacer(doc, ch.pre_gap - _sect_break_comp(ctx))
                new_section(WD_SECTION.CONTINUOUS, ch.n_cols, ch.col_gap)
            drop_col_break = _column_one_overflows(ch, cw_ctx, lay)
            for el in ch.elements:
                if isinstance(el, ColBreak):
                    if drop_col_break:
                        # Column one is predicted to overflow. Forcing the
                        # break here would fire it from column TWO and abandon
                        # that column; letting the content flow costs a few
                        # lines instead of a whole column. See
                        # _column_one_overflows for the matrix.
                        continue
                    par = doc.add_paragraph()
                    pf = par.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    pf.line_spacing = Pt(1)
                    par.add_run().add_break(WD_BREAK.COLUMN)
                    continue
                if isinstance(el, Para):
                    write_para(doc, el, cw_ctx, ctx=ctx,
                               space_before=spill_plan.get(id(el)))
                    continue
                bookmark = getattr(el, "_bookmark", None)
                if bookmark and bookmark in ctx.anchor_ids:
                    # Not a paragraph: mark the spot between block elements.
                    _add_block_bookmark(doc, bookmark, ctx.anchor_ids[bookmark])
                if isinstance(el, TableEl):
                    write_table(doc, el, cw_ctx, ctx=ctx)
                elif isinstance(el, FigureEl):
                    write_figure(doc, el, ctx=ctx)
                elif isinstance(el, ImageEl):
                    write_image(doc, el, ctx=ctx)
                elif isinstance(el, RuleEl):
                    write_rule(doc, el, cw_ctx)

    # drop the initial empty paragraph python-docx puts in a fresh document
    # (never touch section-break paragraphs: removing one deletes a section)
    body = doc.element.body
    paras = body.findall(qn("w:p"))
    if paras:
        p0 = paras[0]
        has_content = p0.findall(qn("w:r")) or p0.findall(qn("w:hyperlink")) \
            or p0.findall(qn("w:fldSimple"))
        ppr = p0.find(qn("w:pPr"))
        has_sectpr = ppr is not None and ppr.find(qn("w:sectPr")) is not None
        if not has_content and not has_sectpr and len(list(body)) > 2:
            body.remove(p0)

    doc.save(out_path)
    return out_path
