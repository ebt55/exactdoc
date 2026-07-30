"""Fidelity verification: DOCX -> PDF (LibreOffice) -> image diff vs source.

Product diagnostics, not release evidence: `testkit/` is the independent harness
and deliberately shares no code with the converter. This module exists so that
`--verify` can tell a user something about their own document.

Page rasterisation goes through the backend seam (`Backend.render_page`) rather
than through `fitz` directly. The import used to be at module scope, which put
PyMuPDF on the default runtime path of a stage that only wants pixels.
"""
import io
import os
import re
import subprocess
import tempfile
from typing import List, Optional

import numpy as np


def _find_soffice():
    """Locate LibreOffice on any platform.

    The previous list held POSIX paths only, so `--verify` silently reported
    'LibreOffice not found' on every Windows and macOS machine even with a
    working install.
    """
    import shutil
    env = os.environ.get("SOFFICE")
    if env and os.path.exists(env):
        return env
    for cand in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/opt/libreoffice26.2/program/soffice",
            "/usr/bin/soffice", "/usr/local/bin/soffice"):
        if os.path.exists(cand):
            return cand
    return shutil.which("soffice") or shutil.which("libreoffice")


SOFFICE = _find_soffice()


def docx_to_pdf(docx_path: str, out_dir: str, profile: Optional[str] = None
                ) -> Optional[str]:
    """Render a DOCX to PDF. Returns None if LibreOffice is unavailable.

    A dedicated user profile is required, not a nicety: soffice refuses rapid
    successive starts against a shared default profile and exits 0 without
    writing anything, which looks exactly like a silent conversion failure.

    It was a single fixed path under the temp directory, shared by every
    conversion in every process on the machine, so two concurrent conversions
    contended for one profile and reproduced the very failure the profile
    existed to prevent. Each invocation now gets a **fresh** directory and
    removes it afterwards, so concurrency is safe by construction and no state
    survives from one render to influence the next.

    `profile=` overrides that for a caller who wants one profile across a batch.
    `testkit/harness.py` deliberately does exactly that: soffice also refuses
    *rapid* restarts against differing profiles, so a tight batch loop wants one
    warm profile, while a product conversion wants isolation. Those are different
    trade-offs and both are now expressible.
    """
    if SOFFICE is None:
        return None
    env = dict(os.environ)
    env.setdefault("HOME", tempfile.gettempdir())
    owned = profile is None
    prof = profile or tempfile.mkdtemp(prefix="exactdoc_soffice_")
    out = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if os.path.exists(out):
        os.remove(out)
    cmd = [SOFFICE, "--headless", "--norestore", "--invisible", "--nolockcheck",
           "-env:UserInstallation=file:///" + prof.replace("\\", "/"),
           "--convert-to", "pdf", "--outdir", out_dir, docx_path]
    try:
        subprocess.run(cmd, capture_output=True, timeout=300, env=env)
    except (subprocess.TimeoutExpired, OSError):
        return None
    finally:
        if owned:
            import shutil
            shutil.rmtree(prof, ignore_errors=True)
    return out if os.path.exists(out) else None


def _page_count(pdf_path: str, backend) -> int:
    """Page count via the seam. `page_lines` returns one entry per page, text or
    not, so its length is the count -- extracting text to count pages is wasteful,
    but this is a diagnostic that goes on to rasterise every page, and a fourth
    seam operation for it would be a worse trade than the wasted extraction."""
    return len(backend.page_lines(pdf_path))


def _page_arrays(pdf_path: str, dpi: int = 96, backend=None) -> List[np.ndarray]:
    """Rasterise every page to an RGB array, through the backend.

    Decoding the backend's PNG with Pillow costs one encode/decode per page that
    reading a MuPDF pixmap's raw samples did not. That is the price of the seam
    being a byte format both backends can honestly produce, and it is paid by a
    diagnostic path, not by conversion. The independent harness in `testkit/` is
    where render performance would matter, and it is free to use whatever it likes
    -- it is not shipped.
    """
    if backend is None:
        from .backend import get_backend
        backend = get_backend()
    import PIL.Image as Image
    out = []
    for i in range(_page_count(pdf_path, backend)):
        data = backend.render_page(pdf_path, i + 1, dpi=dpi)
        if not data:
            break
        arr = np.asarray(Image.open(io.BytesIO(data)).convert("RGB"))
        out.append(arr.astype(np.float64))
    return out


def _pad_to(a: np.ndarray, h: int, w: int) -> np.ndarray:
    out = np.full((h, w, a.shape[2]), 255.0)
    out[:a.shape[0], :a.shape[1], :] = a[:h, :w, :]
    return out


def ssim(a: np.ndarray, b: np.ndarray) -> float:
    """Mean SSIM over 8x8 windows on the grayscale image."""
    ga = a.mean(axis=2)
    gb = b.mean(axis=2)
    k = 8
    H = (ga.shape[0] // k) * k
    W = (ga.shape[1] // k) * k
    ga = ga[:H, :W].reshape(H // k, k, W // k, k).transpose(0, 2, 1, 3).reshape(-1, k * k)
    gb = gb[:H, :W].reshape(H // k, k, W // k, k).transpose(0, 2, 1, 3).reshape(-1, k * k)
    mu_a = ga.mean(1)
    mu_b = gb.mean(1)
    va = ga.var(1)
    vb = gb.var(1)
    cov = ((ga - mu_a[:, None]) * (gb - mu_b[:, None])).mean(1)
    C1, C2 = (0.01 * 255) ** 2, (0.03 * 255) ** 2
    s = ((2 * mu_a * mu_b + C1) * (2 * cov + C2)) / \
        ((mu_a ** 2 + mu_b ** 2 + C1) * (va + vb + C2))
    return float(s.mean())


def compare(src_pdf: str, converted_pdf: str, out_dir: Optional[str] = None,
            dpi: int = 96, backend=None):
    A = _page_arrays(src_pdf, dpi, backend=backend)
    B = _page_arrays(converted_pdf, dpi, backend=backend)
    n = max(len(A), len(B))
    rows = []
    for i in range(n):
        a = A[i] if i < len(A) else None
        b = B[i] if i < len(B) else None
        if a is None or b is None:
            rows.append({"page": i + 1, "ssim": 0.0, "note": "page count mismatch"})
            continue
        h = max(a.shape[0], b.shape[0])
        w = max(a.shape[1], b.shape[1])
        a2, b2 = _pad_to(a, h, w), _pad_to(b, h, w)
        s = ssim(a2, b2)
        mad = float(np.abs(a2 - b2).mean())
        rows.append({"page": i + 1, "ssim": round(s, 4), "mad": round(mad, 2)})
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)
            import PIL.Image as Image
            gapw = 12
            canvas = np.full((h, w * 2 + gapw, 3), 200.0)
            canvas[:, :w, :] = a2
            canvas[:, w + gapw:, :] = b2
            Image.fromarray(canvas.astype(np.uint8)).save(
                os.path.join(out_dir, "cmp_p%02d.png" % (i + 1)))
    return rows


def audit(src_pdf: str, docx_path: str, backend=None):
    """Text-coverage audit: is every source character present in the DOCX?

    Rasterized figure regions (charts) legitimately carry their labels as
    pixels, so their text is excluded from the source side. That exclusion is
    exactly why this is a *diagnostic* and not evidence: it lets the converter
    define its own denominator, and the original `verify.py` scored the résumé at
    `src_chars: 0` because everything it rasterised vanished from its own score.
    `testkit/harness.py` is raster-blind on purpose and shares no code with this.
    """
    from .infer import infer
    from .layout import FigureEl
    from .model import bbox_overlap, bbox_area
    import docx as _docx
    from collections import Counter

    if backend is None:
        from .backend import get_backend
        backend = get_backend()
    ir = backend.parse_pdf(src_pdf, keep_image_data=False)
    lay = infer(ir)
    fig_clips = {}
    for pg in lay.pages:
        for ch in pg.chunks:
            for el in ch.elements:
                if isinstance(el, FigureEl):
                    fig_clips.setdefault(el.page_no, []).append(el.clip)
    src_parts = []
    for p in ir.pages:
        clips = fig_clips.get(p.number, [])
        for b in p.blocks:
            for l in b.lines:
                in_fig = any(bbox_overlap(l.bbox, c) > 0.5 * max(1e-6, bbox_area(l.bbox))
                             for c in clips)
                if not in_fig:
                    src_parts.append(l.text)
    d = _docx.Document(docx_path)
    out_parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                out_parts.append(c.text)
    for s in d.sections:
        for hf in (s.header, s.footer, s.first_page_header, s.first_page_footer):
            try:
                for p in hf.paragraphs:
                    out_parts.append(p.text)
                for t in hf.tables:
                    for row in t.rows:
                        for c in row.cells:
                            out_parts.append(c.text)
            except Exception:
                pass

    def norm(t):
        return re.sub(r"[\s ]+", "", t)

    def grams(s, k=3):
        return Counter(s[i:i + k] for i in range(max(0, len(s) - k + 1)))

    src_n = norm("".join(src_parts))
    out_n = norm("".join(out_parts))
    g1, g2 = grams(src_n), grams(out_n)
    inter = sum(min(c, g2[g]) for g, c in g1.items())
    cov = inter / max(1, sum(g1.values()))
    return {"src_chars": len(src_n), "docx_chars": len(out_n),
            "text_coverage": round(cov, 4)}


def verify(src_pdf: str, docx_path: str, out_dir: Optional[str] = None,
           backend=None):
    with tempfile.TemporaryDirectory() as td:
        pdf2 = docx_to_pdf(docx_path, td)
        if pdf2 is None:
            return {"available": False, "rows": []}
        rows = compare(src_pdf, pdf2, out_dir=out_dir, backend=backend)
        keep = os.path.join(out_dir, "converted.pdf") if out_dir else None
        if keep:
            import shutil
            shutil.copy(pdf2, keep)
    mean = sum(r.get("ssim", 0) for r in rows) / max(1, len(rows))
    return {"available": True, "rows": rows, "mean_ssim": round(mean, 4)}
