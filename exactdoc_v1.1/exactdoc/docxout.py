"""DocLayout -> .docx writer (Google Docs-safe vocabulary).

Constructs used: styled paragraphs, fixed-layout tables with per-side borders
and cell shading, section geometry + true column sections, inline images,
headers/footers with PAGE/NUMPAGES fields, hyperlinks, tab stops.
No floating text boxes, no embedded fonts, no VML.
"""
import io
import re
from typing import Optional, List

import fitz
from docx import Document
from docx.shared import Pt, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .layout import (DocLayout, Para, Run, Cell, TableEl, FigureEl, ImageEl,
                     RuleEl, ColBreak, HFPart)
from .fonts import map_font

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
TABAL = {"left": WD_TAB_ALIGNMENT.LEFT, "center": WD_TAB_ALIGNMENT.CENTER,
         "right": WD_TAB_ALIGNMENT.RIGHT}


def _hex(c: str) -> str:
    return (c or "#000000").lstrip("#").upper()


def _set_borders(el_pr, borders: dict, tag: str):
    """Apply border dict to tcPr/pPr. tag='w:tcBorders' or 'w:pBdr'."""
    bel = OxmlElement(tag)
    order = ("top", "left", "bottom", "right")
    for side in order:
        spec = borders.get(side)
        b = OxmlElement("w:" + side)
        if spec:
            w, color = spec[0], spec[1]
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(max(2, int(round(w * 8)))))
            b.set(qn("w:space"), str(int(spec[2])) if len(spec) > 2 else "0")
            b.set(qn("w:color"), _hex(color))
        else:
            b.set(qn("w:val"), "nil")
        bel.append(b)
    el_pr.append(bel)


def _style_run(r, run: Run):
    f = r.font
    fam = map_font(run.font, mono=run.mono, serif=run.serif)
    f.name = fam
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), fam)
    f.size = Pt(round(run.size * 2) / 2)
    f.bold = run.bold
    f.italic = run.italic
    if run.underline:
        f.underline = True
    if run.superscript:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)
    try:
        f.color.rgb = RGBColor.from_string(_hex(run.color))
    except Exception:
        pass


def _add_field(par, instr: str, sample: str, style_from: Run):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " %s " % instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = sample
    r.append(t)
    fld.append(r)
    par._p.append(fld)
    # style the inner run
    from docx.text.run import Run as DRun
    dr = DRun(r, par)
    _style_run(dr, style_from)


def _add_hyperlink(par, url: str, runs_and_styles):
    part = par.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    par._p.append(h)
    from docx.text.run import Run as DRun
    for text, style in runs_and_styles:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        h.append(r)
        _style_run(DRun(r, par), style)


def write_para(container, p: Para, content_w: float, par=None):
    """Write a Para into container (doc/cell/header). Returns the paragraph."""
    if par is None:
        par = container.add_paragraph()
    pf = par.paragraph_format
    par.alignment = ALIGN.get(p.align, WD_ALIGN_PARAGRAPH.LEFT)
    if p.space_before > 0.05:
        pf.space_before = Pt(round(p.space_before, 1))
    else:
        pf.space_before = Pt(0)
    pf.space_after = Pt(round(max(0.0, p.space_after), 1))
    if p.leading and p.leading > 1:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(round(p.leading, 1))
    if p.left_indent > 0.05:
        pf.left_indent = Pt(round(p.left_indent, 1))
    if abs(p.first_indent) > 0.05:
        pf.first_line_indent = Pt(round(p.first_indent, 1))
    if p.right_indent > 0.05:
        pf.right_indent = Pt(round(p.right_indent, 1))
    for pos, al in p.tab_stops:
        pf.tab_stops.add_tab_stop(Pt(round(pos, 1)), TABAL.get(al, WD_TAB_ALIGNMENT.LEFT))
    # keep heading with following content
    if p.heading:
        pf.keep_with_next = True
        ppr = par._p.get_or_add_pPr()
        lvl = OxmlElement("w:outlineLvl")
        lvl.set(qn("w:val"), str(min(8, p.heading - 1)))
        ppr.append(lvl)
    # paragraph borders from header/footer rules
    bt = getattr(p, "border_top", None)
    bb = getattr(p, "border_bottom", None)
    if bt or bb:
        ppr = par._p.get_or_add_pPr()
        bd = {}
        if bt:
            bd["top"] = bt
        if bb:
            bd["bottom"] = bb
        _set_borders(ppr, bd, "w:pBdr")

    i = 0
    runs = p.runs
    while i < len(runs):
        run = runs[i]
        if run.link:
            grp = []
            while i < len(runs) and runs[i].link == run.link:
                grp.append((runs[i].text, runs[i]))
                i += 1
            _add_hyperlink(par, run.link, grp)
            continue
        if run.field:
            _add_field(par, run.field, "1", run)
            i += 1
            continue
        if run.is_tab:
            r = par.add_run()
            r.add_tab()
            _style_run(r, run)
            i += 1
            continue
        # split on newlines -> soft breaks
        parts = run.text.split("\n")
        for j, chunk in enumerate(parts):
            if chunk:
                r = par.add_run(chunk)
                _style_run(r, run)
            if j < len(parts) - 1:
                br = par.add_run()
                br.add_break(WD_BREAK.LINE)
                _style_run(br, run)
        i += 1
    return par


def _spacer(container, height_pt: float):
    """Tiny exact-height paragraph used as vertical spacing before tables."""
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(max(1.0, round(height_pt, 1)))
    r = par.add_run("")
    r.font.size = Pt(1)
    return par


def write_table(container, t: TableEl, content_w: float):
    n_rows = len(t.rows)
    n_cols = len(t.col_widths)
    if n_rows == 0 or n_cols == 0:
        return None
    if t.space_before > 0.5:
        _spacer(container, t.space_before)
    try:
        tbl = container.add_table(rows=n_rows, cols=n_cols)
    except TypeError:  # header/footer/cell containers require a width argument
        tbl = container.add_table(rows=n_rows, cols=n_cols,
                                  width=Emu(int(sum(t.col_widths) * 12700)))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = tbl._tbl.tblPr
    # fixed layout
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(int(round(sum(t.col_widths) * 20))))
    tw.set(qn("w:type"), "dxa")
    tblPr.append(tw)
    if t.left_indent > 0.5:
        ind = OxmlElement("w:tblInd")
        ind.set(qn("w:w"), str(int(round(t.left_indent * 20))))
        ind.set(qn("w:type"), "dxa")
        tblPr.append(ind)
    # no default borders / spacing; zero default cell margins
    tb = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + side)
        b.set(qn("w:val"), "nil")
        tb.append(b)
    tblPr.append(tb)
    mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement("w:" + side)
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tblPr.append(mar)
    # grid
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc, wpt in zip(grid.findall(qn("w:gridCol")), t.col_widths):
        gc.set(qn("w:w"), str(int(round(wpt * 20))))

    for ri, rowspec in enumerate(t.rows):
        row = tbl.rows[ri]
        h = t.row_heights[ri] if ri < len(t.row_heights) else None
        # Only pin height on rows with no text: text rows are content-driven
        # (cell pads + exact-leading paragraphs sum to the source height,
        # which renders identically in Word, Google Docs and LibreOffice).
        row_has_text = any(c and any(p.text.strip() for p in c.paras)
                           for c in rowspec)
        if h and not row_has_text:
            trPr = row._tr.get_or_add_trPr()
            th = OxmlElement("w:trHeight")
            th.set(qn("w:val"), str(int(round(h * 20))))
            th.set(qn("w:hRule"), "atLeast")
            trPr.append(th)
        for ci in range(n_cols):
            spec = rowspec[ci] if ci < len(rowspec) else None
            cell = tbl.cell(ri, ci)
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcw = OxmlElement("w:tcW")
            tcw.set(qn("w:w"), str(int(round(t.col_widths[ci] * 20))))
            tcw.set(qn("w:type"), "dxa")
            tcPr.append(tcw)
            if spec is None:
                _blank_cell(cell)
                _set_borders(tcPr, {}, "w:tcBorders")
                continue
            if spec.shading:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), _hex(spec.shading))
                tcPr.append(shd)
            _set_borders(tcPr, spec.borders, "w:tcBorders")
            tmar = OxmlElement("w:tcMar")
            pads = spec.pad  # (top, left, bottom, right)
            for side, val in zip(("top", "left", "bottom", "right"),
                                 (pads[0], pads[1], pads[2], pads[3])):
                m = OxmlElement("w:" + side)
                m.set(qn("w:w"), str(max(0, int(round(val * 20)))))
                m.set(qn("w:type"), "dxa")
                tmar.append(m)
            tcPr.append(tmar)
            if spec.valign and spec.valign != "top":
                va = OxmlElement("w:vAlign")
                va.set(qn("w:val"), spec.valign)
                tcPr.append(va)
            if spec.paras:
                first = cell.paragraphs[0]
                write_para(cell, spec.paras[0], t.col_widths[ci], par=first)
                for p in spec.paras[1:]:
                    write_para(cell, p, t.col_widths[ci])
            else:
                _blank_cell(cell)
    return tbl


def _blank_cell(cell):
    par = cell.paragraphs[0]
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(2)
    r = par.add_run("")
    r.font.size = Pt(1)


def write_figure(container, fig: FigureEl, src_doc, dpi: int = 240):
    page = src_doc[fig.page_no - 1]
    clip = fitz.Rect(*fig.clip)
    pix = page.get_pixmap(clip=clip, dpi=dpi, alpha=False)
    data = pix.tobytes("png")
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(round(max(0.0, fig.space_before), 1))
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    pf.line_spacing = Pt(round(fig.height, 1))
    par.alignment = ALIGN.get(fig.align, WD_ALIGN_PARAGRAPH.CENTER)
    if fig.align == "left" and fig.left_indent > 0.5:
        pf.left_indent = Pt(round(fig.left_indent, 1))
    r = par.add_run()
    r.add_picture(io.BytesIO(data), width=Emu(int(fig.width * 12700)),
                  height=Emu(int(fig.height * 12700)))
    return par


def write_image(container, im: ImageEl):
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(round(max(0.0, im.space_before), 1))
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    pf.line_spacing = Pt(round(im.height, 1))
    par.alignment = ALIGN.get(im.align, WD_ALIGN_PARAGRAPH.CENTER)
    if im.align == "left" and im.left_indent > 0.5:
        pf.left_indent = Pt(round(im.left_indent, 1))
    r = par.add_run()
    r.add_picture(io.BytesIO(im.data), width=Emu(int(im.width * 12700)),
                  height=Emu(int(im.height * 12700)))
    return par


def write_rule(container, rule: RuleEl, content_w: float):
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(round(max(0.0, rule.space_before), 1))
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(2)
    length = rule.length or (content_w * rule.width_pct / 100.0)
    if rule.left_indent > 0.5:
        pf.left_indent = Pt(round(rule.left_indent, 1))
    ri = content_w - rule.left_indent - length
    if ri > 0.5:
        pf.right_indent = Pt(round(ri, 1))
    r = par.add_run("")
    r.font.size = Pt(1)
    ppr = par._p.get_or_add_pPr()
    _set_borders(ppr, {"bottom": (rule.thickness, rule.color)}, "w:pBdr")
    return par


# ------------------------------------------------------------------ sections
def _config_section(sec, lay: DocLayout, margin_t=None, cols: int = 1,
                    col_gap: float = 24.0, margin_lr=None):
    sec.page_width = Emu(int(lay.page_w * 12700))
    sec.page_height = Emu(int(lay.page_h * 12700))
    ml = lay.margin_l if margin_lr is None else margin_lr
    mr = lay.margin_r if margin_lr is None else margin_lr
    sec.left_margin = Emu(int(ml * 12700))
    sec.right_margin = Emu(int(mr * 12700))
    sec.top_margin = Emu(int((lay.margin_t if margin_t is None else margin_t) * 12700))
    sec.bottom_margin = Emu(int(lay.margin_b * 12700))
    hd = lay.header_default.distance if lay.header_default else 36.0
    fd = lay.footer_default.distance if lay.footer_default else 36.0
    sec.header_distance = Emu(int(max(0.0, hd) * 12700))
    sec.footer_distance = Emu(int(max(0.0, fd) * 12700))
    sectPr = sec._sectPr
    cols_el = sectPr.find(qn("w:cols"))
    if cols_el is None:
        cols_el = OxmlElement("w:cols")
        sectPr.append(cols_el)
    if cols > 1:
        cols_el.set(qn("w:num"), str(cols))
        cols_el.set(qn("w:space"), str(int(round(col_gap * 20))))
        cols_el.set(qn("w:equalWidth"), "1")
    else:
        cols_el.set(qn("w:num"), "1")
        for a in ("w:space", "w:equalWidth"):
            if cols_el.get(qn(a)):
                cols_el.attrib.pop(qn(a))


def _shifted_part(part: Optional[HFPart], dl: float, dr: float) -> Optional[HFPart]:
    """Clone a header/footer part with indents/tabs shifted (bleed sections)."""
    if part is None:
        return None
    import copy
    np = copy.deepcopy(part)
    for el in np.elements:
        if isinstance(el, Para):
            el.left_indent = round(el.left_indent + dl, 1)
            el.right_indent = round((el.right_indent or 0.0) + dr, 1)
            el.tab_stops = [(round(p + dl, 1), a) for p, a in el.tab_stops]
        elif isinstance(el, TableEl):
            el.left_indent = round(el.left_indent + dl, 1)
    return np


def _fill_hf(hf_obj, part: Optional[HFPart], lay: DocLayout):
    """Fill a python-docx header/footer object with an HFPart."""
    hf_obj.is_linked_to_previous = False
    # clear default paragraph content
    first = hf_obj.paragraphs[0]
    first.clear()
    if part is None or not part.elements:
        pf = first.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(2)
        return
    used_first = False
    for el in part.elements:
        if isinstance(el, TableEl):
            write_table(hf_obj, el, lay.content_w)
        elif isinstance(el, Para):
            if not used_first:
                write_para(hf_obj, el, lay.content_w, par=first)
                used_first = True
            else:
                write_para(hf_obj, el, lay.content_w)
        elif isinstance(el, RuleEl):
            write_rule(hf_obj, el, lay.content_w)
    if not used_first:
        # first paragraph unused: make it invisible
        pf = first.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(2)


# ------------------------------------------------------------------ main
def write_docx(lay: DocLayout, out_path: str, dpi: int = 240) -> str:
    doc = Document()
    src_doc = fitz.open(lay.src_path) if lay.src_path else None
    content_w = lay.content_w

    # neutralize the template's Normal style (1.08 line, 8pt after) so nothing
    # inherits spacing we didn't ask for
    try:
        npf = doc.styles["Normal"].paragraph_format
        npf.space_before = Pt(0)
        npf.space_after = Pt(0)
        npf.line_spacing = 1.0
    except Exception:
        pass
    if lay.hyphenated:
        # source justifies with hyphenation: let Word/Docs hyphenate too so
        # line packing (and therefore paragraph heights) stay comparable
        st = doc.settings.element
        if st.find(qn("w:autoHyphenation")) is None:
            ah = OxmlElement("w:autoHyphenation")
            ah.set(qn("w:val"), "1")
            st.append(ah)

    sec = doc.sections[0]
    has_cover = lay.cover_band is not None
    band_bleed = 4.0  # cover band section: near-zero side margins => full-bleed band
    if has_cover:
        _config_section(sec, lay, margin_t=lay.cover_top, cols=1, margin_lr=band_bleed)
        sec.header_distance = Emu(0)  # body must start flush at the band
    else:
        _config_section(sec, lay, cols=1)

    # headers/footers for section 1 (never create empty parts: an empty header
    # still reserves a line and pushes the body down)
    if has_cover:
        dl = lay.margin_l - band_bleed
        dr = lay.margin_r - band_bleed
        if lay.header_first is not None:
            _fill_hf(sec.header, _shifted_part(lay.header_first, dl, dr), lay)
        if (lay.footer_first or lay.footer_default) is not None:
            _fill_hf(sec.footer,
                     _shifted_part(lay.footer_first or lay.footer_default, dl, dr), lay)
    else:
        if lay.header_default is not None:
            _fill_hf(sec.header, lay.header_default, lay)
        if lay.footer_default is not None:
            _fill_hf(sec.footer, lay.footer_default, lay)
        if lay.different_first:
            sec.different_first_page_header_footer = True
            _fill_hf(sec.first_page_header, lay.header_first, lay)
            _fill_hf(sec.first_page_footer,
                     lay.footer_first or lay.footer_default, lay)

    cur_cols = 1
    # config of the currently-open section; re-applied after each break because
    # python-docx's add_section clones sectPr elements in ways that can shuffle
    # previously-applied properties between sections
    cur_cfg = {"margin_t": (lay.cover_top if has_cover else None), "cols": 1,
               "gap": 24.0, "margin_lr": (band_bleed if has_cover else None),
               "hdr0": has_cover}

    def new_section(kind, cols, gap=24.0, margin_t=None, margin_lr=None):
        nonlocal sec, cur_cols, cur_cfg
        doc.add_section(kind)
        # re-apply the finished section's geometry to whatever element now
        # represents it
        fin = doc.sections[-2]
        _config_section(fin, lay, margin_t=cur_cfg["margin_t"], cols=cur_cfg["cols"],
                        col_gap=cur_cfg["gap"], margin_lr=cur_cfg["margin_lr"])
        if cur_cfg.get("hdr0"):
            fin.header_distance = Emu(0)
        sec = doc.sections[-1]
        _config_section(sec, lay, margin_t=margin_t, cols=cols, col_gap=gap,
                        margin_lr=margin_lr)
        cur_cfg = {"margin_t": margin_t, "cols": cols, "gap": gap,
                   "margin_lr": margin_lr, "hdr0": False}
        cur_cols = cols
        # crush section-break paragraphs so they consume no vertical space
        for p_el in doc.element.body.findall(qn("w:p")):
            ppr = p_el.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                if ppr.find(qn("w:spacing")) is None:
                    sp = OxmlElement("w:spacing")
                    sp.set(qn("w:before"), "0")
                    sp.set(qn("w:after"), "0")
                    sp.set(qn("w:line"), "20")
                    sp.set(qn("w:lineRule"), "exact")
                    ppr.append(sp)
        return sec

    if has_cover:
        # The whole cover page lives in ONE bleed-margin section (renderers do
        # not honor L/R margin changes at mid-page continuous breaks). The band
        # spans the bleed width; every other page-1 element is shifted right by
        # indents so it keeps its original x-position.
        delta_l = lay.margin_l - band_bleed
        delta_r = lay.margin_r - band_bleed
        band = lay.cover_band
        band.col_widths = [lay.page_w - 2 * band_bleed]
        for c in band.rows[0]:
            if c:
                c.pad = (c.pad[0], round(c.pad[1] + delta_l, 1), c.pad[2], c.pad[3])
        for ch in lay.pages[0].chunks:
            for el in ch.elements:
                if isinstance(el, Para):
                    el.left_indent = round(el.left_indent + delta_l, 1)
                    el.right_indent = round((el.right_indent or 0.0) + delta_r, 1)
                    el.tab_stops = [(round(p + delta_l, 1), a) for p, a in el.tab_stops]
                elif isinstance(el, TableEl):
                    el.left_indent = round(el.left_indent + delta_l, 1)
                elif isinstance(el, (FigureEl, ImageEl)):
                    if el.align == "left":
                        el.left_indent = round(el.left_indent + delta_l, 1)
                elif isinstance(el, RuleEl):
                    el.left_indent = round(el.left_indent + delta_l, 1)
        write_table(doc, band, lay.page_w - 2 * band_bleed)

    last_el_par = None
    for pi, pg in enumerate(lay.pages):
        if pi > 0:
            # page boundary
            after_cover = has_cover and pi == 1
            next_cols = pg.chunks[0].n_cols if pg.chunks else 1
            if after_cover or cur_cols != next_cols:
                gap = pg.chunks[0].col_gap if pg.chunks else 24.0
                pre = pg.chunks[0].pre_gap if pg.chunks else 0.0
                mt = (lay.margin_t + pre) if (next_cols > 1 and pre > 0.5) else None
                s = new_section(WD_SECTION.NEW_PAGE, next_cols, gap, margin_t=mt)
                if after_cover:
                    _fill_hf(s.header, lay.header_default, lay)
                    _fill_hf(s.footer, lay.footer_default, lay)
            else:
                par = doc.add_paragraph()
                pf = par.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(1)
                par.add_run().add_break(WD_BREAK.PAGE)
        cw_ctx = (lay.page_w - 2 * band_bleed) if (has_cover and pi == 0) else content_w
        for ci, ch in enumerate(pg.chunks):
            if ch.n_cols != cur_cols:
                if ch.pre_gap > 0.5:
                    _spacer(doc, ch.pre_gap)
                new_section(WD_SECTION.CONTINUOUS, ch.n_cols, ch.col_gap)
            for el in ch.elements:
                if isinstance(el, ColBreak):
                    par = doc.add_paragraph()
                    pf = par.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    pf.line_spacing = Pt(1)
                    par.add_run().add_break(WD_BREAK.COLUMN)
                elif isinstance(el, Para):
                    write_para(doc, el, cw_ctx)
                elif isinstance(el, TableEl):
                    write_table(doc, el, cw_ctx)
                elif isinstance(el, FigureEl):
                    write_figure(doc, el, src_doc, dpi=dpi)
                elif isinstance(el, ImageEl):
                    write_image(doc, el)
                elif isinstance(el, RuleEl):
                    write_rule(doc, el, cw_ctx)

    # drop the initial empty paragraph python-docx puts in a fresh document
    # (never touch section-break paragraphs: removing one deletes a section)
    body = doc.element.body
    paras = body.findall(qn("w:p"))
    if paras:
        p0 = paras[0]
        has_content = p0.findall(qn("w:r")) or p0.findall(qn("w:hyperlink")) \
            or p0.findall(qn("w:fldSimple"))
        ppr = p0.find(qn("w:pPr"))
        has_sectpr = ppr is not None and ppr.find(qn("w:sectPr")) is not None
        if not has_content and not has_sectpr and len(list(body)) > 2:
            body.remove(p0)

    if src_doc is not None:
        src_doc.close()
    doc.save(out_path)
    return out_path
