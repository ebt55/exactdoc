"""Probe: what does Docs put between a page-leading band table and the body?

    python testkit/probe_band_gap.py build [out.docx]
    python testkit/probe_band_gap.py analyse <google-export.pdf>

WHY THIS EXISTS
---------------
01_whitepaper_market's page-1 SSIM is the last blocking finding, and its text
below the cover band sits low. Measured on BASELINES in the pass-3/4 exports --
not on glyph tops, see the note below -- the band's own three lines land within
0.6pt of the source, and then the very next element jumps 3.7pt. So nothing
inside the band is wrong; something at its lower boundary is.

That 3.7pt splits into two quantities the writer cannot currently account for:

    +1.12pt  between the band's last text baseline and the band's bottom edge,
             where the only thing requested is the cell's tcMar bottom (24pt)
             and the last line's descent.
    +2.77pt  between the band's bottom edge and the following paragraph's
             baseline, where the only thing requested is that paragraph's
             space_before (36.8pt) and its own line-box ascent.

Both are computed against a line-box model that is otherwise accurate to 0.31pt
on this document, so they are real and not modelling slack. This probe isolates
each one by varying it alone.

ON MEASURING BASELINES
----------------------
Glyph tops cannot be used here. PyMuPDF derives a span's bbox top from the
font's own reported ascender, and the source's Helvetica reports 1.070-1.075 em
where Docs' ArialMT reports 0.905 -- so the two disagree about where a line
starts by 0.17 x size even when the baseline is identical. That is 1.8pt at
10.5pt and 4.1pt at 25pt, and it lands in a glyph-top comparison as drift that
never happened. Baselines are what the renderer actually positioned.

WHAT IT VARIES
--------------
One variable per page, in the style of probe_cover_band.py:

    tcmar  band cell tcMar bottom in {0, 12, 24}pt, following space_before 24
    space  following paragraph space_before in {0, 18, 36}pt, tcMar bottom 12
    extra  the constructs 01 might be adding on top: a spacer paragraph between
           the table and the body, and a bottom border on the band cell
    plain  no table at all -- paragraph to paragraph across the same gap, which
           is the control that says whether the table is implicated

Every page carries a marker, a known last line INSIDE the band, and a known
first line after it, so the analyser can read both baselines and the band's
rectangle without guessing which line is which.

HOW TO RUN IT (rides along with the next consented pass -- uploads nothing)
--------------------------------------------------------------------------
1.  python testkit/probe_band_gap.py build
2.  Upload during the next CONSENTED pass and export the PDF.
3.  python testkit/probe_band_gap.py analyse <export.pdf>

Shares no code with the converter, like every other probe here.
"""
import argparse
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

PAGE_W, PAGE_H = 612.0, 792.0
MARGIN = 54.0
BAND_FILL = "1F3864"
BAND_TEXT = "BANDLASTLINE"
BODY_TEXT = "BODYFIRSTLINE"
BAND_SIZE = 9.5          # matches 01's last band line
BODY_SIZE = 16.0         # matches 01's first body element
MARK_RE = re.compile(r"BGPROBE(\d{2})")

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "quirks")
DEFAULT_OUT = os.path.join(OUT_DIR, "probe_band_gap.docx")

# (family, tcmar_bottom_pt, space_before_pt, spacer_pt, border, has_band)
VARIANTS = (
    ("tcmar", 0.0, 24.0, None, False, True),
    ("tcmar", 12.0, 24.0, None, False, True),
    ("tcmar", 24.0, 24.0, None, False, True),
    ("space", 12.0, 0.0, None, False, True),
    ("space", 12.0, 18.0, None, False, True),
    ("space", 12.0, 36.0, None, False, True),
    ("extra", 12.0, 24.0, 6.0, False, True),
    ("extra", 12.0, 24.0, None, True, True),
    ("plain", 12.0, 24.0, None, False, False),
)


def marker(i):
    return "BGPROBE%02d" % i


def parse_marker(text):
    m = MARK_RE.search(text or "")
    return int(m.group(1)) if m else None


def label(i):
    fam, tcmar, before, spacer, border, band = VARIANTS[i]
    bits = ["tcMar=%.0f" % tcmar, "before=%.0f" % before]
    if spacer:
        bits.append("spacer=%.0f" % spacer)
    if border:
        bits.append("border")
    if not band:
        bits.append("NO-BAND")
    return "%s %s" % (fam, " ".join(bits))


# ------------------------------------------------------------------ build
def _run(par, text, size, family="Arial"):
    r = par.add_run(text)
    r.font.size = Pt(size)
    r.font.name = family
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), family)
    return r


def _neutral(par, before=0.0, after=0.0):
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0
    return par


def _band(doc, width, tcmar_bottom, border):
    tbl = doc.add_table(rows=1, cols=1)
    tblPr = tbl._tbl.tblPr
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(int(round(width * 20))))
    tw.set(qn("w:type"), "dxa")
    tblPr.append(tw)
    tb = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + side)
        b.set(qn("w:val"), "nil")
        tb.append(b)
    tblPr.append(tb)
    mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement("w:" + side)
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tblPr.append(mar)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), str(int(round(width * 20))))

    cell = tbl.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcw = OxmlElement("w:tcW")
    tcw.set(qn("w:w"), str(int(round(width * 20))))
    tcw.set(qn("w:type"), "dxa")
    tcPr.append(tcw)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), BAND_FILL)
    tcPr.append(shd)
    tmar = OxmlElement("w:tcMar")
    for side, val in (("top", 18.0), ("left", 12.0),
                      ("bottom", tcmar_bottom), ("right", 0.0)):
        m = OxmlElement("w:" + side)
        m.set(qn("w:w"), str(max(0, int(round(val * 20)))))
        m.set(qn("w:type"), "dxa")
        tmar.append(m)
    tcPr.append(tmar)
    if border:
        bd = OxmlElement("w:tcBorders")
        for side in ("top", "left", "right"):
            e = OxmlElement("w:" + side)
            e.set(qn("w:val"), "nil")
            bd.append(e)
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "8")
        bot.set(qn("w:space"), "0")
        bot.set(qn("w:color"), "E5A000")
        bd.append(bot)
        tcPr.append(bd)

    par = _neutral(cell.paragraphs[0])
    _run(par, BAND_TEXT, BAND_SIZE)
    return tbl


def build(out_path=None):
    out_path = out_path or DEFAULT_OUT
    doc = Document()
    npf = doc.styles["Normal"].paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)
    npf.line_spacing = 1.0

    for i, (fam, tcmar, before, spacer, border, has_band) in enumerate(VARIANTS):
        if i == 0:
            sec = doc.sections[0]
        else:
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
        sec.page_width = Pt(PAGE_W)
        sec.page_height = Pt(PAGE_H)
        for attr in ("left_margin", "right_margin", "top_margin",
                     "bottom_margin"):
            setattr(sec, attr, Pt(MARGIN))
        sec.header_distance = Pt(0)

        head = _neutral(doc.add_paragraph())
        _run(head, "%s  %s" % (marker(i), label(i)), 8.0)

        if has_band:
            _band(doc, PAGE_W - 2 * MARGIN, tcmar, border)
        else:
            # the control: a paragraph standing in for the band, so the same
            # gap is measured without a table anywhere near it
            p = _neutral(doc.add_paragraph(), before=18.0, after=tcmar)
            _run(p, BAND_TEXT, BAND_SIZE)
        if spacer:
            sp = _neutral(doc.add_paragraph())
            pf = sp.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(spacer)
            _run(sp, "", 1.0)
        body = _neutral(doc.add_paragraph(), before=before)
        _run(body, BODY_TEXT, BODY_SIZE)

    parent = os.path.dirname(os.path.abspath(out_path))
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------- analyse
def analyse(pdf_path):
    """Per page: band rect, the two baselines, and the two residuals."""
    import fitz

    doc = fitz.open(pdf_path)
    rows = []
    try:
        for pno in range(doc.page_count):
            page = doc[pno]
            idx = parse_marker(page.get_text("text"))
            row = {"page": pno + 1, "variant": None}
            if idx is None or idx >= len(VARIANTS):
                rows.append(row)
                continue
            fam, tcmar, before, spacer, border, has_band = VARIANTS[idx]
            row.update(variant=label(idx), family=fam, tcmar=tcmar,
                       before=before, spacer=spacer, has_band=has_band)

            for b in page.get_text("dict")["blocks"]:
                if b.get("type"):
                    continue
                for ln in b["lines"]:
                    for s in ln["spans"]:
                        t = s["text"].strip()
                        if t.startswith(BAND_TEXT):
                            row["band_base"] = s["origin"][1]
                            row["band_desc"] = s["bbox"][3] - s["origin"][1]
                            row["band_size"] = s["size"]
                        elif t.startswith(BODY_TEXT):
                            row["body_base"] = s["origin"][1]
                            row["body_asc"] = s["origin"][1] - s["bbox"][1]
                            row["body_size"] = s["size"]

            best = None
            for dr in page.get_drawings():
                f = dr.get("fill")
                if not f or sum(f[:3]) / 3.0 > 0.55:
                    continue
                r = dr["rect"]
                if r.get_area() < 2000 or r.width < 200:
                    continue
                if best is None or r.get_area() > best.get_area():
                    best = r
            if best is not None:
                row["band_top"] = best.y0
                row["band_bottom"] = best.y1

            # residual 1: below the band's last baseline, what exceeds the
            # requested tcMar plus that line's own descent
            if "band_bottom" in row and "band_base" in row:
                got = row["band_bottom"] - row["band_base"]
                row["below_line"] = got
                row["resid_bottom"] = got - (tcmar + row["band_desc"])
            # residual 2: from the band's bottom to the body baseline, what
            # exceeds the requested space_before plus that line's own ascent
            if "body_base" in row:
                anchor = row.get("band_bottom")
                if anchor is None and "band_base" in row:
                    anchor = row["band_base"] + row.get("band_desc", 0)
                if anchor is not None:
                    got = row["body_base"] - anchor
                    row["gap"] = got
                    row["resid_gap"] = got - (before + (spacer or 0.0)
                                              + row["body_asc"])
            rows.append(row)
    finally:
        doc.close()
    return rows


def render(rows):
    L = ["%-34s %9s %9s %10s %10s"
         % ("variant", "below", "resid", "gap", "resid"), "-" * 76]
    for r in rows:
        if not r.get("variant"):
            L.append("page %d: no probe marker found" % r["page"])
            continue
        L.append("%-34s %9s %9s %10s %10s" % (
            r["variant"],
            "-" if "below_line" not in r else "%.2f" % r["below_line"],
            "-" if "resid_bottom" not in r else "%+.2f" % r["resid_bottom"],
            "-" if "gap" not in r else "%.2f" % r["gap"],
            "-" if "resid_gap" not in r else "%+.2f" % r["resid_gap"]))
    L.append("-" * 76)
    L.append("resid = what Docs added beyond what the OOXML asked for.")
    L.append("01_whitepaper needs +1.12 below the band and +2.77 in the gap")
    L.append("explained; a residual that reproduces those is the mechanism.")

    banded = [r["resid_gap"] for r in rows
              if r.get("has_band") and "resid_gap" in r]
    plain = [r["resid_gap"] for r in rows
             if r.get("has_band") is False and "resid_gap" in r]
    if banded and plain:
        mb = sorted(banded)[len(banded) // 2]
        L.append("")
        L.append("table present: median gap residual %+.2f" % mb)
        L.append("no table     : gap residual %+.2f" % plain[0])
        L.append("A difference here says the TABLE boundary is the cause; the")
        L.append("same residual both ways says it is ordinary paragraph space.")
    by_tcmar = [(r["tcmar"], r.get("resid_bottom")) for r in rows
                if r.get("family") == "tcmar" and r.get("resid_bottom") is not None]
    if len(by_tcmar) > 1:
        L.append("")
        L.append("tcMar bottom requested -> residual: %s"
                 % ", ".join("%.0f:%+.2f" % t for t in by_tcmar))
        L.append("A residual flat across the three says Docs adds a constant;")
        L.append("one that grows with the request says it scales it.")
    return "\n".join(L)


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    sub = ap.add_subparsers(dest="cmd", required=True)
    b = sub.add_parser("build", help="write the probe DOCX (uploads nothing)")
    b.add_argument("out", nargs="?", default=DEFAULT_OUT)
    a = sub.add_parser("analyse", help="read a Google export of the probe")
    a.add_argument("pdf")
    args = ap.parse_args(argv)

    if args.cmd == "build":
        path = build(args.out)
        print("wrote %s (%d variants)" % (path, len(VARIANTS)))
        print("Upload during the next CONSENTED pass, export PDF, then:")
        print("  python testkit/probe_band_gap.py analyse <export.pdf>")
        return 0
    if not os.path.isfile(args.pdf):
        sys.stderr.write("error: no such file: %s\n" % args.pdf)
        return 2
    print(render(analyse(args.pdf)))
    return 0


if __name__ == "__main__":
    sys.exit(main())
