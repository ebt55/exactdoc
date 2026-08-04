"""Probe: what floor does Google Docs enforce above a page-leading bleed band?

    python testkit/probe_cover_band.py build [out.docx]
    python testkit/probe_cover_band.py analyse <google-export.pdf>

WHY THIS EXISTS
---------------
01_whitepaper_market scores mean_ssim 0.6723, and page 1 (0.6183) is the worst
of its three even though its text is placed well (|dx|p50 0.58, |dy|p50 2.68)
and the page looks right side by side. Measured from Google's own export of the
2026-08-04 pass:

    source band   rows 0.0 .. 172.8pt   cols   0.0 .. 611.3pt   (true full bleed)
    Docs band     rows 14.4 .. 176.7pt  cols   3.9 .. 606.8pt

So Docs pushed the band down 14.4pt and left a 3.9pt white frame on each side.
3.9pt is exactly the `band_bleed = 4.0` the writer asks for as a side margin,
which says Docs honours that margin instead of bleeding. The 14.4pt is
unexplained: the writer already asks for `header_distance = 0` and a top margin
of `cover_top`. A large solid rectangle displaced by 14.4pt over the top quarter
of the page is a heavy SSIM penalty, and it is the one part of that page this
project has not accounted for.

`_GDOCS_COVER_BEFORE_COMP_TWIPS = 290` (14.5pt) in docxout is the same quantity
seen from inside: it pulls the band's first paragraph up, which repositions the
TEXT but cannot move the band rectangle. Compensating the rectangle needs to
know which knob Docs is clamping, and that is what this probe measures.

WHAT IT VARIES
--------------
One variable per page, everything else fixed, in the style of docs_quirks.py.

    top   family: section top margin in {0, 4, 8, 14.4, 20}pt, header_distance 0
    hdr   family: header_distance in {0, 8, 14.4}pt, top margin 0
    side  family: left/right margin in {0, 2, 4, 8}pt, top margin 0

If the rendered band top is `max(requested, floor)`, the top family reads the
floor off directly. If instead it tracks header_distance, the hdr family says
so and the fix is a header property rather than a margin. The side family
answers the other half: whether a true full-bleed band is reachable at all, or
whether every band keeps a white frame equal to the side margin.

Like testkit/docs_quirks.py this shares no code with the converter under test:
it builds the OOXML shape directly, so a writer bug cannot hide itself by being
present on both sides of the comparison.

HOW TO RUN IT (rides along with the next consented pass -- uploads nothing)
--------------------------------------------------------------------------
1.  python testkit/probe_cover_band.py build
    Writes testkit/quirks/probe_cover_band.docx -- the same gitignored output
    directory docs_quirks.py uses, so building the probe cannot dirty the tree.
2.  During the next CONSENTED live pass, upload that .docx to Google Docs the
    same way the corpus documents go up, and export it back as PDF. This step
    needs the user's consent; the probe never contacts anything by itself.
3.  python testkit/probe_cover_band.py analyse testkit/quirks/probe_cover_band.gdocs.pdf

Step 3 prints requested versus rendered band geometry per variant. Nothing here
touches the gate, the oracle, or the quality policy.
"""
import argparse
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Letter, matching the corpus fixtures.
PAGE_W, PAGE_H = 612.0, 792.0
BAND_H = 150.0
BAND_FILL = "1F3864"          # dark navy, unambiguous against white
MARK_RE = re.compile(r"PROBE(TOP|HDR|SIDE)(\d{4})")
# Same gitignored directory docs_quirks.py writes to: building a probe must not
# leave an untracked artefact in the working tree.
OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "quirks")
DEFAULT_OUT = os.path.join(OUT_DIR, "probe_cover_band.docx")

# (family, requested_pt) -- one page each, one variable each.
VARIANTS = (
    [("TOP", v) for v in (0.0, 4.0, 8.0, 14.4, 20.0)]
    + [("HDR", v) for v in (0.0, 8.0, 14.4)]
    + [("SIDE", v) for v in (0.0, 2.0, 4.0, 8.0)]
)


def _marker(family, value):
    """PROBETOP0144 == the TOP variant that asked for 14.4pt."""
    return "PROBE%s%04d" % (family, int(round(value * 10)))


def parse_marker(text):
    m = MARK_RE.search(text or "")
    if not m:
        return None
    return m.group(1), int(m.group(2)) / 10.0


# ------------------------------------------------------------------ build
def _configure(sec, top, side, header_distance):
    sec.page_width = Pt(PAGE_W)
    sec.page_height = Pt(PAGE_H)
    sec.top_margin = Pt(top)
    sec.bottom_margin = Pt(36)
    sec.left_margin = Pt(side)
    sec.right_margin = Pt(side)
    sec.header_distance = Pt(header_distance)
    sec.footer_distance = Pt(36)


def _band(doc, width_pt, label):
    """A one-cell shaded fixed-layout table: the writer's cover-band shape."""
    tbl = doc.add_table(rows=1, cols=1)
    tblPr = tbl._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(int(round(width_pt * 20))))
    tw.set(qn("w:type"), "dxa")
    tblPr.append(tw)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + side)
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)
    mar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement("w:" + side)
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tblPr.append(mar)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), str(int(round(width_pt * 20))))

    cell = tbl.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcw = OxmlElement("w:tcW")
    tcw.set(qn("w:w"), str(int(round(width_pt * 20))))
    tcw.set(qn("w:type"), "dxa")
    tcPr.append(tcw)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), BAND_FILL)
    tcPr.append(shd)
    tmar = OxmlElement("w:tcMar")
    for side, val in (("top", 24.0), ("left", 24.0), ("bottom", 24.0), ("right", 0.0)):
        m = OxmlElement("w:" + side)
        m.set(qn("w:w"), str(int(round(val * 20))))
        m.set(qn("w:type"), "dxa")
        tmar.append(m)
    tcPr.append(tmar)

    # Pin the height so a clamped top margin cannot be mistaken for a band that
    # simply grew: the rectangle's height is then a constant of the probe.
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:trHeight")
    th.set(qn("w:val"), str(int(round(BAND_H * 20))))
    th.set(qn("w:hRule"), "atLeast")
    trPr.append(th)

    par = cell.paragraphs[0]
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(24)
    run = par.add_run(label)
    run.font.size = Pt(18)
    run.font.name = "Arial"
    run.font.bold = True
    return tbl


def _crush(par):
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(1)


def build(out_path):
    """Write the probe DOCX. Returns the path."""
    doc = Document()
    npf = doc.styles["Normal"].paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)
    npf.line_spacing = 1.0

    for i, (family, value) in enumerate(VARIANTS):
        top = value if family == "TOP" else 0.0
        side = value if family == "SIDE" else 4.0
        hdr = value if family == "HDR" else 0.0
        if i == 0:
            sec = doc.sections[0]
        else:
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
        _configure(sec, top, side, hdr)
        _band(doc, PAGE_W - 2 * side, _marker(family, value))
        # a crushed tail paragraph keeps the table off the section boundary
        _crush(doc.add_paragraph())

    # The template's leading empty paragraph would sit above the first band and
    # silently invalidate variant 1, which is the one that asks for 0pt.
    body = doc.element.body
    paras = body.findall(qn("w:p"))
    if paras:
        first = paras[0]
        ppr = first.find(qn("w:pPr"))
        has_sect = ppr is not None and ppr.find(qn("w:sectPr")) is not None
        if not first.findall(qn("w:r")) and not has_sect:
            body.remove(first)
    for p_el in body.findall(qn("w:p")):
        ppr = p_el.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            if ppr.find(qn("w:spacing")) is None:
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:before"), "0")
                sp.set(qn("w:after"), "0")
                sp.set(qn("w:line"), "20")
                sp.set(qn("w:lineRule"), "exact")
                ppr.append(sp)

    parent = os.path.dirname(os.path.abspath(out_path))
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------- analyse
def _dark_rects(page, min_area=5000.0):
    """Filled rectangles dark enough to be the band, largest first.

    Vector fills are read rather than pixels: the answer is a position in
    points, and rasterising first would quantise it to the dpi and then need
    converting back.
    """
    out = []
    for d in page.get_drawings():
        fill = d.get("fill")
        if not fill or d.get("fill_opacity", 1) == 0:
            continue
        if sum(fill[:3]) / 3.0 > 0.55:      # not a dark band
            continue
        r = d["rect"]
        if r.get_area() < min_area:
            continue
        out.append((r.get_area(), r))
    out.sort(key=lambda t: -t[0])
    return [r for _, r in out]


def analyse(pdf_path):
    """[{family, requested_pt, top, bottom, left, right, delta}] per page."""
    import fitz          # imported here so `build` works without PyMuPDF

    doc = fitz.open(pdf_path)
    rows = []
    try:
        for pno in range(doc.page_count):
            page = doc[pno]
            found = parse_marker(page.get_text("text"))
            rects = _dark_rects(page)
            row = {"page": pno + 1,
                   "family": found[0] if found else None,
                   "requested_pt": found[1] if found else None}
            if rects:
                r = rects[0]
                row.update({"top": round(r.y0, 2), "bottom": round(r.y1, 2),
                            "left": round(r.x0, 2), "right": round(r.x1, 2),
                            "height": round(r.y1 - r.y0, 2),
                            "page_w": round(page.rect.width, 2)})
                if found and found[0] in ("TOP", "HDR"):
                    row["delta"] = round(r.y0 - found[1], 2)
                elif found and found[0] == "SIDE":
                    row["delta"] = round(r.x0 - found[1], 2)
            rows.append(row)
    finally:
        doc.close()
    return rows


def render(rows):
    L = ["%-6s %10s %9s %9s %9s %9s %8s"
         % ("family", "requested", "top", "left", "right", "height", "delta"),
         "-" * 68]
    for r in rows:
        if r["family"] is None:
            L.append("page %d: no probe marker found" % r["page"])
            continue
        if "top" not in r:
            L.append("%-6s %10.1f   (no band rectangle found)"
                     % (r["family"], r["requested_pt"]))
            continue
        L.append("%-6s %10.1f %9.2f %9.2f %9.2f %9.2f %8.2f"
                 % (r["family"], r["requested_pt"], r["top"], r["left"],
                    r["right"], r["height"], r.get("delta", float("nan"))))
    L.append("-" * 68)
    tops = [r for r in rows if r.get("family") == "TOP" and "top" in r]
    if tops:
        floor = min(r["top"] for r in tops)
        honoured = [r for r in tops if abs(r.get("delta", 9e9)) <= 0.75]
        L.append("lowest band top achieved: %.2fpt (asked for %.1fpt)"
                 % (floor, min(r["requested_pt"] for r in tops)))
        L.append("variants Docs honoured exactly: %s"
                 % (", ".join("%.1f" % r["requested_pt"] for r in honoured)
                    or "none"))
        L.append("If every top clamps to the same value, that value is the floor "
                 "and the writer can stop asking for less.")
    sides = [r for r in rows if r.get("family") == "SIDE" and "left" in r]
    if sides:
        L.append("side frame: " + ", ".join(
            "%.1f->%.2f" % (r["requested_pt"], r["left"]) for r in sides))
    return "\n".join(L)


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    sub = ap.add_subparsers(dest="cmd", required=True)
    b = sub.add_parser("build", help="write the probe DOCX (uploads nothing)")
    b.add_argument("out", nargs="?", default=DEFAULT_OUT)
    a = sub.add_parser("analyse", help="read a Google export of the probe")
    a.add_argument("pdf")
    args = ap.parse_args(argv)

    if args.cmd == "build":
        path = build(args.out)
        print("wrote %s (%d variants)" % (path, len(VARIANTS)))
        print("Upload it during the next CONSENTED pass, export PDF, then:")
        print("  python testkit/probe_cover_band.py analyse <export.pdf>")
        return 0
    if not os.path.isfile(args.pdf):
        sys.stderr.write("error: no such file: %s\n" % args.pdf)
        return 2
    print(render(analyse(args.pdf)))
    return 0


if __name__ == "__main__":
    sys.exit(main())
