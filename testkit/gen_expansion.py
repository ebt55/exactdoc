"""Corpus expansion, tranche 1: the ordinary documents the frozen 16 under-sample.

    python testkit/gen_expansion.py /work/build      # generate into a scratch dir

**This does not touch the gate.** It writes nowhere near `testkit/fixtures/`,
knows nothing about `corpus_manifest.json`, and produces no gated number. It is
deliberately a separate script from `gen_corpus.py`: `tests/test_corpus_generation.py`
runs that one under `--strict` and fails on any document it did not expect, so
the expansion tranche must not be generated from the same entry point.

What it makes, and why these and not others. The 16 frozen fixtures spend 3
documents on designed stress (nested tables, CJK/RTL, gradients) and carry the
whole blocking `ordinary_digital` claim on 13 that are mostly one prose layout
with a table in it. Meanwhile the word-processor dialect -- the single most
common shape of PDF a user will actually hand this converter -- is one document,
`l1_word_native`. One sample cannot tell a general rule from a fixture accident.

So this tranche is 100% `ordinary_digital` and weights the thin producer:

    LibreOffice (word processor)   x01-x06    running heads, TOC, lists,
                                              plain tables, quotes/notes, scripts
    Chromium (browser print)       x07-x12    running heads, browser-furnished
                                              header/footer, lists, tables, TOC,
                                              European scripts
    ReportLab / fpdf2 (reports)    x13-x16    page-numbered reports, statement
                                              line items, handbook, bulletin

Every document is born-digital, left-to-right, rectangular-tabled, and uses only
scripts the pinned font set in `scripts/fonts.conf` covers without fallback --
which is exactly the `ordinary_digital` rule in `docs/corpus-expansion.md` §4.
The European-scripts documents are the deliberate contrast with `c4_i18n`:
Cyrillic and Greek are *ordinary* here because Liberation and DejaVu cover them,
where CJK and RTL are stress because they move under font fallback.

Output is a scratch directory plus `expansion_provenance.json`, which carries
the tier, dialect, rationale, licence and recipe of each document but **no
hashes**. `corpus_manifest.py expansion-seal` computes the identity. A human
authors the claims, the tool computes the hash, and neither can forge the other.
"""
import datetime
import json
import os
import shutil
import subprocess
import sys

import _paths  # noqa: F401  (sets sys.path, finds soffice/chrome)
from _paths import CHROME, SOFFICE

OUT = sys.argv[1] if len(sys.argv) > 1 else "expansion_build"
HTML = os.path.join(OUT, "_html")
WORK = os.path.join(OUT, "_work")

# Stamped into `provenance.license` on every fixture this script generates, so
# it must track the repository's own LICENSE. Changed with the Apache-2.0
# migration; the 16 already-generated entries in corpus_expansion.json were
# relabelled in the same commit, because an AGPL-labelled corpus inside an
# Apache-2.0 repository is a contradiction a reader trips over long before a
# lawyer does. Sole authorship is what makes relabelling a correction rather
# than a re-licence of someone else's work.
LICENSE = "Apache-2.0"
ACQUIRED = os.environ.get("EXPANSION_DATE") or datetime.date.today().isoformat()

# Prose deliberately unlike gen_corpus.py's retrieval/RAG material, so the two
# corpora do not share vocabulary and a text-recall metric cannot be flattered by
# having seen the same sentences sixteen times already.
P1 = ("The depot replacement programme was approved on the understanding that "
      "service levels would be maintained throughout construction. That "
      "assumption has not survived contact with the schedule: the temporary "
      "layover facility is smaller than the site it replaces, and peak-hour "
      "vehicle circulation is now the binding constraint. ")
P2 = ("Mitigation is available but unglamorous. Two additional relief runs, "
      "shifted twelve minutes earlier, absorb most of the shortfall at a cost "
      "well inside the contingency already held for this phase. The remainder "
      "is best handled by publishing the revised layover times rather than by "
      "adding vehicles. ")
P3 = ("Members should note that the figures below are drawn from the automatic "
      "passenger counters, which under-report boardings at the two rear doors "
      "by a margin the vendor states as three per cent. No adjustment has been "
      "applied; the raw counts are reported so that later revisions remain "
      "traceable to their source. ")
P4 = ("The recommendation is therefore to proceed, with the reporting "
      "requirement attached, and to bring a further report in the autumn once "
      "a full quarter of post-occupancy data is available for comparison. ")


# ------------------------------------------------------------------ producers
def soffice_pdf(src, out_pdf):
    """DOCX -> PDF through LibreOffice.

    The profile path must be ABSOLUTE. `-env:UserInstallation` takes a file URL
    and `file:///` + a relative path resolves against the filesystem root, at
    which point soffice cannot create its profile, exits 1, and writes nothing.
    `gen_corpus.py` records that this once silently produced a 15-document
    corpus that every downstream figure was then computed over.
    """
    prof = os.path.abspath(os.path.join(WORK, "_loprof"))
    outdir = os.path.dirname(os.path.abspath(out_pdf))
    r = subprocess.run([SOFFICE, "--headless", "--norestore",
                        "-env:UserInstallation=file:///" + prof.replace("\\", "/"),
                        "--convert-to", "pdf", "--outdir", outdir,
                        os.path.abspath(src)],
                       capture_output=True, timeout=300)
    made = os.path.join(outdir, os.path.splitext(os.path.basename(src))[0] + ".pdf")
    if os.path.exists(made):
        if os.path.abspath(made) != os.path.abspath(out_pdf):
            shutil.move(made, out_pdf)
        return out_pdf
    raise RuntimeError("soffice exited %d without writing a PDF: %s" % (
        r.returncode, (r.stderr or b"").decode("utf-8", "replace").strip()[-300:]))


def chrome_pdf(html_path, out_pdf, browser_furniture=False):
    """HTML -> PDF through headless Chromium.

    `browser_furniture=True` keeps Chromium's own printed header and footer --
    the date, document title, source URL and `page/total` that a person gets
    when they hit Print in a browser. Every existing `c*` fixture passes
    `--no-pdf-header-footer` and so has none of it, which means the corpus has
    never contained the most recognisable artifact of a browser-printed PDF.
    """
    if not CHROME:
        raise RuntimeError("no Chromium found (set CHROME=/path/to/chrome)")
    url = "file:///" + os.path.abspath(html_path).replace("\\", "/")
    cmd = [CHROME, "--headless=new", "--disable-gpu", "--no-sandbox",
           "--run-all-compositor-stages-before-draw",
           "--virtual-time-budget=4000"]
    if not browser_furniture:
        cmd.append("--no-pdf-header-footer")
    cmd += ["--print-to-pdf=" + os.path.abspath(out_pdf), url]
    r = subprocess.run(cmd, capture_output=True, timeout=180)
    if os.path.exists(out_pdf):
        return out_pdf
    raise RuntimeError("chromium exited %d: %s" % (
        r.returncode, (r.stderr or b"").decode("utf-8", "replace").strip()[-300:]))


BASE_CSS = """
@page { size: Letter; margin: 0.9in 0.8in; }
* { box-sizing: border-box; }
body { font-family: 'Liberation Serif', Georgia, serif; font-size: 11pt;
       line-height: 1.42; color: #111; margin: 0; }
h1 { font-family: 'Liberation Sans', Arial, sans-serif; font-size: 19pt; margin: 0 0 8pt; }
h2 { font-family: 'Liberation Sans', Arial, sans-serif; font-size: 13.5pt;
     margin: 16pt 0 5pt; }
h3 { font-family: 'Liberation Sans', Arial, sans-serif; font-size: 11.5pt;
     margin: 12pt 0 4pt; }
h4 { font-family: 'Liberation Sans', Arial, sans-serif; font-size: 10.5pt;
     margin: 10pt 0 3pt; font-style: italic; }
p { margin: 0 0 8pt; }
ul, ol { margin: 0 0 8pt 0; padding-left: 22pt; }
li { margin-bottom: 3pt; }
table { border-collapse: collapse; margin: 8pt 0 12pt; font-size: 10pt; width: 100%; }
td, th { padding: 4pt 7pt; text-align: left; vertical-align: top; }
blockquote { margin: 10pt 0 10pt 18pt; padding-left: 12pt;
             border-left: 2pt solid #888; font-style: italic; color: #333; }
a { color: #14417a; }
.notes { font-size: 8.5pt; color: #333; border-top: 0.5pt solid #999;
         padding-top: 5pt; margin-top: 14pt; }
.notes p { margin: 0 0 3pt; }
sup { font-size: 7pt; vertical-align: super; }
"""


def html_doc(name, body, extra_css=""):
    os.makedirs(HTML, exist_ok=True)
    p = os.path.join(HTML, name + ".html")
    with open(p, "w", encoding="utf-8") as f:
        f.write("<!doctype html><html><head><meta charset='utf-8'>"
                "<title>%s</title><style>%s\n%s</style></head><body>\n%s"
                "</body></html>" % (name.replace("_", " "), BASE_CSS, extra_css, body))
    return p


# ------------------------------------------------------- word-processor helpers
def _oxml(tag):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def _qn(tag):
    from docx.oxml.ns import qn
    return qn(tag)


def field(paragraph, instr, placeholder="1"):
    """A real Word field (PAGE, NUMPAGES). LibreOffice evaluates these on export.

    python-docx has no field API, so this is built from OOXML directly. The
    placeholder text is what a reader that does not evaluate fields would show;
    LibreOffice replaces it with the computed value when it writes the PDF.
    """
    fld = _oxml("w:fldSimple")
    fld.set(_qn("w:instr"), instr)
    run = _oxml("w:r")
    text = _oxml("w:t")
    text.text = placeholder
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)
    return paragraph


def hyperlink(paragraph, url, text):
    """An external hyperlink. Also absent from python-docx's API."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = _oxml("w:hyperlink")
    link.set(_qn("r:id"), rid)
    run = _oxml("w:r")
    props = _oxml("w:rPr")
    colour = _oxml("w:color")
    colour.set(_qn("w:val"), "14417A")
    under = _oxml("w:u")
    under.set(_qn("w:val"), "single")
    props.append(colour)
    props.append(under)
    run.append(props)
    node = _oxml("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)
    return paragraph


def table_borders(table, val="single", colour="808080", sz=4):
    """Set or clear every border on a table.

    Explicit rather than style-driven: the borderless case has to be borderless
    for a reason the file states, not because a template default happened to
    carry no borders on the machine that ran this.
    """
    props = table._tbl.tblPr
    for old in props.findall(_qn("w:tblBorders")):
        props.remove(old)
    borders = _oxml("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = _oxml("w:" + edge)
        el.set(_qn("w:val"), val)
        if val != "none":
            el.set(_qn("w:sz"), str(sz))
            el.set(_qn("w:color"), colour)
        borders.append(el)
    props.append(borders)
    return table


def shade(cell, colour="D9D9D9"):
    fill = _oxml("w:shd")
    fill.set(_qn("w:val"), "clear")
    fill.set(_qn("w:fill"), colour)
    cell._tc.get_or_add_tcPr().append(fill)
    return cell


def leader_tab(paragraph, position_pt=468):
    """A right tab stop with a dot leader -- how a word processor draws a TOC."""
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.shared import Pt
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Pt(position_pt), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    return paragraph


def new_docx(title=None, header_text=None, footer_page_field=False):
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = docx.Document()
    for s in d.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.9)
    style = d.styles["Normal"]
    style.font.name = "Liberation Serif"
    style.font.size = Pt(11)
    if header_text:
        hp = d.sections[0].header.paragraphs[0]
        hp.text = header_text
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.size = Pt(8.5)
    if footer_page_field:
        fp = d.sections[0].footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Page ")
        run.font.size = Pt(8.5)
        field(fp, " PAGE ")
        run = fp.add_run(" of ")
        run.font.size = Pt(8.5)
        field(fp, " NUMPAGES ")
    if title:
        d.add_heading(title, 0)
    return d


def build_docx(name, build):
    os.makedirs(WORK, exist_ok=True)
    src = os.path.join(WORK, name + ".docx")
    doc = build()
    doc.save(src)
    return soffice_pdf(src, os.path.join(OUT, name + ".pdf"))


# --------------------------------------------------- LibreOffice documents x01-x06
def x01_lo_memo_pageno():
    """Three-page memo: running header, footer PAGE of NUMPAGES, H1/H2/H3."""
    def build():
        d = new_docx(header_text="Transit Authority — Operations Committee — Restricted",
                     footer_page_field=True)
        d.add_heading("Depot Replacement: Interim Service Report", 1)
        d.add_paragraph("To: Operations Committee").runs[0].bold = True
        d.add_paragraph("From: Network Planning")
        d.add_paragraph("Date: 14 March")
        d.add_heading("1. Purpose", 2)
        d.add_paragraph(P1)
        d.add_heading("2. Background", 2)
        d.add_paragraph(P2)
        d.add_heading("2.1 Counting methodology", 3)
        d.add_paragraph(P3)
        d.add_heading("2.2 Comparison with the previous phase", 3)
        d.add_paragraph(P1 + P2)
        d.add_page_break()
        d.add_heading("3. Findings", 2)
        d.add_paragraph(P2 + P3)
        d.add_heading("3.1 Peak-hour circulation", 3)
        d.add_paragraph(P1)
        d.add_heading("3.2 Layover capacity", 3)
        d.add_paragraph(P4 + P1)
        d.add_page_break()
        d.add_heading("4. Recommendation", 2)
        d.add_paragraph(P4)
        d.add_paragraph(P2 + P3)
        return d
    return build_docx("x01_lo_memo_pageno", build)


def x02_lo_report_toc():
    """Contents page with dot leaders, then a numbered multi-page report."""
    def build():
        d = new_docx(header_text="Annual Review — Network Planning",
                     footer_page_field=True)
        d.add_heading("Annual Service Review", 0)
        d.add_paragraph("Network Planning Directorate")
        d.add_paragraph("Published under the transparency requirement")
        d.add_page_break()
        d.add_heading("Contents", 1)
        for label, page in [("1. Introduction", "1"), ("2. Method", "2"),
                            ("2.1 Data sources", "2"), ("2.2 Adjustments", "2"),
                            ("3. Results", "3"), ("3.1 Ridership", "3"),
                            ("3.2 Punctuality", "3"), ("4. Conclusions", "4"),
                            ("Appendix A. Counter calibration", "4")]:
            p = d.add_paragraph()
            leader_tab(p)
            p.add_run(label)
            p.add_run("\t" + page)
        d.add_page_break()
        for num, head, body in [
                ("1.", "Introduction", P1 + P2),
                ("2.", "Method", P3),
                ("2.1", "Data sources", P3 + P1),
                ("2.2", "Adjustments", P2),
                ("3.", "Results", P1 + P3),
                ("3.1", "Ridership", P2 + P4),
                ("3.2", "Punctuality", P3),
                ("4.", "Conclusions", P4 + P1)]:
            d.add_heading("%s %s" % (num, head), 1 if "." == num[-1] else 2)
            d.add_paragraph(body)
        d.add_heading("Appendix A. Counter calibration", 1)
        d.add_paragraph(P3 + P4)
        return d
    return build_docx("x02_lo_report_toc", build)


def x03_lo_lists_nested():
    """Bulleted and numbered lists nested three deep, and interleaved."""
    def build():
        d = new_docx(header_text="Standing Orders — Extract")
        d.add_heading("Depot Commissioning Checklist", 1)
        d.add_paragraph(P1)
        d.add_heading("Bulleted, three levels", 2)
        for text, style in [
                ("Site preparation", "List Bullet"),
                ("Confirm the hoarding line against the approved drawing", "List Bullet 2"),
                ("Photograph the boundary before any plant arrives", "List Bullet 2"),
                ("Both photographs are retained for the duration", "List Bullet 3"),
                ("Utilities", "List Bullet"),
                ("Isolate the disused feeder at the substation", "List Bullet 2"),
                ("Record the isolation certificate number", "List Bullet 3"),
                ("Countersign the record in the site diary", "List Bullet 3"),
                ("Prove dead before any cutting begins", "List Bullet 2"),
                ("Handover", "List Bullet")]:
            d.add_paragraph(text, style=style)
        d.add_heading("Numbered, three levels", 2)
        for text, style in [
                ("Establish the temporary layover", "List Number"),
                ("Mark the bay positions", "List Number 2"),
                ("Set the stop lines two metres back from the gate", "List Number 3"),
                ("Check the swept path with a full-length vehicle", "List Number 3"),
                ("Install the driver information board", "List Number 2"),
                ("Revise the running board", "List Number"),
                ("Shift the two relief runs twelve minutes earlier", "List Number 2"),
                ("Publish the revised layover times", "List Number"),
                ("Notify the trade union representatives", "List Number 2")]:
            d.add_paragraph(text, style=style)
        d.add_heading("Numbered with bulleted sub-points", 2)
        for text, style in [
                ("Confirm the contingency draw", "List Number"),
                ("Two relief runs, costed at the framework rate", "List Bullet 2"),
                ("No additional vehicles are required", "List Bullet 2"),
                ("Report to the committee in the autumn", "List Number"),
                ("A full quarter of post-occupancy data", "List Bullet 2")]:
            d.add_paragraph(text, style=style)
        d.add_paragraph(P4)
        return d
    return build_docx("x03_lo_lists_nested", build)


def x04_lo_tables_plain():
    """Bordered grid, fully borderless, and header-shaded-only tables."""
    def build():
        d = new_docx(header_text="Quarterly Figures")
        d.add_heading("Service Performance Tables", 1)
        d.add_paragraph(P3)

        d.add_heading("Table 1. Ridership by corridor (bordered grid)", 2)
        rows = [("Corridor", "Q1", "Q2", "Q3", "Change"),
                ("North radial", "412,880", "428,110", "441,207", "+6.9%"),
                ("Orbital west", "208,441", "203,900", "199,884", "-4.1%"),
                ("Cross-river", "331,209", "349,772", "362,015", "+9.3%"),
                ("Depot shuttle", "18,442", "17,880", "16,995", "-7.8%")]
        t = d.add_table(rows=len(rows), cols=5)
        table_borders(t, "single")
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
                if i == 0:
                    t.cell(i, j).paragraphs[0].runs[0].bold = True
        d.add_paragraph(P1)

        d.add_heading("Table 2. Contingency drawdown (borderless)", 2)
        rows = [("Item", "Committed", "Drawn", "Remaining"),
                ("Relief running", "142,000", "96,400", "45,600"),
                ("Signage and information", "38,500", "38,500", "0"),
                ("Temporary lighting", "61,200", "12,050", "49,150"),
                ("Contingency held", "250,000", "146,950", "103,050")]
        t = d.add_table(rows=len(rows), cols=4)
        table_borders(t, "none")
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
                if i == 0 or i == len(rows) - 1:
                    t.cell(i, j).paragraphs[0].runs[0].bold = True
        d.add_paragraph(P2)

        d.add_heading("Table 3. Punctuality (shaded header, no rules)", 2)
        rows = [("Period", "On time", "Within 5 min", "Missed"),
                ("January", "88.2%", "96.1%", "0.4%"),
                ("February", "86.9%", "95.4%", "0.6%"),
                ("March", "89.5%", "97.0%", "0.3%")]
        t = d.add_table(rows=len(rows), cols=4)
        table_borders(t, "none")
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
                if i == 0:
                    shade(t.cell(i, j))
                    t.cell(i, j).paragraphs[0].runs[0].bold = True
        d.add_paragraph(P4)
        return d
    return build_docx("x04_lo_tables_plain", build)


def x05_lo_quotes_notes():
    """Block quotes, an end-of-page note apparatus, and inline hyperlinks."""
    def build():
        from docx.shared import Pt
        d = new_docx(header_text="Consultation Response")
        d.add_heading("Response to the Draft Layover Policy", 1)
        d.add_paragraph(P1)
        d.add_paragraph("The draft states:")
        q = d.add_paragraph(
            "Operators shall not be required to accept a layover position that "
            "cannot accommodate a full-length vehicle without reversing.",
            style="Quote")
        q.paragraph_format.left_indent = Pt(24)
        p = d.add_paragraph(P2.rstrip() + " ")
        p.add_run("1").font.superscript = True
        d.add_heading("Second observation", 2)
        d.add_paragraph(P3)
        q = d.add_paragraph(
            "Where a reversing manoeuvre is unavoidable, a banksman shall be "
            "provided for every movement, without exception.", style="Quote")
        q.paragraph_format.left_indent = Pt(24)
        p = d.add_paragraph(P4.rstrip() + " ")
        p.add_run("2").font.superscript = True
        p = d.add_paragraph("The full consultation text is published at ")
        hyperlink(p, "https://example.org/consultations/layover-policy",
                  "example.org/consultations/layover-policy")
        p.add_run(" and responses may be sent to ")
        hyperlink(p, "mailto:consultation@example.org", "consultation@example.org")
        p.add_run(".")
        note = d.add_paragraph()
        note.add_run("_" * 34).font.size = Pt(8)
        for marker, text in [
                ("1", "Costed at the framework rate current at the date of this "
                      "response; no allowance is made for indexation."),
                ("2", "The operator's own standard already requires this on all "
                      "sites, so no additional cost is anticipated.")]:
            n = d.add_paragraph()
            run = n.add_run(marker)
            run.font.superscript = True
            run.font.size = Pt(8)
            run = n.add_run(" " + text)
            run.font.size = Pt(8.5)
        return d
    return build_docx("x05_lo_quotes_notes", build)


def x06_lo_euro_scripts():
    """Latin diacritics, Cyrillic and Greek -- all inside the pinned font set.

    The deliberate contrast with `c4_i18n`, which is `designed_stress` because
    CJK and RTL move under font fallback. These scripts do not: Liberation and
    DejaVu cover them outright, so this is an ordinary document that happens not
    to be in English.
    """
    def build():
        d = new_docx(header_text="Multilingual Passenger Information")
        d.add_heading("Passenger Notice — European Language Set", 1)
        d.add_heading("Latin with diacritics", 2)
        d.add_paragraph(
            "Čeština: Náhradní autobusová doprava je zajištěna po celou dobu "
            "výluky. Prosíme cestující, aby dbali pokynů personálu.")
        d.add_paragraph(
            "Polski: Zastępcza komunikacja autobusowa kursuje przez cały okres "
            "zamknięcia. Prosimy o stosowanie się do poleceń obsługi.")
        d.add_paragraph(
            "Tiếng Việt: Xe buýt thay thế hoạt động trong suốt thời gian đóng "
            "cửa. Hành khách vui lòng làm theo hướng dẫn của nhân viên.")
        d.add_paragraph(
            "Norsk / Dansk: Erstatningsbusser kjører i hele perioden. "
            "Ærøskøbing, Straße, naïve façade, Voilà.")
        d.add_heading("Cyrillic", 2)
        d.add_paragraph(
            "Русский: Замещающие автобусы курсируют в течение всего периода "
            "закрытия. Просим пассажиров следовать указаниям персонала.")
        d.add_paragraph(
            "Українська: Замінні автобуси курсують упродовж усього періоду "
            "закриття. Просимо пасажирів дотримуватися вказівок персоналу.")
        d.add_paragraph(
            "Српски: Замењујући аутобуси саобраћају током целог периода "
            "затварања депоа.")
        d.add_heading("Greek", 2)
        d.add_paragraph(
            "Ελληνικά: Τα λεωφορεία αντικατάστασης λειτουργούν καθ' όλη τη "
            "διάρκεια της διακοπής. Παρακαλούνται οι επιβάτες να ακολουθούν "
            "τις οδηγίες του προσωπικού.")
        d.add_heading("Mixed numerals and symbols", 2)
        d.add_paragraph("№ 14 · 12–18 min · 3 °C · ± 2 % · € 2,40 · £ 1.90 · "
                        "α β γ Δ Σ Ω · А Б В Г Д")
        d.add_paragraph(P4)
        return d
    return build_docx("x06_lo_euro_scripts", build)


# ------------------------------------------------------ Chromium documents x07-x12
def x07_chrome_memo_running():
    """Fixed running header and footer repeated on every printed page."""
    css = """
.runhead { position: fixed; top: -0.55in; left: 0; right: 0;
           font-family: 'Liberation Sans', Arial, sans-serif; font-size: 8pt;
           color: #444; border-bottom: 0.5pt solid #bbb; padding-bottom: 3pt; }
.runfoot { position: fixed; bottom: -0.55in; left: 0; right: 0;
           font-family: 'Liberation Sans', Arial, sans-serif; font-size: 8pt;
           color: #444; border-top: 0.5pt solid #bbb; padding-top: 3pt;
           display: flex; justify-content: space-between; }
"""
    secs = []
    for i, head in enumerate(
            ["Purpose", "Background", "Site constraints", "Circulation",
             "Layover capacity", "Cost", "Consultation", "Recommendation"], 1):
        secs.append("<h2>%d. %s</h2><p>%s</p><p>%s</p>"
                    % (i, head, P1 if i % 2 else P3, P2 if i % 2 else P4))
    body = ("<div class='runhead'>Transit Authority &middot; Operations Committee "
            "&middot; Interim Service Report</div>"
            "<div class='runfoot'><span>Network Planning</span>"
            "<span>Not for onward circulation</span></div>"
            "<h1>Depot Replacement: Interim Service Report</h1>"
            "<p>%s</p>%s" % (P1 + P2, "".join(secs)))
    return chrome_pdf(html_doc("x07_chrome_memo_running", body, css),
                      os.path.join(OUT, "x07_chrome_memo_running.pdf"))


def x08_chrome_print_default():
    """Printed with Chromium's own header/footer: date, title, URL, page N/M.

    No existing fixture has this. Every `c*` document passes
    `--no-pdf-header-footer`, so the corpus has never contained the margin
    furniture that a browser puts on a printed page -- which is the single most
    recognisable signature of a print-to-PDF document in the wild.
    """
    secs = []
    for i, head in enumerate(
            ["What is changing", "When", "Which services are affected",
             "Replacement buses", "Accessibility", "Refunds", "Contact"], 1):
        secs.append("<h2>%d. %s</h2><p>%s</p><p>%s</p>"
                    % (i, head, P3 if i % 2 else P1, P4 if i % 2 else P2))
    body = ("<h1>Planned Depot Closure — Passenger Information</h1>"
            "<p>%s</p>%s<p>%s</p>" % (P1, "".join(secs), P4))
    return chrome_pdf(html_doc("x08_chrome_print_default", body),
                      os.path.join(OUT, "x08_chrome_print_default.pdf"),
                      browser_furniture=True)


def x09_chrome_lists_nested():
    """Bulleted and numbered lists nested three deep, and interleaved."""
    body = """
<h1>Depot Commissioning Checklist</h1>
<p>%s</p>
<h2>Bulleted, three levels</h2>
<ul>
  <li>Site preparation
    <ul><li>Confirm the hoarding line against the approved drawing</li>
        <li>Photograph the boundary before any plant arrives
          <ul><li>Both photographs are retained for the duration</li>
              <li>Filed against the commissioning record</li></ul></li></ul></li>
  <li>Utilities
    <ul><li>Isolate the disused feeder at the substation
          <ul><li>Record the isolation certificate number</li>
              <li>Countersign the record in the site diary</li></ul></li>
        <li>Prove dead before any cutting begins</li></ul></li>
  <li>Handover</li>
</ul>
<h2>Numbered, three levels</h2>
<ol>
  <li>Establish the temporary layover
    <ol><li>Mark the bay positions
          <ol><li>Set the stop lines two metres back from the gate</li>
              <li>Check the swept path with a full-length vehicle</li></ol></li>
        <li>Install the driver information board</li></ol></li>
  <li>Revise the running board
    <ol><li>Shift the two relief runs twelve minutes earlier</li>
        <li>Reprint the duty cards</li></ol></li>
  <li>Publish the revised layover times</li>
</ol>
<h2>Numbered with bulleted sub-points</h2>
<ol>
  <li>Confirm the contingency draw
    <ul><li>Two relief runs, costed at the framework rate</li>
        <li>No additional vehicles are required</li></ul></li>
  <li>Report to the committee in the autumn
    <ul><li>A full quarter of post-occupancy data</li></ul></li>
</ol>
<p>%s</p>
""" % (P1, P4)
    return chrome_pdf(html_doc("x09_chrome_lists_nested", body),
                      os.path.join(OUT, "x09_chrome_lists_nested.pdf"))


def x10_chrome_tables_plain():
    """Bordered, borderless and shaded-header tables. No nesting anywhere."""
    css = """
table.grid td, table.grid th { border: 0.5pt solid #808080; }
table.bare td, table.bare th { border: none; }
table.head th { background: #d9d9d9; border: none; }
table.head td { border: none; }
td.n, th.n { text-align: right; }
"""
    body = """
<h1>Service Performance Tables</h1>
<p>%s</p>
<h2>Table 1. Ridership by corridor (bordered grid)</h2>
<table class="grid">
<tr><th>Corridor</th><th class="n">Q1</th><th class="n">Q2</th>
    <th class="n">Q3</th><th class="n">Change</th></tr>
<tr><td>North radial</td><td class="n">412,880</td><td class="n">428,110</td>
    <td class="n">441,207</td><td class="n">+6.9%%</td></tr>
<tr><td>Orbital west</td><td class="n">208,441</td><td class="n">203,900</td>
    <td class="n">199,884</td><td class="n">-4.1%%</td></tr>
<tr><td>Cross-river</td><td class="n">331,209</td><td class="n">349,772</td>
    <td class="n">362,015</td><td class="n">+9.3%%</td></tr>
<tr><td>Depot shuttle</td><td class="n">18,442</td><td class="n">17,880</td>
    <td class="n">16,995</td><td class="n">-7.8%%</td></tr>
</table>
<p>%s</p>
<h2>Table 2. Contingency drawdown (borderless)</h2>
<table class="bare">
<tr><th>Item</th><th class="n">Committed</th><th class="n">Drawn</th>
    <th class="n">Remaining</th></tr>
<tr><td>Relief running</td><td class="n">142,000</td><td class="n">96,400</td>
    <td class="n">45,600</td></tr>
<tr><td>Signage and information</td><td class="n">38,500</td><td class="n">38,500</td>
    <td class="n">0</td></tr>
<tr><td>Temporary lighting</td><td class="n">61,200</td><td class="n">12,050</td>
    <td class="n">49,150</td></tr>
<tr><td><b>Contingency held</b></td><td class="n"><b>250,000</b></td>
    <td class="n"><b>146,950</b></td><td class="n"><b>103,050</b></td></tr>
</table>
<p>%s</p>
<h2>Table 3. Punctuality (shaded header, no rules)</h2>
<table class="head">
<tr><th>Period</th><th class="n">On time</th><th class="n">Within 5 min</th>
    <th class="n">Missed</th></tr>
<tr><td>January</td><td class="n">88.2%%</td><td class="n">96.1%%</td>
    <td class="n">0.4%%</td></tr>
<tr><td>February</td><td class="n">86.9%%</td><td class="n">95.4%%</td>
    <td class="n">0.6%%</td></tr>
<tr><td>March</td><td class="n">89.5%%</td><td class="n">97.0%%</td>
    <td class="n">0.3%%</td></tr>
</table>
<p>%s</p>
""" % (P3, P1, P2, P4)
    return chrome_pdf(html_doc("x10_chrome_tables_plain", body, css),
                      os.path.join(OUT, "x10_chrome_tables_plain.pdf"))


def x11_chrome_toc_headings():
    """Contents with dot leaders, four heading levels, internal and external links."""
    css = """
.toc { list-style: none; padding-left: 0; }
.toc li { display: flex; align-items: baseline; margin-bottom: 4pt; }
.toc .dots { flex: 1; border-bottom: 0.8pt dotted #999; margin: 0 4pt;
             transform: translateY(-2pt); }
.toc a { text-decoration: none; }
"""
    entries = [("s1", "1. Introduction", "1"), ("s2", "2. Method", "1"),
               ("s21", "2.1 Data sources", "2"), ("s22", "2.2 Adjustments", "2"),
               ("s3", "3. Results", "2"), ("s31", "3.1 Ridership", "3"),
               ("s4", "4. Conclusions", "3")]
    toc = "".join("<li><a href='#%s'>%s</a><span class='dots'></span>"
                  "<span>%s</span></li>" % e for e in entries)
    body = """
<h1>Annual Service Review</h1>
<p>%s</p>
<h2>Contents</h2>
<ul class="toc">%s</ul>
<h2 id="s1">1. Introduction</h2><p>%s</p>
<p>The statutory basis is set out in <a href="https://example.org/standards/ptr-2019">
the passenger transport regulations</a>; the counter specification is published
<a href="https://example.org/specs/apc">separately</a>.</p>
<h2 id="s2">2. Method</h2><p>%s</p>
<h3 id="s21">2.1 Data sources</h3><p>%s</p>
<h4>2.1.1 Automatic passenger counters</h4><p>%s</p>
<h4>2.1.2 Manual verification</h4><p>%s</p>
<h3 id="s22">2.2 Adjustments</h3><p>%s</p>
<blockquote>No adjustment has been applied; the raw counts are reported so that
later revisions remain traceable to their source.</blockquote>
<h2 id="s3">3. Results</h2><p>%s</p>
<h3 id="s31">3.1 Ridership</h3><p>%s</p>
<h2 id="s4">4. Conclusions</h2><p>%s</p>
<p>Return to <a href="#s1">the introduction</a>, or write to
<a href="mailto:planning@example.org">planning@example.org</a>.</p>
<div class="notes">
<p><sup>1</sup> Counter under-reporting is stated by the vendor as three per cent
at the rear doors.</p>
<p><sup>2</sup> Punctuality is measured at the timing points only.</p>
</div>
""" % (P1, toc, P2, P3, P1, P2, P3, P4, P1, P2, P4)
    return chrome_pdf(html_doc("x11_chrome_toc_headings", body, css),
                      os.path.join(OUT, "x11_chrome_toc_headings.pdf"))


def x12_chrome_euro_scripts():
    """The browser producer's take on Latin-diacritic, Cyrillic and Greek text."""
    body = """
<h1>Passenger Notice — European Language Set</h1>
<h2>Latin with diacritics</h2>
<p><b>Čeština:</b> Náhradní autobusová doprava je zajištěna po celou dobu
výluky. Prosíme cestující, aby dbali pokynů personálu.</p>
<p><b>Polski:</b> Zastępcza komunikacja autobusowa kursuje przez cały okres
zamknięcia. Prosimy o stosowanie się do poleceń obsługi.</p>
<p><b>Tiếng Việt:</b> Xe buýt thay thế hoạt động trong suốt thời gian đóng cửa.
Hành khách vui lòng làm theo hướng dẫn của nhân viên.</p>
<p><b>Norsk / Dansk:</b> Erstatningsbusser kjører i hele perioden.
Ærøskøbing, Straße, naïve façade, Voilà.</p>
<h2>Cyrillic</h2>
<p><b>Русский:</b> Замещающие автобусы курсируют в течение всего периода
закрытия. Просим пассажиров следовать указаниям персонала.</p>
<p><b>Українська:</b> Замінні автобуси курсують упродовж усього періоду
закриття. Просимо пасажирів дотримуватися вказівок персоналу.</p>
<p><b>Српски:</b> Замењујући аутобуси саобраћају током целог периода
затварања депоа.</p>
<h2>Greek</h2>
<p><b>Ελληνικά:</b> Τα λεωφορεία αντικατάστασης λειτουργούν καθ' όλη τη
διάρκεια της διακοπής. Παρακαλούνται οι επιβάτες να ακολουθούν τις οδηγίες
του προσωπικού.</p>
<h2>Mixed numerals and symbols</h2>
<p>№ 14 &middot; 12–18 min &middot; 3 °C &middot; ± 2 %% &middot; € 2,40 &middot;
£ 1.90 &middot; α β γ Δ Σ Ω &middot; А Б В Г Д</p>
<p>%s</p>
""" % P4
    return chrome_pdf(html_doc("x12_chrome_euro_scripts", body),
                      os.path.join(OUT, "x12_chrome_euro_scripts.pdf"))


# ------------------------------------------- report-generator documents x13-x16
class _Numbered(object):
    """ReportLab two-pass canvas so a footer can say 'Page N of M'.

    A single-pass canvas cannot know M while it is drawing page 1. Pages are
    buffered, the total is counted at save time, and each page is then replayed
    with the furniture drawn on it.
    """

    def __call__(self, *args, **kwargs):
        raise NotImplementedError


def _numbered_canvas(header_left, header_right):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import LETTER

    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            canvas.Canvas.__init__(self, *args, **kwargs)
            self._saved = []

        def showPage(self):
            self._saved.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._saved)
            for state in self._saved:
                self.__dict__.update(state)
                self._furniture(total)
                canvas.Canvas.showPage(self)
            canvas.Canvas.save(self)

        def _furniture(self, total):
            width, height = LETTER
            self.setFont("Helvetica", 8)
            self.setFillGray(0.35)
            self.drawString(64, height - 44, header_left)
            self.drawRightString(width - 64, height - 44, header_right)
            self.setStrokeGray(0.75)
            self.line(64, height - 50, width - 64, height - 50)
            self.line(64, 52, width - 64, 52)
            self.drawString(64, 40, "Network Planning Directorate")
            self.drawRightString(width - 64, 40,
                                 "Page %d of %d" % (self._pageNumber, total))
    return NumberedCanvas


def _rl_styles():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    ss = getSampleStyleSheet()
    return {
        "body": ParagraphStyle("body", parent=ss["BodyText"], fontName="Times-Roman",
                               fontSize=10.5, leading=14.5, alignment=4, spaceAfter=8),
        "h1": ParagraphStyle("h1", parent=ss["Heading1"], fontName="Helvetica-Bold",
                             fontSize=17, spaceAfter=10),
        "h2": ParagraphStyle("h2", parent=ss["Heading2"], fontName="Helvetica-Bold",
                             fontSize=12.5, textColor=HexColor("#1a1a1a"), spaceAfter=6),
        "h3": ParagraphStyle("h3", parent=ss["Heading3"], fontName="Helvetica-Oblique",
                             fontSize=10.5, spaceAfter=4),
        "quote": ParagraphStyle("quote", parent=ss["BodyText"], fontName="Times-Italic",
                                fontSize=10.5, leading=14.5, leftIndent=26,
                                rightIndent=14, spaceBefore=6, spaceAfter=8),
        "note": ParagraphStyle("note", parent=ss["BodyText"], fontName="Times-Roman",
                               fontSize=8, leading=10.5, spaceAfter=2),
        "toc": ParagraphStyle("toc", parent=ss["BodyText"], fontName="Times-Roman",
                              fontSize=10.5, leading=15, spaceAfter=2),
    }


def _grid_style(bordered=True, shade_header=True):
    from reportlab.platypus import TableStyle
    from reportlab.lib.colors import HexColor
    cmds = [("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6)]
    if bordered:
        cmds.append(("GRID", (0, 0), (-1, -1), 0.4, HexColor("#808080")))
    if shade_header:
        cmds.append(("BACKGROUND", (0, 0), (-1, 0), HexColor("#d9d9d9")))
    return TableStyle(cmds)


def x13_rl_report_running():
    """Multi-page ReportLab report: running furniture, 'Page N of M', two tables."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
    out = os.path.join(OUT, "x13_rl_report_running.pdf")
    st = _rl_styles()
    doc = SimpleDocTemplate(out, pagesize=LETTER, leftMargin=0.9 * inch,
                            rightMargin=0.9 * inch, topMargin=0.95 * inch,
                            bottomMargin=0.9 * inch,
                            title="Interim Service Report")
    story = [Paragraph("Depot Replacement: Interim Service Report", st["h1"]),
             Paragraph(P1 + P2, st["body"])]
    for i, head in enumerate(["Purpose", "Background", "Site constraints",
                              "Circulation", "Layover capacity"], 1):
        story += [Paragraph("%d. %s" % (i, head), st["h2"]),
                  Paragraph(P3 if i % 2 else P1, st["body"]),
                  Paragraph(P2 if i % 2 else P4, st["body"])]
    story += [Paragraph("6. Ridership (bordered)", st["h2"])]
    rows = [["Corridor", "Q1", "Q2", "Q3"],
            ["North radial", "412,880", "428,110", "441,207"],
            ["Orbital west", "208,441", "203,900", "199,884"],
            ["Cross-river", "331,209", "349,772", "362,015"]]
    t = Table(rows, colWidths=[150, 85, 85, 85])
    t.setStyle(_grid_style(bordered=True))
    story += [t, Spacer(1, 10),
              Paragraph("7. Contingency drawdown (borderless)", st["h2"])]
    rows = [["Item", "Committed", "Drawn", "Remaining"],
            ["Relief running", "142,000", "96,400", "45,600"],
            ["Signage and information", "38,500", "38,500", "0"],
            ["Temporary lighting", "61,200", "12,050", "49,150"]]
    t = Table(rows, colWidths=[170, 80, 80, 80])
    t.setStyle(_grid_style(bordered=False, shade_header=False))
    story += [t, Spacer(1, 10),
              Paragraph("8. Recommendation", st["h2"]),
              Paragraph(P4 + P1, st["body"]),
              Paragraph(P2 + P3, st["body"])]
    doc.build(story, canvasmaker=_numbered_canvas(
        "Transit Authority — Operations Committee", "Interim Service Report"))
    return out


def x14_rl_statement_lines():
    """Statement/invoice shape: borderless line items, right-aligned money, totals.

    The commonest output of a report generator in the wild, and absent from the
    frozen 16 entirely.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)
    out = os.path.join(OUT, "x14_rl_statement_lines.pdf")
    st = _rl_styles()
    doc = SimpleDocTemplate(out, pagesize=LETTER, leftMargin=0.9 * inch,
                            rightMargin=0.9 * inch, topMargin=0.9 * inch,
                            bottomMargin=0.9 * inch, title="Statement of Account")
    story = [Paragraph("Statement of Account", st["h1"]),
             Paragraph("Account 44-20817 &middot; Period ending 31 March", st["body"]),
             Spacer(1, 6)]
    rows = [["Date", "Reference", "Description", "Qty", "Rate", "Amount"]]
    items = [("04 Jan", "RR-1181", "Relief running, north radial", "18", "412.00", "7,416.00"),
             ("11 Jan", "RR-1194", "Relief running, orbital west", "12", "412.00", "4,944.00"),
             ("22 Jan", "SG-0042", "Temporary signage, depot gate", "1", "3,880.00", "3,880.00"),
             ("03 Feb", "RR-1210", "Relief running, north radial", "18", "412.00", "7,416.00"),
             ("14 Feb", "TL-0071", "Temporary lighting hire", "6", "1,004.00", "6,024.00"),
             ("27 Feb", "RR-1233", "Relief running, cross-river", "9", "438.00", "3,942.00"),
             ("09 Mar", "SG-0051", "Driver information board", "2", "1,240.00", "2,480.00"),
             ("21 Mar", "RR-1259", "Relief running, north radial", "18", "412.00", "7,416.00")]
    rows += [list(i) for i in items]
    rows += [["", "", "", "", "Subtotal", "43,518.00"],
             ["", "", "", "", "Contingency applied", "-6,568.00"],
             ["", "", "", "", "Total due", "36,950.00"]]
    t = Table(rows, colWidths=[52, 62, 178, 34, 74, 78])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTNAME", (4, len(rows) - 1), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (3, 0), (-1, -1), "RIGHT"),
        ("LINEBELOW", (0, 0), (-1, 0), 0.6, HexColor("#333333")),
        ("LINEABOVE", (4, len(rows) - 3), (-1, len(rows) - 3), 0.4, HexColor("#333333")),
        ("LINEABOVE", (4, len(rows) - 1), (-1, len(rows) - 1), 0.9, HexColor("#333333")),
        ("TOPPADDING", (0, 0), (-1, -1), 3.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4)]))
    story += [t, Spacer(1, 14),
              Paragraph("Notes", st["h2"]),
              Paragraph(P3, st["body"]),
              Paragraph("Payment terms are thirty days from the date of this "
                        "statement. Queries should quote the account number "
                        "above.", st["body"])]
    doc.build(story)
    return out


def x15_rl_handbook_toc():
    """Contents page with dot leaders, heading hierarchy, block quote, notes."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    out = os.path.join(OUT, "x15_rl_handbook_toc.pdf")
    st = _rl_styles()
    doc = SimpleDocTemplate(out, pagesize=LETTER, leftMargin=0.95 * inch,
                            rightMargin=0.95 * inch, topMargin=0.9 * inch,
                            bottomMargin=0.9 * inch, title="Layover Handbook")
    dots = '<font color="#999999">' + "." * 60 + "</font>"
    story = [Paragraph("Layover Handbook", st["h1"]),
             Paragraph("Network Planning Directorate", st["body"]),
             Spacer(1, 12), Paragraph("Contents", st["h2"])]
    for label, page in [("1. Scope", "1"), ("2. Definitions", "1"),
                        ("2.1 Layover", "2"), ("2.2 Relief run", "2"),
                        ("3. Standards", "2"), ("3.1 Swept path", "3"),
                        ("3.2 Reversing", "3"), ("4. Records", "3")]:
        story.append(Paragraph(
            '<para>%s %s %s</para>' % (label, dots, page), st["toc"]))
    story.append(PageBreak())
    for num, head, level, text in [
            ("1.", "Scope", "h2", P1),
            ("2.", "Definitions", "h2", P2),
            ("2.1", "Layover", "h3", P3),
            ("2.2", "Relief run", "h3", P1),
            ("3.", "Standards", "h2", P2 + P3),
            ("3.1", "Swept path", "h3", P1),
            ("3.2", "Reversing", "h3", P4)]:
        story.append(Paragraph("%s %s" % (num, head), st[level]))
        story.append(Paragraph(text, st["body"]))
        if num == "3.2":
            story.append(Paragraph(
                "Where a reversing manoeuvre is unavoidable, a banksman shall "
                "be provided for every movement, without exception.<super>1</super>",
                st["quote"]))
    story += [Paragraph("4. Records", st["h2"]), Paragraph(P4 + P2, st["body"]),
              Spacer(1, 16),
              Paragraph("<font color='#999999'>%s</font>" % ("_" * 40), st["note"]),
              Paragraph("<super>1</super> The operator's own standard already "
                        "requires this on all sites, so no additional cost is "
                        "anticipated.", st["note"]),
              Paragraph("<super>2</super> Counter under-reporting is stated by "
                        "the vendor as three per cent at the rear doors.",
                        st["note"])]
    doc.build(story)
    return out


def x16_fpdf_bulletin():
    """fpdf2 with core-14 fonts: no FontDescriptor, footer page numbers.

    Text here is Latin-1 only, on purpose. The core-14 fonts carry no Unicode
    cmap, and fpdf2 raises rather than silently dropping a character it cannot
    encode -- an em dash is enough to stop it. That constraint is exactly what
    makes this document worth having: it is the dialect a PDF acquires when its
    producer used the base fonts and nothing else.
    """
    from fpdf import FPDF
    out = os.path.join(OUT, "x16_fpdf_bulletin.pdf")

    class Bulletin(FPDF):
        def header(self):
            self.set_font("Helvetica", "", 8)
            self.set_text_color(90)
            self.cell(0, 12, "Operations Bulletin - Network Planning", ln=1)
            self.set_draw_color(180)
            self.line(self.l_margin, 40, self.w - self.r_margin, 40)
            self.ln(6)
            self.set_text_color(0)

        def footer(self):
            self.set_y(-40)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(90)
            self.cell(0, 10, "Page %d of {nb}" % self.page_no(), align="C")
            self.set_text_color(0)

    p = Bulletin(format="letter", unit="pt")
    p.set_auto_page_break(True, margin=56)
    p.add_page()
    p.set_font("Helvetica", "B", 17)
    p.cell(0, 24, "Depot Replacement Bulletin", ln=1)
    p.ln(4)
    p.set_font("Times", "", 10.5)
    p.multi_cell(0, 14, P1 + P2, align="J")
    p.ln(6)
    for head, text in [("Current position", P3), ("Mitigation", P2),
                       ("Next report", P4)]:
        p.set_font("Helvetica", "B", 12)
        p.cell(0, 18, head, ln=1)
        p.set_font("Times", "", 10.5)
        p.multi_cell(0, 14, text, align="J")
        p.ln(4)
    p.set_font("Helvetica", "B", 12)
    p.cell(0, 18, "Punctuality (bordered)", ln=1)
    p.set_font("Helvetica", "", 9)
    for row in [("Period", "On time", "Within 5 min", "Missed"),
                ("January", "88.2%", "96.1%", "0.4%"),
                ("February", "86.9%", "95.4%", "0.6%"),
                ("March", "89.5%", "97.0%", "0.3%")]:
        for cell in row:
            p.cell(105, 16, cell, border=1)
        p.ln()
    p.ln(10)
    p.set_font("Times", "", 10.5)
    p.multi_cell(0, 14, P1 + P3 + P4, align="J")
    p.output(out)
    return out


# --------------------------------------------------------------- the registry
# name -> (function, tier, dialect, why). Every entry is `ordinary_digital` by
# the rules in docs/corpus-expansion.md §4; see that file for why tranche 1 adds
# no stress documents.
DOCS = [
    (x01_lo_memo_pageno, "libreoffice",
     "running header and footer with real PAGE/NUMPAGES fields over three pages"),
    (x02_lo_report_toc, "libreoffice",
     "contents page with dot-leader tab stops ahead of a numbered multi-page report"),
    (x03_lo_lists_nested, "libreoffice",
     "bulleted and numbered lists nested three deep, and interleaved"),
    (x04_lo_tables_plain, "libreoffice",
     "bordered grid, fully borderless, and shaded-header-only tables"),
    (x05_lo_quotes_notes, "libreoffice",
     "block quotes, a page-bottom note apparatus, and inline hyperlinks"),
    (x06_lo_euro_scripts, "libreoffice",
     "Latin diacritics, Cyrillic and Greek inside the pinned font set"),
    (x07_chrome_memo_running, "chromium",
     "fixed running header and footer repeated on every printed page"),
    (x08_chrome_print_default, "chromium",
     "Chromium's own printed furniture: date, title, URL and page N/M"),
    (x09_chrome_lists_nested, "chromium",
     "nested ul/ol three levels deep, and interleaved"),
    (x10_chrome_tables_plain, "chromium",
     "bordered, borderless and shaded-header tables with no nesting"),
    (x11_chrome_toc_headings, "chromium",
     "contents with dot leaders, four heading levels, internal and external links"),
    (x12_chrome_euro_scripts, "chromium",
     "the browser producer's Latin-diacritic, Cyrillic and Greek text"),
    (x13_rl_report_running, "reportlab",
     "report-generator running furniture with a two-pass 'Page N of M' footer"),
    (x14_rl_statement_lines, "reportlab",
     "statement line items: borderless rules, right-aligned money, totals"),
    (x15_rl_handbook_toc, "reportlab",
     "contents with dot leaders, heading hierarchy, block quote and notes"),
    (x16_fpdf_bulletin, "fpdf2",
     "core-14 fonts with no FontDescriptor, and a page-numbered footer"),
]


def toolchain():
    """What produced these bytes. Recorded per document, in its provenance."""
    def first_line(cmd):
        try:
            r = subprocess.run(cmd, capture_output=True, timeout=60)
            return (r.stdout or b"").decode("utf-8", "replace").strip().splitlines()[0]
        except Exception:
            return "unavailable"
    parts = ["python %d.%d.%d" % sys.version_info[:3]]
    if SOFFICE:
        parts.append(first_line([SOFFICE, "--version"]))
    if CHROME:
        parts.append(first_line([CHROME, "--version"]))
    for mod, label in (("reportlab", "reportlab"), ("fpdf", "fpdf2"),
                       ("docx", "python-docx")):
        try:
            import importlib.metadata as md
            parts.append("%s %s" % (label, md.version(label)))
        except Exception:
            parts.append("%s unknown" % label)
    image = "/etc/exactdoc-image.txt"
    if os.path.exists(image):
        parts.append("image=exactdoc-gate")
    return " / ".join(parts)


def main():
    os.makedirs(OUT, exist_ok=True)
    os.makedirs(WORK, exist_ok=True)
    print("capabilities: chromium=%s  soffice=%s"
          % (CHROME or "MISSING", SOFFICE or "MISSING"))
    if not CHROME or not SOFFICE:
        print("\nBoth producers are required for this tranche. A partial "
              "generation would freeze a corpus whose producer mix is not the "
              "mix on record, which is the whole point of the exercise.")
        return 1

    chain = toolchain()
    print("toolchain: %s\n" % chain)
    made, failed, provenance = [], [], {}
    for fn, dialect, why in DOCS:
        name = fn.__name__
        try:
            path = fn()
            if not os.path.exists(path) or os.path.getsize(path) < 500:
                raise RuntimeError("wrote no usable PDF")
            made.append(path)
            provenance[os.path.basename(path)] = {
                "tier": "ordinary_digital",
                "dialect": dialect,
                "why": why,
                "provenance": {
                    "origin": "generated",
                    "recipe": "testkit/gen_expansion.py::%s" % name,
                    "source_url": None,
                    "license": LICENSE,
                    "acquired": ACQUIRED,
                    "toolchain": chain,
                },
            }
            print("  OK   %s" % os.path.basename(path))
        except Exception as exc:
            failed.append((name, "%s: %s" % (type(exc).__name__, exc)))
            print("  FAIL %s" % name)

    if failed:
        print("\nFAILED %d document(s):" % len(failed))
        for name, why in failed:
            print("  %-26s %s" % (name, why))
        print("\nNothing is sealed. A tranche is frozen whole or not at all: a "
              "partial freeze mixes two generation runs under one provenance "
              "record.")
        return 1

    side = os.path.join(OUT, "expansion_provenance.json")
    with open(side, "w", encoding="utf-8") as fh:
        json.dump({"schema": "exactdoc.expansion-provenance.v1",
                   "documents": provenance}, fh, indent=1, sort_keys=True)
        fh.write("\n")

    print("\n%d PDFs in %s" % (len(made), os.path.abspath(OUT)))
    try:
        import fitz
        for path in sorted(made):
            doc = fitz.open(path)
            print("  %-30s %d pages  %7d bytes  producer=%s"
                  % (os.path.basename(path), doc.page_count,
                     os.path.getsize(path), doc.metadata.get("producer")))
            doc.close()
    except ImportError:
        pass
    print("\nprovenance sidecar: %s" % side)
    print("next:  python testkit/corpus_manifest.py expansion-seal %s" % OUT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
