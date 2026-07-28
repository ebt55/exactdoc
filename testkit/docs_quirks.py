"""Characterise Google Docs' layout quirks with minimal single-variable DOCX.

The closed loop can *correct* the Docs offset per document. That is a
compensator, not an explanation, and it costs a network round trip per pass.
The offset is regular (one-off after the first heading, ~3pt per paragraph
boundary, accumulating), and regularity implies a generative rule. Find the
rule and it becomes a static, offline, deterministic fix for every document.

Each probe changes exactly ONE variable, is rendered by both LibreOffice and
Google Docs, and reports the gap each renderer produced.

    python testkit/docs_quirks.py            # all probes
    python testkit/docs_quirks.py h1         # one family
"""
import os
import re
import sys

import _paths  # noqa: F401
import fitz
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import harness

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "quirks")
MARK_A = "AAAAMARKERTOP"
MARK_B = "ZZZZMARKERBOTTOM"


def _neutralise(doc):
    pf = doc.styles["Normal"].paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _text(doc, txt, size=10.0, leading=12.0, before=0.0, after=0.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if leading:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(leading)
    r = p.add_run(txt)
    r.font.size = Pt(size)
    r.font.name = "Arial"
    return p


def _spacer(doc, height):
    """The construct exactdoc emits before tables and for section breaks."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(max(1.0, height))
    r = p.add_run()
    r.font.size = Pt(1)
    return p


# ----------------------------------------------------------------- probes
def probe(name, build, expected_pt):
    return {"name": name, "build": build, "expected": expected_pt}


def build_baseline(doc):
    _text(doc, MARK_A)
    _text(doc, MARK_B)


def build_n_spacers(n):
    def b(doc):
        _text(doc, MARK_A)
        for _ in range(n):
            _spacer(doc, 1.0)
        _text(doc, MARK_B)
    return b


def build_space_before(pts):
    def b(doc):
        _text(doc, MARK_A)
        _text(doc, MARK_B, before=pts)
    return b


def build_exact_leading(mult):
    """H2: is lineRule=exact honoured below the font's natural height?"""
    def b(doc):
        _text(doc, MARK_A)
        _text(doc, "filler line one", leading=10.0 * mult)
        _text(doc, "filler line two", leading=10.0 * mult)
        _text(doc, "filler line three", leading=10.0 * mult)
        _text(doc, MARK_B)
    return b


def build_empty_normal_paras(n):
    """Empty paragraphs with NO exact leading -- natural height."""
    def b(doc):
        _text(doc, MARK_A)
        for _ in range(n):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
        _text(doc, MARK_B)
    return b


def build_heading_pair(size_a=22.0, lead_a=25.5, size_b=14.0, lead_b=16.2,
                       before_b=24.1, keepnext=False, outline=None):
    """Reproduce the real c8 case: 22pt heading then a 14pt heading below it.

    LibreOffice renders the gap at 49.4pt (as specified). Docs renders 77.6pt.
    Ablate one property at a time to find which one Docs reacts to.
    """
    def b(doc):
        pa = _text(doc, MARK_A, size=size_a, leading=lead_a)
        pb = _text(doc, MARK_B, size=size_b, leading=lead_b, before=before_b)
        for p, lvl in ((pa, 0), (pb, 1)):
            ppr = p._p.get_or_add_pPr()
            if keepnext:
                kn = OxmlElement("w:keepNext")
                ppr.append(kn)
            if outline is not None:
                ol = OxmlElement("w:outlineLvl")
                ol.set(qn("w:val"), str(lvl))
                ppr.append(ol)
    return b


def build_exact_vs_multiple(size, leading, mode):
    """Same visual intent, three encodings. Which one does Docs honour?

    Google Docs has no 'exact' line spacing in its own model -- only multiples
    -- so the importer must translate w:line/lineRule=exact into something, and
    the translation is where the height is lost.
    """
    def b(doc):
        _text(doc, MARK_A, size=10.0, leading=12.0)
        for _ in range(3):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            if mode == "exact":
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(leading)
            elif mode == "atleast":
                pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                pf.line_spacing = Pt(leading)
            elif mode == "multiple":
                pf.line_spacing = leading / (size * 1.15)   # approx natural
            r = p.add_run("filler text at size %g" % size)
            r.font.size = Pt(size)
            r.font.name = "Arial"
        _text(doc, MARK_B, size=10.0, leading=12.0)
    return b


def build_natural_lines(font, size=20.0, n=4):
    """Measure a font's NATURAL line height in each renderer.

    No line_spacing is set at all, so the gap between the markers is
    n x natural(font, size) and the factor is (gap - 12) / (n x size).
    The gdocs line-height translation divides by this factor, so getting it
    wrong per-font re-introduces the drift the static fix removed.
    """
    def b(doc):
        _text(doc, MARK_A, size=10.0, leading=12.0)
        for _ in range(n):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            r = p.add_run("natural height probe")
            r.font.size = Pt(size)
            r.font.name = font
            rpr = r._element.get_or_add_rPr()
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rpr.append(rf)
            for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rf.set(qn(a), font)
        _text(doc, MARK_B, size=10.0, leading=12.0)
    return b


PROBES = {
    "h5": [
        probe("natural Arial 20pt x4", build_natural_lines("Arial"), None),
        probe("natural TimesNewRoman x4", build_natural_lines("Times New Roman"), None),
        probe("natural CourierNew x4", build_natural_lines("Courier New"), None),
        probe("natural Georgia x4", build_natural_lines("Georgia"), None),
        probe("natural Roboto x4", build_natural_lines("Roboto"), None),
    ],
    "h4": [
        probe("18pt/21pt  exact", build_exact_vs_multiple(18, 21, "exact"), 12 + 63),
        probe("18pt/21pt  atLeast", build_exact_vs_multiple(18, 21, "atleast"), 12 + 63),
        probe("18pt/21pt  multiple", build_exact_vs_multiple(18, 21, "multiple"), 12 + 63),
        probe("22pt/25.5pt exact", build_exact_vs_multiple(22, 25.5, "exact"), 12 + 76.5),
        probe("22pt/25.5pt atLeast", build_exact_vs_multiple(22, 25.5, "atleast"), 12 + 76.5),
        probe("22pt/25.5pt multiple", build_exact_vs_multiple(22, 25.5, "multiple"), 12 + 76.5),
        probe("10pt/12pt  exact", build_exact_vs_multiple(10, 12, "exact"), 12 + 36),
        probe("10pt/12pt  atLeast", build_exact_vs_multiple(10, 12, "atleast"), 12 + 36),
    ],
    "h3": [
        # exact replica of the real failing pair, then ablations
        probe("replica (22pt->14pt, kn+outline)",
              build_heading_pair(keepnext=True, outline=True), 49.6),
        probe("  minus keepNext", build_heading_pair(outline=True), 49.6),
        probe("  minus outlineLvl", build_heading_pair(keepnext=True), 49.6),
        probe("  minus both", build_heading_pair(), 49.6),
        probe("  both sizes 14pt", build_heading_pair(size_a=14.0, lead_a=16.2),
              40.3),
        probe("  first para lead=natural(22)",
              build_heading_pair(lead_a=0), None),
        probe("  second para lead=natural(14)",
              build_heading_pair(lead_b=0), None),
        probe("  before_b=0", build_heading_pair(before_b=0.0), 25.5),
    ],
    "h1": [
        probe("baseline (no gap)", build_baseline, 12.0),
        probe("1 x 1pt exact spacer", build_n_spacers(1), 13.0),
        probe("5 x 1pt exact spacers", build_n_spacers(5), 17.0),
        probe("10 x 1pt exact spacers", build_n_spacers(10), 22.0),
        probe("5 x empty natural paras", build_empty_normal_paras(5), None),
        probe("space_before 20pt", build_space_before(20.0), 32.0),
        probe("space_before 60pt", build_space_before(60.0), 72.0),
    ],
    "h2": [
        probe("exact leading 0.8x", build_exact_leading(0.8), 12 + 3 * 8.0),
        probe("exact leading 1.0x", build_exact_leading(1.0), 12 + 3 * 10.0),
        probe("exact leading 1.2x", build_exact_leading(1.2), 12 + 3 * 12.0),
        probe("exact leading 2.0x", build_exact_leading(2.0), 12 + 3 * 20.0),
    ],
}


def measure_gap(pdf_path):
    """Vertical distance between the two marker baselines, in points."""
    d = fitz.open(pdf_path)
    ya = yb = None
    for p in d:
        for b in p.get_text("dict")["blocks"]:
            if b.get("type"):
                continue
            for ln in b["lines"]:
                t = "".join(s["text"] for s in ln["spans"])
                if MARK_A in t:
                    ya = ln["bbox"][3]
                if MARK_B in t:
                    yb = ln["bbox"][3]
    d.close()
    if ya is None or yb is None:
        return None
    return round(yb - ya, 2)


def main(which=None):
    os.makedirs(OUT, exist_ok=True)
    families = [which] if which else list(PROBES)

    gsvc = None
    try:
        import gdocs_oracle as G
        gsvc = G._service(interactive=False)
    except Exception as e:
        print("(!) Google Docs oracle unavailable: %s" % str(e)[:80])

    for fam in families:
        print("\n=== %s ===" % fam.upper())
        print("%-28s %9s %9s %9s %9s" % ("probe", "expect", "LibreOff", "GDocs",
                                         "Docs-LO"))
        for pr in PROBES[fam]:
            doc = Document()
            _neutralise(doc)
            pr["build"](doc)
            safe = re.sub(r"[^A-Za-z0-9]+", "_", pr["name"]).strip("_")
            dx = os.path.join(OUT, "%s_%s.docx" % (fam, safe))
            doc.save(dx)

            lo = gd = None
            try:
                p = harness.docx_to_pdf(dx, OUT)
                lo = measure_gap(p)
            except Exception:
                pass
            if gsvc is not None:
                try:
                    import gdocs_oracle as G
                    gp = os.path.join(OUT, os.path.basename(dx) + ".gdocs.pdf")
                    G.roundtrip(gsvc, dx, gp)
                    gd = measure_gap(gp)
                except Exception as e:
                    print("   (gdocs fail: %s)" % str(e)[:60])
            exp = pr["expected"]
            print("%-28s %9s %9s %9s %9s" % (
                pr["name"][:28],
                "-" if exp is None else "%.1f" % exp,
                "-" if lo is None else "%.2f" % lo,
                "-" if gd is None else "%.2f" % gd,
                "-" if (lo is None or gd is None) else "%+.2f" % (gd - lo)))


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else None)
