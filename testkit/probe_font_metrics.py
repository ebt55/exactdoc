"""Probe: which Docs-native serif is metrically closest to DejaVu Serif?

    python testkit/probe_font_metrics.py build [out.docx]
    python testkit/probe_font_metrics.py analyse <google-export.pdf>

WHY THIS EXISTS
---------------
l1_word_native maps a DejaVu Serif source onto the closest family this project
can measure offline, Noto Serif, and that leaves a 6.3% advance-width gap. Live
pass 2 confirmed the substitution works -- Docs rendered NotoSerif-Regular and
the packing ratio went 1.2935 -> 1.0637 -- but the residual still moves line
breaks late in a paragraph, and the mechanism that used to close it (run-level
w:spacing) is now known to be discarded by Docs. So the only remaining lever is
a better family, and the families worth trying are Google Fonts whose files are
not installed here and therefore cannot be measured offline at all.

This probe measures them where they actually matter: inside Google Docs.

WHAT IT MEASURES, AND WHY BOTH
------------------------------
Two quantities per family, because getting one right and the other wrong is
exactly how the last attempt regressed:

  advance width  Six fixed segments of `fonts.METRIC_REFERENCE`, each short
                 enough to occupy one line in any plausible family (the widest
                 is 363.9pt at 10pt in DejaVu Serif against a 540pt text width,
                 so there is 1.48x of headroom before anything wraps). Summing
                 rendered widths over summed characters gives advance per
                 character over exactly the reference character set.

  line pitch     The same six segments are ONE paragraph separated by explicit
                 line breaks, with no spacing properties, so the gaps between
                 them are the family's natural line height and nothing else.
                 Substituting Noto Serif without measuring this is what put
                 l1's rendered pitch at 17.48pt against a source 14.70 --
                 `docxout.NATURAL_FACTORS` needs an entry for any family this
                 project adopts, and this is where it comes from.

  glyph face     The font name Docs reports per page, counted over visible
                 characters only. A family Docs does not have is silently
                 substituted, and a measurement of the substitute tells you
                 nothing about the family you asked for.

WHAT v1 GOT WRONG (live pass 3)
-------------------------------
v1 identified a measurement line by its exact text and measured the line's
bounding box. Both halves failed, and the failure is worth keeping written down
because it did not look like a failure -- it looked like a family measuring
-0.91%.

Docs preserves explicit line breaks but appends U+200B (zero width space) at
each one, so five of the six segments no longer compared equal and were
discarded. The sixth, the last line, has no break after it and did match -- but
its bounding box also contained a trailing space and an 11pt Cambria paragraph
mark, both of which are wider than the 10pt glyphs they follow. Measuring one
short line with that tail attached inflated every family by 3-5%, and inflated
the shortest segment most, which is exactly the kind of error that survives
review because all nine numbers move together.

v2 therefore measures from WORD boxes: the extent runs from the first word's x0
to the last word's x1, so it ends at the last glyph and cannot include a
trailing space or a paragraph mark, and the line's text is rebuilt from the
words themselves so zero-width marks never enter the comparison. Re-run against
the pass-3 export, the three controls came back +0.58/+0.78/+0.76% instead of
+3.57/+4.33/+5.28%.

The residual +0.7% is systematic across all three controls (spread 0.2%), so
`decide` calibrates on them: every candidate is divided by the controls' mean
offset before it is judged. That makes the comparison independent of whatever
causes the offset, and it is itself checkable -- Noto Serif calibrates to
-6.30% against the -6.18% measured from its own font file.

Controls: Noto Serif (today's substitute), Times New Roman (the original
mapping) and Georgia are all measurable offline, so their probe numbers say
whether the analyzer and the round trip can be trusted before any candidate is
believed.

Like testkit/probe_cover_band.py and testkit/docs_quirks.py this shares no code
with the converter under test.

HOW TO RUN IT (rides along with the next consented pass -- uploads nothing)
--------------------------------------------------------------------------
1.  python testkit/probe_font_metrics.py build
    Writes testkit/quirks/probe_font_metrics.docx, the gitignored directory
    docs_quirks.py already uses, so building cannot dirty the tree.
2.  During the next CONSENTED live pass, upload that .docx to Google Docs
    alongside the corpus and export it back as PDF. This step needs the user's
    consent; the probe never contacts anything by itself.
3.  python testkit/probe_font_metrics.py analyse <export.pdf>

Step 3 prints per-family advance ratio, derived natural line-height factor,
whether Docs substituted the face, and the decision the rule below reaches.
Nothing here touches the gate, the oracle, or the writer.
"""
import argparse
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

PAGE_W, PAGE_H = 612.0, 792.0
MARGIN = 36.0
PROBE_SIZE = 10.0
MARK_RE = re.compile(r"FMPROBE(\d{2})")

# Every candidate is asserted Docs-native by exactdoc.fonts.GDOCS_NATIVE. The
# serifs chosen are those with a reputation for wide, large-x-height letterforms,
# because the gap to close runs in that direction: DejaVu Serif is wider than
# every family measurable here.
CANDIDATES = (
    "Noto Serif",           # control: today's substitute, 0.491413 offline
    "Times New Roman",      # control: the original mapping, 0.410575 offline
    "Georgia",              # control: 0.451187 offline
    "Merriweather",
    "Bitter",
    "Libre Baskerville",
    "PT Serif",
    "Lora",
    "IBM Plex Serif",
)

# Offline-measured advance per character at 1pt over CONCATENATED SEGMENTS (not
# over METRIC_REFERENCE itself -- the segments drop the spaces at their joins,
# so the two differ in the fourth decimal). Measured with
# fitz.Font(fontfile=...).text_length(joined, fontsize=1.0)/len(joined).
TARGET_ADVANCE = 0.523808          # DejaVu Serif: what l1's source actually used
# fonts.FAMILY_METRICS is measured over METRIC_REFERENCE itself, whose spaces at
# the segment joins this probe drops; DejaVu Serif reads 0.520900 there against
# 0.523808 here. Multiply a probe advance by this to put it on that scale.
REFERENCE_SCALE = 0.520900 / 0.523808
KNOWN_ADVANCE = {                  # controls, for validating the round trip
    "Noto Serif": 0.491413,
    "Times New Roman": 0.410575,
    "Georgia": 0.451187,
}
# Docs-measured natural line heights already in docxout.NATURAL_FACTORS, for the
# same purpose on the pitch side.
KNOWN_FACTOR = {"Noto Serif": 1.360, "Times New Roman": 1.144, "Georgia": 1.130}

MAX_SEG_CHARS = 68

# A copy of exactdoc.fonts.METRIC_REFERENCE, deliberately not imported. Like
# testkit/docs_quirks.py and testkit/harness.py, this probe shares no code with
# the converter it is measuring for -- and importing `exactdoc` from a testkit
# script resolves against whatever checkout happens to be on sys.path, which is
# not necessarily this worktree. tests/test_font_metrics_probe.py asserts the
# two strings stay identical, so the duplication cannot drift silently.
METRIC_REFERENCE = (
    "The quick brown fox jumps over the lazy dog. Modern inference workloads "
    "exhibit sharply bimodal traffic patterns, with sustained baseline demand "
    "punctuated by bursts that exceed steady-state volume by an order of "
    "magnitude; provisioning for peak wastes capacity, and provisioning for "
    "baseline degrades latency guarantees precisely when demand is highest.")


def segments(text=None):
    """METRIC_REFERENCE split at word boundaries into one-line pieces."""
    if text is None:
        text = METRIC_REFERENCE
    words, out, cur = text.split(), [], ""
    for word in words:
        cand = (cur + " " + word).strip()
        if len(cand) > MAX_SEG_CHARS and cur:
            out.append(cur)
            cur = word
        else:
            cur = cand
    if cur:
        out.append(cur)
    return out


OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "quirks")
DEFAULT_OUT = os.path.join(OUT_DIR, "probe_font_metrics.docx")


def marker(index):
    return "FMPROBE%02d" % index


def parse_marker(text):
    m = MARK_RE.search(text or "")
    return int(m.group(1)) if m else None


# ------------------------------------------------------------------ build
def _family_run(run, family, size):
    run.font.size = Pt(size)
    run.font.name = family
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), family)


def _neutral(par):
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    return par


def build(out_path=None):
    """Write the probe DOCX, one candidate family per page. Returns the path."""
    out_path = out_path or DEFAULT_OUT
    segs = segments()
    doc = Document()
    npf = doc.styles["Normal"].paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)
    npf.line_spacing = 1.0

    for i, family in enumerate(CANDIDATES):
        if i == 0:
            sec = doc.sections[0]
        else:
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
        sec.page_width = Pt(PAGE_W)
        sec.page_height = Pt(PAGE_H)
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(sec, attr, Pt(MARGIN))

        # The marker is deliberately in a family Docs certainly has: if the
        # candidate is missing, the page must still identify itself.
        head = _neutral(doc.add_paragraph())
        _family_run(head.add_run("%s  %s" % (marker(i), family)), "Arial", 9.0)

        # One paragraph, explicit line breaks: the gaps between these lines are
        # the family's natural line height with no paragraph spacing involved.
        body = _neutral(doc.add_paragraph())
        for j, seg in enumerate(segs):
            if j:
                br = body.add_run()
                br.add_break(WD_BREAK.LINE)
                _family_run(br, family, PROBE_SIZE)
            _family_run(body.add_run(seg), family, PROBE_SIZE)

    parent = os.path.dirname(os.path.abspath(out_path))
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------- analyse
def _norm(name):
    return re.sub(r"[^a-z]", "", (name or "").lower())


# Docs appends U+200B at every explicit line break, and soft hyphens and other
# zero-width marks travel the same way. None of them are glyphs to measure.
_ZERO_WIDTH = re.compile(r"[​‌‍﻿­]")


def analyse(pdf_path):
    """[{family, advance, ratio, pitch, factor, face, substituted}] per page."""
    import fitz          # imported here so `build` works without PyMuPDF

    by_text = {s: len(s) for s in segments()}
    doc = fitz.open(pdf_path)
    rows = []
    try:
        for pno in range(doc.page_count):
            page = doc[pno]
            idx = parse_marker(page.get_text("text"))
            row = {"page": pno + 1, "family": None}
            if idx is None or idx >= len(CANDIDATES):
                rows.append(row)
                continue
            family = CANDIDATES[idx]
            row["family"] = family

            # Word boxes, grouped into lines. The extent of a line is first
            # word x0 to last word x1: it ends at the last glyph, so a trailing
            # space or a paragraph mark cannot widen it.
            lines = {}
            for x0, y0, x1, y1, word, blk, lno, _wn in page.get_text("words"):
                clean = _ZERO_WIDTH.sub("", word)
                if clean:
                    lines.setdefault((blk, lno), []).append((x0, y0, x1, clean))

            width_sum = char_sum = 0.0
            tops = []
            for parts in lines.values():
                parts.sort(key=lambda t: t[0])
                text = " ".join(t[3] for t in parts)
                if text not in by_text:
                    continue
                width_sum += max(t[2] for t in parts) - min(t[0] for t in parts)
                char_sum += by_text[text]
                tops.append(min(t[1] for t in parts))

            # Faces, counted over visible characters at the probe size only, so
            # an empty span cannot masquerade as a substitution.
            faces = {}
            for b in page.get_text("dict")["blocks"]:
                if b.get("type"):
                    continue
                for ln in b["lines"]:
                    for s in ln["spans"]:
                        if abs(s["size"] - PROBE_SIZE) > 0.3:
                            continue
                        n = len(_ZERO_WIDTH.sub("", s["text"]).strip())
                        if n:
                            faces[s["font"]] = faces.get(s["font"], 0) + n

            row["lines_found"] = len(tops)
            if char_sum:
                row["advance"] = width_sum / char_sum / PROBE_SIZE
                row["ratio"] = row["advance"] / TARGET_ADVANCE
            if len(tops) > 1:
                tops.sort()
                gaps = sorted(b - a for a, b in zip(tops, tops[1:]) if b > a)
                row["pitch"] = gaps[len(gaps) // 2]
                row["factor"] = row["pitch"] / PROBE_SIZE
            if faces:
                face = max(faces, key=faces.get)
                row["face"] = face
                row["substituted"] = _norm(family) not in _norm(face)
                row["face_share"] = faces[face] / float(sum(faces.values()))
            rows.append(row)
    finally:
        doc.close()
    return rows


def calibration(rows):
    """Mean (measured / offline) over the control families, or None.

    The controls are the only families whose true advance is known here, so
    they are what turns a raw reading into a comparable one. Returning None
    rather than 1.0 when they are missing keeps "uncalibrated" distinguishable
    from "calibrated to no correction".
    """
    offsets = [r["advance"] / KNOWN_ADVANCE[r["family"]]
               for r in rows
               if r.get("advance") and r.get("family") in KNOWN_ADVANCE
               and not r.get("substituted")]
    if not offsets:
        return None
    return sum(offsets) / len(offsets)


# ------------------------------------------------------------ decision rule
# A family is adoptable only if it is measurably better on advance width AND
# Docs actually rendered it AND its natural line height came back, because
# adopting a family without a docxout.NATURAL_FACTORS entry is what turned the
# last substitution into a 19% line-pitch error.
#
# ADOPT_DEV: the mappings that already work sit at 0.0% (the Liberation faces
# are metric clones of their Microsoft counterparts) and the worst honest
# residual in the corpus is DejaVu Sans Mono -> Courier New at -0.32%. Over
# l1's 85-character line, 2% is about 1.7 characters -- below the level that
# systematically moves a break.
ADOPT_DEV = 0.02
# Failing that, a candidate still earns the swap if it at least halves the
# deviation Noto Serif leaves (6.2% over these segments).
IMPROVE_FACTOR = 0.5
# If the best measured candidate is still worse than this, no family available
# to Docs closes the gap and the question is settled the other way.
CLOSE_BOOK_DEV = 0.045
NOTO_DEV = abs(KNOWN_ADVANCE["Noto Serif"] / TARGET_ADVANCE - 1.0)


def decide(rows):
    """Apply the documented rule. Returns a verdict dict.

    Judged on CALIBRATED ratios: the controls' mean offset is divided out
    first, so the comparison does not depend on whatever small systematic the
    round trip introduces. Uncalibrated raw ratios are kept for reporting.
    """
    cal = calibration(rows)
    for r in rows:
        if r.get("ratio"):
            r["cal_ratio"] = r["ratio"] / cal if cal else r["ratio"]
    usable = [r for r in rows
              if r.get("cal_ratio") and not r.get("substituted")
              and r.get("factor")]
    scored = sorted(usable, key=lambda r: abs(r["cal_ratio"] - 1.0))
    best = scored[0] if scored else None
    verdict = {"n_measured": len(rows), "n_usable": len(usable),
               "noto_dev": NOTO_DEV, "best": best, "calibration": cal}
    if best is None:
        verdict.update(action="inconclusive",
                       reason="no candidate produced a usable measurement")
        return verdict
    dev = abs(best["cal_ratio"] - 1.0)
    verdict["best_dev"] = dev
    if best["family"] == "Noto Serif":
        verdict.update(action="keep-noto-serif",
                       reason="no candidate beat the family already in use")
    elif dev <= ADOPT_DEV:
        verdict.update(action="adopt",
                       reason="deviation %.2f%% is inside the %.0f%% bar the "
                              "already-working mappings sit at"
                              % (100 * dev, 100 * ADOPT_DEV))
    elif dev <= NOTO_DEV * IMPROVE_FACTOR:
        verdict.update(action="adopt",
                       reason="deviation %.2f%% at least halves Noto Serif's "
                              "%.2f%%" % (100 * dev, 100 * NOTO_DEV))
    elif dev >= CLOSE_BOOK_DEV:
        verdict.update(action="close-the-book",
                       reason="best available deviation %.2f%% is no real "
                              "improvement on Noto Serif's %.2f%%; the residual "
                              "is inherent and l1's dx bound needs a ratified "
                              "waiver rather than another substitution"
                              % (100 * dev, 100 * NOTO_DEV))
    else:
        verdict.update(action="marginal",
                       reason="deviation %.2f%% beats Noto Serif but not by the "
                              "margin that justifies changing typeface; decide "
                              "on visual grounds" % (100 * dev))
    return verdict


def render(rows, verdict=None):
    cal = (verdict or {}).get("calibration")
    L = ["%-20s %6s %9s %9s %9s %8s  %s"
         % ("family", "lines", "adv@1pt", "raw dev", "calib dev", "factor",
            "face"),
         "-" * 92]
    for r in rows:
        if not r.get("family"):
            L.append("page %d: no probe marker found" % r["page"])
            continue
        if not r.get("ratio"):
            L.append("%-20s %6s  (no measurement lines matched)"
                     % (r["family"], r.get("lines_found", 0)))
            continue
        cr = r.get("cal_ratio")
        L.append("%-20s %6d %9.6f %+8.2f%% %+8s %8.3f  %s%s" % (
            r["family"], r["lines_found"], r["advance"],
            100 * (r["ratio"] - 1),
            "-" if cr is None else "%.2f%%" % (100 * (cr - 1)),
            r.get("factor", float("nan")),
            r.get("face", "?"), "  SUBSTITUTED" if r.get("substituted") else ""))
    L.append("-" * 92)
    L.append("dev 0.00% = the family matches DejaVu Serif, l1's source face")
    if cal:
        L.append("calibrated on the controls (offset %+.2f%%); judge on that "
                 "column" % (100 * (cal - 1)))
    else:
        L.append("NO CALIBRATION: no control measured, so the raw column is all "
                 "there is and candidates cannot be trusted against it")
    if any(r.get("lines_found") not in (None, len(segments())) for r in rows):
        L.append("NOTE: a family with fewer than %d matched lines lost segments "
                 "to the extractor; its advance is measured over less text and "
                 "its pitch may be missing." % len(segments()))

    controls = [r for r in rows if r.get("family") in KNOWN_ADVANCE
                and r.get("advance")]
    if controls:
        L.append("")
        L.append("control check (these three are measurable offline):")
        for r in controls:
            want = KNOWN_ADVANCE[r["family"]]
            wf = KNOWN_FACTOR.get(r["family"])
            L.append("  %-18s advance %.6f vs %.6f offline (%+.2f%%)%s"
                     % (r["family"], r["advance"], want,
                        100 * (r["advance"] / want - 1),
                        "" if wf is None or not r.get("factor") else
                        "   factor %.3f vs %.3f" % (r["factor"], wf)))
        L.append("  A control that disagrees means the round trip or the")
        L.append("  analyzer is wrong, and no candidate number can be believed.")

    if verdict:
        L.append("")
        L.append("decision: %s" % verdict["action"].upper())
        L.append("  %s" % verdict["reason"])
        if verdict.get("best"):
            b = verdict["best"]
            L.append("  best usable candidate: %s (calibrated %+.2f%%, factor %.3f)"
                     % (b["family"], 100 * (b.get("cal_ratio", 1) - 1),
                        b.get("factor", float("nan"))))
        if verdict["action"] == "adopt":
            b = verdict["best"]
            adv_ref = (b["advance"] / (verdict.get("calibration") or 1.0)
                       * REFERENCE_SCALE)
            L.append("")
            L.append("  To act, register %r in all THREE places:" % b["family"])
            L.append("    1. fonts.FAMILY_METRICS[%r] = (%.6f, 'serif')"
                     % (_norm(b["family"]), adv_ref))
            L.append("       (calibrated advance rescaled from the probe's")
            L.append("       concatenated segments onto METRIC_REFERENCE, which")
            L.append("       is what the rest of that table is measured over)")
            L.append("    2. fonts._CANDIDATES['serif'] -- add the family name")
            L.append("    3. docxout.NATURAL_FACTORS[%r] = %.3f"
                     % (b["family"].lower(), b.get("factor", float("nan"))))
            L.append("  The test tying FAMILY_METRICS to NATURAL_FACTORS fails")
            L.append("  if you do 1 and 2 without 3 -- which is the mistake that")
            L.append("  put l1's line pitch 19% out in the first place.")
    return "\n".join(L)


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    sub = ap.add_subparsers(dest="cmd", required=True)
    b = sub.add_parser("build", help="write the probe DOCX (uploads nothing)")
    b.add_argument("out", nargs="?", default=DEFAULT_OUT)
    a = sub.add_parser("analyse", help="read a Google export of the probe")
    a.add_argument("pdf")
    args = ap.parse_args(argv)

    if args.cmd == "build":
        path = build(args.out)
        print("wrote %s (%d families, %d measurement lines each)"
              % (path, len(CANDIDATES), len(segments())))
        print("Upload it during the next CONSENTED pass, export PDF, then:")
        print("  python testkit/probe_font_metrics.py analyse <export.pdf>")
        return 0
    if not os.path.isfile(args.pdf):
        sys.stderr.write("error: no such file: %s\n" % args.pdf)
        return 2
    rows = analyse(args.pdf)
    print(render(rows, decide(rows)))
    return 0


if __name__ == "__main__":
    sys.exit(main())
