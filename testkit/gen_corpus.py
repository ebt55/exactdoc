"""Adversarial multi-producer PDF corpus for PDF->DOCX fidelity testing.

Producers: Chromium/Skia (headless Chrome), ReportLab, fpdf2, LibreOffice.
The Chromium set is styled like Claude-generated whitepapers/reports.
"""
import _paths  # noqa: F401  (sets sys.path, finds soffice/chrome)
from _paths import CHROME, SOFFICE
import os, sys, subprocess, textwrap, shutil

OUT = sys.argv[1] if len(sys.argv) > 1 else "adv"
HTML = os.path.join(OUT, "_html")
os.makedirs(HTML, exist_ok=True)

LOREM = ("Retrieval quality degrades non-linearly as the corpus grows past the "
         "point where the embedding model was calibrated, and the failure is "
         "quiet: recall stays flat while precision collapses in the tail. "
         "Teams that measure only average relevance will not see it. ")
LOREM2 = ("The mitigation is unglamorous. Chunk boundaries must respect document "
          "structure rather than token counts, and the reranker has to be trained "
          "on the same distribution it will serve. Everything else is tuning. ")

BASE_CSS = """
@page { size: Letter; margin: 0.9in 0.85in; }
* { box-sizing: border-box; }
body { font-family: Georgia, 'Times New Roman', serif; font-size: 10.5pt;
       line-height: 1.45; color: #1a1a1a; margin: 0; }
h1 { font-family: Helvetica, Arial, sans-serif; font-size: 22pt; margin: 0 0 6pt; }
h2 { font-family: Helvetica, Arial, sans-serif; font-size: 14pt; margin: 18pt 0 6pt;
     color: #123a5e; }
h3 { font-family: Helvetica, Arial, sans-serif; font-size: 11.5pt; margin: 12pt 0 4pt; }
p  { margin: 0 0 8pt; text-align: justify; }
ul, ol { margin: 0 0 8pt 0; padding-left: 20pt; }
li { margin-bottom: 4pt; }
table { border-collapse: collapse; width: 100%; font-size: 9.5pt; margin: 8pt 0 12pt; }
th { background: #123a5e; color: #fff; text-align: left; padding: 5pt 7pt;
     font-family: Helvetica, Arial, sans-serif; font-size: 9pt; }
td { padding: 4pt 7pt; border-bottom: 0.5pt solid #d8dee5; }
tr.alt td { background: #f2f5f8; }
.callout { border-left: 3pt solid #2b7a4b; background: #eef7f1; padding: 8pt 12pt;
           margin: 10pt 0; }
.warn { border-left: 3pt solid #b3541e; background: #fdf1e8; padding: 8pt 12pt;
        margin: 10pt 0; }
pre { font-family: 'Courier New', monospace; font-size: 8.5pt; background: #f6f8fa;
      border: 0.5pt solid #d8dee5; padding: 8pt 10pt; line-height: 1.35;
      white-space: pre; margin: 8pt 0; }
.band { background: #123a5e; color: #fff; padding: 26pt 24pt; margin: -0.9in -0.85in 18pt;
        font-family: Helvetica, Arial, sans-serif; }
.band .sub { font-size: 11pt; opacity: .85; margin-top: 4pt; }
.cards { display: flex; gap: 10pt; margin: 12pt 0; }
.card { flex: 1; border: 0.5pt solid #c9d3dd; border-radius: 4pt; padding: 10pt;
        text-align: center; background: #fafcfe; }
.card .n { font-size: 19pt; font-weight: bold; color: #123a5e;
           font-family: Helvetica, Arial, sans-serif; }
.card .l { font-size: 8pt; color: #556; text-transform: uppercase; }
blockquote { border-left: 2pt solid #999; margin: 10pt 0 10pt 12pt; padding-left: 12pt;
             font-style: italic; color: #444; }
.footnote { font-size: 8pt; color: #444; border-top: 0.5pt solid #ccc;
            padding-top: 4pt; margin-top: 10pt; }
"""


def html_doc(name, body, extra_css=""):
    p = os.path.join(HTML, name + ".html")
    with open(p, "w", encoding="utf-8") as f:
        f.write("<!doctype html><meta charset='utf-8'><style>%s\n%s</style>\n%s"
                % (BASE_CSS, extra_css, body))
    return p


def chrome_pdf(html_path, out_pdf):
    url = "file:///" + os.path.abspath(html_path).replace("\\", "/")
    cmd = [CHROME, "--headless=new", "--disable-gpu", "--no-sandbox",
           "--no-pdf-header-footer", "--run-all-compositor-stages-before-draw",
           "--virtual-time-budget=4000",
           "--print-to-pdf=" + os.path.abspath(out_pdf), url]
    subprocess.run(cmd, capture_output=True, timeout=180)
    return os.path.exists(out_pdf)


# ------------------------------------------------------------------ documents
def c1_whitepaper():
    b = """
<div class="band"><div style="font-size:26pt;font-weight:bold">Retrieval at Scale</div>
<div class="sub">Why production RAG degrades quietly &mdash; and what to measure instead</div>
<div class="sub" style="font-size:9pt">Technical Whitepaper &middot; July 2026</div></div>
<h2>1. Executive summary</h2>
<p>%s%s</p>
<div class="cards">
  <div class="card"><div class="n">41%%</div><div class="l">precision drop</div></div>
  <div class="card"><div class="n">2.3x</div><div class="l">tail latency</div></div>
  <div class="card"><div class="n">$0.72</div><div class="l">cost per query</div></div>
</div>
<h2>2. Where the failure hides</h2>
<p>%s</p>
<ul>
  <li>Average relevance is a lagging indicator; it moves last.</li>
  <li>Tail queries carry most of the business value and least of the traffic.</li>
  <li>Chunk boundaries drawn on token counts sever the exact clauses users ask about.</li>
</ul>
<div class="callout"><b>Key finding.</b> Precision in the bottom decile of query
frequency fell 41%% while mean nDCG moved less than two points.</div>
<h2>3. Measured results</h2>
<table>
<tr><th>Configuration</th><th>nDCG@10</th><th>Tail precision</th><th>p95 latency</th></tr>
<tr><td>Baseline dense</td><td>0.612</td><td>0.301</td><td>410 ms</td></tr>
<tr class="alt"><td>+ structural chunking</td><td>0.634</td><td>0.447</td><td>425 ms</td></tr>
<tr><td>+ in-domain reranker</td><td>0.671</td><td>0.588</td><td>690 ms</td></tr>
<tr class="alt"><td>+ query rewriting</td><td>0.688</td><td>0.601</td><td>940 ms</td></tr>
</table>
<div class="warn"><b>Caveat.</b> Latency is measured on warm caches. Cold-start
numbers are roughly 1.8x higher and are reported in Appendix B.</div>
<h2>4. Recommendations</h2>
<ol>
  <li>Instrument tail precision as a first-class SLO.</li>
  <li>Re-chunk on document structure, not token windows.</li>
  <li>Train the reranker on the served distribution.</li>
</ol>
<blockquote>The systems that failed did not fail loudly. They returned
plausible documents that were subtly wrong.</blockquote>
<p>%s</p>
""" % (LOREM, LOREM2, LOREM, LOREM2)
    return html_doc("c1_whitepaper", b)


def c2_paper2col():
    css = """
.two { column-count: 2; column-gap: 22pt; }
.title { text-align: center; font-family: Helvetica, Arial, sans-serif; }
.abstract { font-style: italic; margin: 0 42pt 14pt; text-align: justify; font-size: 9.5pt; }
sup { font-size: 7pt; }
"""
    b = """
<div class="title"><div style="font-size:17pt;font-weight:bold">Structural Chunking for
Long-Context Retrieval</div>
<div style="font-size:10pt;margin-top:6pt">A. Researcher<sup>1</sup>, B. Coauthor<sup>2</sup></div>
<div style="font-size:9pt;color:#555">1. Institute of Things &nbsp; 2. Other Lab</div></div>
<hr>
<div class="abstract"><b>Abstract.</b> %s%s</div>
<div class="two">
<h3>1 Introduction</h3><p>%s</p><p>%s</p>
<h3>2 Method</h3><p>%s</p>
<p>We define the boundary score <i>s(i)</i> as a weighted sum of heading depth and
sentence terminality<sup>3</sup>.</p>
<p>%s</p>
<h3>3 Results</h3><p>%s</p><p>%s</p>
<h3>4 Discussion</h3><p>%s</p>
<div class="footnote">3. Terminality is estimated with a small classifier;
see supplementary material.</div>
</div>
""" % (LOREM, LOREM2, LOREM, LOREM2, LOREM, LOREM2, LOREM, LOREM2, LOREM)
    return html_doc("c2_paper2col", b, css)


def c3_tables():
    b = """
<h1>Quarterly Operations Review</h1>
<h2>Regional breakdown (merged headers)</h2>
<table>
<tr><th rowspan="2">Region</th><th colspan="3">Revenue</th><th colspan="2">Headcount</th></tr>
<tr><th>Q1</th><th>Q2</th><th>Q3</th><th>Eng</th><th>Ops</th></tr>
<tr><td>North America</td><td>4.10</td><td>4.55</td><td>5.02</td><td>112</td><td>44</td></tr>
<tr class="alt"><td>EMEA</td><td>2.88</td><td>3.01</td><td>3.44</td><td>78</td><td>31</td></tr>
<tr><td>APAC</td><td>1.92</td><td>2.30</td><td>2.71</td><td>55</td><td>19</td></tr>
<tr class="alt"><td><b>Total</b></td><td><b>8.90</b></td><td><b>9.86</b></td>
<td><b>11.17</b></td><td><b>245</b></td><td><b>94</b></td></tr>
</table>
<h2>Nested detail</h2>
<table><tr><th>Programme</th><th>Detail</th></tr>
<tr><td>Platform</td><td>
  <table style="margin:0"><tr><td>Ingest</td><td>green</td></tr>
  <tr><td>Index</td><td>amber</td></tr><tr><td>Serve</td><td>green</td></tr></table>
</td></tr>
<tr class="alt"><td>Safety</td><td>Review scheduled for week 34.</td></tr></table>
<h2>Long table spanning a page break</h2>
<table><tr><th>#</th><th>Item</th><th>Owner</th><th>Status</th></tr>
%s
</table>
""" % "".join("<tr%s><td>%d</td><td>Work item number %d in the backlog</td>"
              "<td>Team %s</td><td>%s</td></tr>" %
              (" class='alt'" if i % 2 else "", i, i, "ABCDE"[i % 5],
               ["open", "in review", "done"][i % 3]) for i in range(1, 46))
    return html_doc("c3_tables", b)


def c4_i18n():
    css = "body{font-size:11pt} .rtl{direction:rtl;text-align:right}"
    b = """
<h1>Multilingual Coverage Test</h1>
<h2>Chinese (Simplified)</h2>
<p>检索质量随着语料库规模的增长呈非线性下降，而且这种失败是安静的：
召回率保持平稳，精确率却在长尾崩塌。只测量平均相关性的团队看不到这一点。</p>
<h2>Japanese</h2>
<p>コーパスが埋め込みモデルの較正点を超えて大きくなると、検索品質は
非線形に低下します。しかもその失敗は静かです。</p>
<h2>Korean</h2>
<p>말뭉치가 임베딩 모델이 보정된 지점을 넘어 커지면 검색 품질은
비선형적으로 저하됩니다.</p>
<h2>Arabic (RTL)</h2>
<p class="rtl">تتدهور جودة الاسترجاع بشكل غير خطي مع نمو مجموعة النصوص
إلى ما بعد النقطة التي تمت فيها معايرة نموذج التضمين.</p>
<h2>Hebrew (RTL)</h2>
<p class="rtl">איכות האחזור יורדת באופן לא ליניארי ככל שהקורפוס גדל
מעבר לנקודה שבה כויל מודל ההטמעה.</p>
<h2>Accented Latin &amp; symbols</h2>
<p>Voilà — naïve façade, Ærøskøbing, Straße, ¿cómo estás? ½ ¾ × ÷ ≤ ≥ ≠ ∑ ∫ √ π µ Ω
&nbsp;•&nbsp;→&nbsp;⇒&nbsp;★</p>
"""
    return html_doc("c4_i18n", b, css)


def c5_graphics():
    css = """
.grad { background: linear-gradient(120deg,#123a5e,#2b7a9b); color:#fff;
        padding: 20pt; margin: 0 0 14pt; border-radius: 6pt; }
.round { border: 1pt solid #123a5e; border-radius: 10pt; padding: 12pt; margin: 10pt 0; }
"""
    b = """
<div class="grad"><b style="font-size:18pt">Gradient Band</b>
<div>Vector gradients have no OOXML paragraph equivalent.</div></div>
<h2>Inline SVG chart</h2>
<svg width="440" height="200" xmlns="http://www.w3.org/2000/svg">
  <line x1="40" y1="170" x2="430" y2="170" stroke="#333" stroke-width="1"/>
  <line x1="40" y1="10" x2="40" y2="170" stroke="#333" stroke-width="1"/>
  <rect x="60"  y="90"  width="40" height="80"  fill="#123a5e"/>
  <rect x="130" y="60"  width="40" height="110" fill="#2b7a9b"/>
  <rect x="200" y="40"  width="40" height="130" fill="#2b7a4b"/>
  <rect x="270" y="105" width="40" height="65"  fill="#b3541e"/>
  <rect x="340" y="25"  width="40" height="145" fill="#6b3fa0"/>
  <text x="62"  y="185" font-size="10">Q1</text>
  <text x="132" y="185" font-size="10">Q2</text>
  <text x="202" y="185" font-size="10">Q3</text>
  <text x="272" y="185" font-size="10">Q4</text>
  <text x="342" y="185" font-size="10">Q5</text>
  <text x="10"  y="174" font-size="9">0</text>
  <text x="10"  y="95"  font-size="9">50</text>
  <text x="6"   y="16"  font-size="9">100</text>
</svg>
<div class="round"><b>Rounded card.</b> Border radius is not expressible on a
Word table cell; the corners must either be squared off or rasterized.</div>
<h2>Rotated text</h2>
<div style="transform:rotate(-90deg);transform-origin:left top;height:120pt;
   font-family:Helvetica;font-size:12pt;color:#123a5e">Rotated axis label</div>
<h2>Body after graphics</h2>
<p>%s</p>
""" % LOREM
    return html_doc("c5_graphics", b, css)


def c6_long():
    secs = []
    for i in range(1, 26):
        secs.append("<h2>%d. Section heading number %d</h2><p>%s</p><p>%s</p>"
                    "<ul><li>Point one for section %d.</li>"
                    "<li>Point two for section %d, somewhat longer so that it "
                    "wraps onto a second line in most column widths.</li></ul>"
                    % (i, i, LOREM, LOREM2, i, i))
    return html_doc("c6_long", "<h1>Long Document Stress Test</h1>" + "".join(secs))


def c7_code():
    b = """
<h1>Implementation Notes</h1>
<p>%s</p>
<pre>def rerank(query: str, docs: list[Doc]) -&gt; list[Doc]:
    \"\"\"Score each doc against the query with the in-domain cross-encoder.\"\"\"
    pairs = [(query, d.text) for d in docs]

    scores = model.predict(pairs, batch_size=32)

    order = sorted(range(len(docs)), key=lambda i: -scores[i])
    return [docs[i] for i in order]

# NOTE: batch_size above is tuned for a 24GB card; halve it on 12GB.
</pre>
<h2>A longer listing</h2>
<pre>class StructuralChunker:
    def __init__(self, max_tokens: int = 512, overlap: int = 48):
        self.max_tokens, self.overlap = max_tokens, overlap

    def split(self, doc):
        out, buf = [], []
        for node in doc.walk():
            if node.is_heading and buf:
                out.append(self._flush(buf)); buf = []
            buf.append(node)
        if buf:
            out.append(self._flush(buf))
        return out
</pre>
<p>%s</p>
""" % (LOREM, LOREM2)
    return html_doc("c7_code", b)


def c8_toc_links():
    css = """
.toc a { text-decoration: none; color: #123a5e; }
.toc li { margin-bottom: 3pt; }
"""
    b = """
<h1>Design Document</h1>
<h2>Contents</h2>
<ol class="toc">
  <li><a href="#s1">Motivation</a></li>
  <li><a href="#s2">Architecture</a></li>
  <li><a href="#s3">Evaluation</a></li>
</ol>
<h2 id="s1">1. Motivation</h2><p>%s See <a href="https://example.com/spec">the
specification</a> and the <a href="https://example.com/rfc-2119">RFC</a>.</p>
<h2 id="s2">2. Architecture</h2><p>%s</p>
<p>Contact <a href="mailto:team@example.com">team@example.com</a> for access.</p>
<h2 id="s3">3. Evaluation</h2><p>%s</p>
""" % (LOREM, LOREM2, LOREM)
    return html_doc("c8_toc_links", b, css)


# ------------------------------------------------------------- other producers
def r1_reportlab():
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, ListFlowable, ListItem)
    out = os.path.join(OUT, "r1_reportlab_report.pdf")
    ss = getSampleStyleSheet()
    body = ParagraphStyle("b", parent=ss["BodyText"], fontName="Times-Roman",
                          fontSize=10.5, leading=14.5, alignment=4, spaceAfter=8)
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontName="Helvetica-Bold",
                        fontSize=18, spaceAfter=10)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontName="Helvetica-Bold",
                        fontSize=13, textColor=HexColor("#123a5e"), spaceAfter=6)
    doc = SimpleDocTemplate(out, pagesize=LETTER, leftMargin=0.9 * inch,
                            rightMargin=0.9 * inch, topMargin=0.8 * inch,
                            bottomMargin=0.8 * inch)
    story = [Paragraph("Vector Index Maintenance", h1),
             Paragraph("1. Background", h2), Paragraph(LOREM + LOREM2, body),
             Paragraph("2. Findings", h2), Paragraph(LOREM2, body),
             ListFlowable([ListItem(Paragraph("Rebuild cost scales with deletions.", body)),
                           ListItem(Paragraph("Tombstones dominate memory after week six.", body)),
                           ListItem(Paragraph("Compaction windows must be scheduled off-peak.", body))],
                          bulletType="bullet", leftIndent=18),
             Spacer(1, 8)]
    data = [["Shard", "Vectors", "Deleted", "Rebuild (s)"],
            ["s-01", "4,100,220", "812,004", "196"],
            ["s-02", "3,980,117", "1,204,551", "244"],
            ["s-03", "4,210,880", "98,220", "121"]]
    t = Table(data, colWidths=[90, 110, 100, 90])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#123a5e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.4, HexColor("#c9d3dd")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#ffffff"), HexColor("#f2f5f8")]),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6)]))
    story += [t, Spacer(1, 10), Paragraph("3. Next steps", h2), Paragraph(LOREM, body)]
    doc.build(story)
    return out


def f1_fpdf():
    from fpdf import FPDF
    out = os.path.join(OUT, "f1_fpdf_brief.pdf")
    p = FPDF(format="letter", unit="pt")
    p.set_auto_page_break(True, margin=56)
    p.add_page()
    p.set_font("Helvetica", "B", 18); p.cell(0, 24, "Operations Brief", ln=1)
    p.set_font("Times", "", 10.5)
    p.multi_cell(0, 14, LOREM + LOREM2, align="J")
    p.ln(6)
    p.set_font("Helvetica", "B", 12); p.cell(0, 18, "Status table", ln=1)
    p.set_font("Helvetica", "", 9)
    for row in [("Region", "Status", "Owner"), ("NA", "green", "Team A"),
                ("EMEA", "amber", "Team B"), ("APAC", "green", "Team C")]:
        for c in row:
            p.cell(110, 16, c, border=1)
        p.ln()
    p.ln(8); p.set_font("Times", "", 10.5)
    p.multi_cell(0, 14, LOREM2, align="J")
    p.output(out)
    return out


def l1_libreoffice():
    """Word-native dialect: build a DOCX, let LibreOffice render it to PDF."""
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    tmp = os.path.join(OUT, "_l1.docx")
    out = os.path.join(OUT, "l1_word_native.pdf")
    d = docx.Document()
    for s in d.sections:
        s.left_margin = s.right_margin = Inches(0.9)
    d.add_heading("Programme Status", 0)
    d.add_heading("Summary", 1)
    pr = d.add_paragraph(LOREM + LOREM2)
    pr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    d.add_heading("Risks", 1)
    for t in ("Index rebuild window is contended.",
              "Reranker training data is stale by two quarters.",
              "Cold-start latency is unmeasured in production."):
        d.add_paragraph(t, style="List Bullet")
    t = d.add_table(rows=4, cols=3); t.style = "Table Grid"
    for i, row in enumerate([("Risk", "Impact", "Owner"), ("Rebuild", "High", "A"),
                             ("Staleness", "Medium", "B"), ("Cold start", "Medium", "C")]):
        for j, c in enumerate(row):
            t.cell(i, j).text = c
    d.add_paragraph(LOREM2).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    d.save(tmp)
    prof = os.path.join(OUT, "_loprof2")
    subprocess.run([SOFFICE, "--headless", "--norestore",
                    "-env:UserInstallation=file:///" + prof.replace("\\", "/"),
                    "--convert-to", "pdf", "--outdir", OUT, tmp],
                   capture_output=True, timeout=300)
    src = os.path.join(OUT, "_l1.pdf")
    if os.path.exists(src):
        shutil.move(src, out)
    return out if os.path.exists(out) else None


if __name__ == "__main__":
    made = []
    for fn in (c1_whitepaper, c2_paper2col, c3_tables, c4_i18n, c5_graphics,
               c6_long, c7_code, c8_toc_links):
        h = fn()
        name = os.path.splitext(os.path.basename(h))[0]
        pdf = os.path.join(OUT, name + ".pdf")
        ok = chrome_pdf(h, pdf)
        print(("  OK " if ok else "FAIL ") + name)
        if ok:
            made.append(pdf)
    for fn in (r1_reportlab, f1_fpdf, l1_libreoffice):
        try:
            p = fn()
            if p and os.path.exists(p):
                made.append(p); print("  OK " + os.path.basename(p))
        except Exception as e:
            print("FAIL %s: %s" % (fn.__name__, e))
    print("\n%d PDFs" % len(made))
    import fitz
    for m in sorted(made):
        d = fitz.open(m)
        print("  %-28s %d pages  producer=%s" %
              (os.path.basename(m), d.page_count, d.metadata.get("producer")))
        d.close()
